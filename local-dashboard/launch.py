#!/usr/bin/env python3
"""
Akibwa Life — Local Server and Hosted Export

Launches the repo copy of public/life.html and wires it to private local data.

  python3 launch.py
  python3 launch.py --export-cloudflare ./cloudflare-dist
"""

import argparse
import http.server
import base64
import copy
import json
import os
import re
import shutil
import subprocess
import tempfile
import threading
import urllib.parse
import urllib.request
import webbrowser
import xml.etree.ElementTree as ET
import zipfile
from collections import Counter, defaultdict
from datetime import datetime, timedelta, timezone
from pathlib import Path
from zoneinfo import ZoneInfo

PORT = 8787
AI_CONTEXT_SYNC_INTERVAL_SECONDS = 45

_script_dir = os.path.dirname(os.path.abspath(__file__))
_repo_dir = os.path.dirname(_script_dir)
_public_dir = os.path.join(_repo_dir, 'public')
_dist_dir = os.path.join(_repo_dir, 'dist')
_home_dir = str(Path.home())
_docs_candidates = [
    os.environ.get('AKIBWA_DOCUMENTS_DIR'),
    os.path.join(_home_dir, 'Downloads', 'Akibwa Documents'),
    os.path.join(_home_dir, 'Downloads', 'Aibwa Documents'),
]
WORD_NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
JOURNAL_DATE_RE = re.compile(
    r'^(?P<date>\d{1,2}/\d{1,2}/\d{4})(?: (?P<time>\d{1,2}:\d{2})(?: [A-Z]{2,4})?)?:\s*(?P<body>.+)$'
)
JOURNAL_TZ = ZoneInfo('Europe/London')
TRANSCRIBE_CLI = Path(_home_dir) / '.codex' / 'skills' / 'transcribe' / 'scripts' / 'transcribe_diarize.py'
MAX_AUDIO_BYTES = 25 * 1024 * 1024
EMPLOYMENT_RANGE_RE = re.compile(
    r'^(?P<start_month>[A-Z][a-z]{2}) (?P<start_year>\d{4})\s*[–-]\s*(?:(?P<end_month>[A-Z][a-z]{2}) (?P<end_year>\d{4})|(?P<present>Present))$'
)
MARKET_DATA_USER_AGENT = (
    'Mozilla/5.0'
)
MARKET_CHART_HOSTS = (
    'https://query1.finance.yahoo.com',
    'https://query2.finance.yahoo.com',
)
MARKET_SYMBOLS = {
    'Alphabet (Class A)': {'symbol': 'GOOGL', 'label': 'Alphabet', 'currency': 'USD'},
    'Tesla': {'symbol': 'TSLA', 'label': 'Tesla', 'currency': 'USD'},
    'Amazon': {'symbol': 'AMZN', 'label': 'Amazon', 'currency': 'USD'},
    'iShares S&P 500 Info Tech (Acc)': {
        'symbol': 'IITU.L',
        'label': 'iShares S&P 500 Info Tech',
        'currency': 'GBp',
    },
    'TSMC': {'symbol': 'TSM', 'label': 'TSMC', 'currency': 'USD'},
}
MARKET_FX_SYMBOL = 'GBPUSD=X'
SHARES_RE = re.compile(r'(?P<quantity>\d+(?:\.\d+)?)\s*shares?', re.IGNORECASE)
AS_OF_DATE_RE = re.compile(r'As of (?P<date>\d{1,2} \w{3} \d{4})', re.IGNORECASE)
SOUL_SECTION_HEADINGS = {
    'What I want',
    'What matters to me',
    'Love',
    'Home',
    'Embodiment',
    'Beauty',
    'Truthfulness',
    'Compounding',
    'What I am becoming',
    'I am becoming a man who',
    'Non-negotiables',
    'The question underneath all of this',
}
MEMORY_CATEGORY_LABELS = {
    'working-on': 'Working on',
    'intent': 'Intent',
    'constraint': 'Constraint',
    'preference': 'Preference',
    'project': 'Project',
    'fact': 'Fact',
}
MEMORY_FOCUS_CATEGORIES = ('working-on', 'intent', 'project', 'constraint')
JOURNAL_PERSON_STOPWORDS = {
    'A', 'An', 'The', 'This', 'That', 'These', 'Those', 'I', 'My', 'Me', 'We', 'Our', 'It', 'He', 'She',
    'They', 'You', 'His', 'Her', 'Their', 'When', 'What', 'Why', 'How', 'Where', 'Who', 'Unknown', 'Day',
    'Dates', 'Date', 'Today', 'Tomorrow', 'Yesterday', 'Recovery', 'Notes', 'Note', 'Time', 'Reasons',
    'Goals', 'Summer', 'Work', 'Wellness', 'Diet', 'Dating', 'Exercise', 'Travel', 'Reflections', 'Solo',
    'Travelling', 'Public', 'Presentations', 'Running', 'Clubs', 'Club', 'Crew', 'Born', 'To', 'Run',
    'Human', 'Performance', 'Outliers', 'Peak', 'Games', 'Scoop', 'Podcast', 'Beyond', 'Croft', 'Avenue',
    'Tiger', 'Mountain', 'Life', 'Love', 'Home', 'Beauty', 'Truthfulness', 'Compounding', 'Loyalty',
    'Self', 'Worth', 'Witnessable', 'Improvement', 'Audio', 'Log', 'Interesting', 'Just', 'Another', 'Still',
    'Then', 'There', 'Here', 'Again', 'Later', 'Nothing', 'Something', 'Everything', 'Anywhere', 'Anyway',
    'Maybe', 'Great', 'Good', 'Current', 'Evening', 'Morning', 'Reflection', 'Social', 'Recap', 'Build',
    'Remember', 'Improve', 'Spend', 'Find', 'Do', 'For', 'At', 'In', 'On', 'Of', 'And', 'But', 'As', 'Not',
    'No', 'So', 'Last', 'First', 'Keep', 'While', 'Like', 'Had', 'Let', 'Woke', 'Ok', 'After', 'Upon',
    'Meeting', 'Man', 'Grandad', 'Granny', 'Grandparents', 'Family', 'Projects',
    'Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday',
    'January', 'February', 'March', 'April', 'May', 'June', 'July', 'August', 'September', 'October',
    'November', 'December',
}
JOURNAL_PERSON_STOP_ENTITIES = {
    'Greenwich Park', 'Canary Wharf', 'Battersea Park', 'Old Naval', 'St Peter', 'Lord Of', 'The Thames',
    'Shoreditch', 'San Francisco', 'New York', 'Leeds Festival', 'National Wealth Fund', 'Yorkshire Building Society',
    'Google Documents', 'Richmond Park', 'Cologne Cathedral', 'Old Cock', 'Black Horse', 'Castle Bank',
    'Vampire Weekend', 'Ram Dass', 'Jordan Peterson', 'Ben Shapiro', 'Sam Harris', 'Nick Cave', 'David Bowie',
    'James Clear', 'Alex Hormozi', 'Joe Rogan', 'Paul Tierney', 'Naval Ravikant', 'William Blake', 'Fiona Apple',
    'Alex Cameron', 'Rich Roll', 'Carl Jung', 'Abraham Maslow', 'Dan Sullivan', 'Humbert Humbert',
}
JOURNAL_PUBLIC_CONTEXT_HINTS = (
    'listened to', 'listening to', 'reading', 'watched', 'watching', 'podcast', 'album', 'song', 'film',
    'book', 'documentary', 'author', 'presentation by', 'quote', 'record', 'episode',
)
JOURNAL_RELATIONSHIP_HINTS = {
    'family': (
        ' dad', ' mum', ' brother', ' sister', ' grandparents', ' granny', ' grandad',
        ' aunt', ' cousin', ' family',
    ),
    'romantic': (
        'dating', 'dated ', 'the girl i\'m dating', 'the girl i’m dating', 'boyfriend', 'girlfriend',
        'kiss', 'falling for', 'what love feels like', 'romantic', 'sex', 'life partner', 'my love',
        'ask about dating', 'date with', 'date tomorrow with',
    ),
    'work': (
        'manager', 'colleague', 'office', 'team', 'interview', 'coworker', 'boss',
    ),
    'mentor': (
        'mentor', 'prompted by', 'opened my portal', 'learn from', 'inspiration', 'guidance',
    ),
    'friend': (
        'friend', 'friends', 'conversation', 'chat', 'spoke', 'speak', 'texted', 'bumped into',
        'went for a pint with', 'went out with', 'walking to work with', 'deep conversation with',
        'great to speak to', 'good chat with', 'opened up', 'bond', 'visit ', 'visited ',
    ),
    'community': ('climbing with', 'run with', 'goodgym', 'park run'),
}
JOURNAL_FAMILY_ENTITY_PATTERNS = (
    (re.compile(r'\b(?:my )?dad\b', re.IGNORECASE), 'Dad'),
    (re.compile(r'\bmum\b', re.IGNORECASE), 'Mum'),
    (re.compile(r'\bgrandparents\b', re.IGNORECASE), 'Grandparents'),
    (re.compile(r'\bgrandad\b', re.IGNORECASE), 'Grandad'),
    (re.compile(r'\bgranny\b', re.IGNORECASE), 'Granny'),
)
JOURNAL_RELATION_ROLE_PATTERNS = (
    (re.compile(r'\bmy manager (?P<name>[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\b'), {'work'}),
    (re.compile(r'\bmanager (?P<name>[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\b'), {'work'}),
    (re.compile(r'\bmy colleague (?P<name>[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\b'), {'work'}),
    (re.compile(r'\bcolleague (?P<name>[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\b'), {'work'}),
    (re.compile(r'\bmy boss (?P<name>[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\b'), {'work'}),
    (re.compile(r'\bmy cousin (?P<name>[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\b'), {'family'}),
    (re.compile(r'\bmy brother (?P<name>[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\b'), {'family'}),
    (re.compile(r'\bmy sister (?P<name>[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\b'), {'family'}),
    (re.compile(r'\bmy aunt (?P<name>[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\b'), {'family'}),
    (re.compile(r'\bmy friend (?P<name>[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\b'), {'friend'}),
    (re.compile(r'\bfriend (?P<name>[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\b'), {'friend'}),
    (re.compile(r'\bmentor (?P<name>[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\b'), {'mentor'}),
)
JOURNAL_DIRECT_NAME_PATTERNS = (
    re.compile(r'^(?P<name>[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\s*,'),
    re.compile(r'\b(?P<name>[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\s+(?:has|is|was|left|challenged|seemed|reminds?)\b'),
    re.compile(r'\bI found you (?P<name>[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\b'),
    re.compile(r'\bFinding (?P<name>[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\b'),
)
JOURNAL_SAFE_ALIAS_MAP = {
    'Anton': 'Anton Metro',
    'Chris': 'Chris Power',
    'Cecilia': 'Cecilia Bull',
    'Matilda': 'Matilde',
}
JOURNAL_CANONICAL_PERSON_ALIASES = {
    'dad': {
        'name': 'Ian Atkinson',
        'relation_to_dan': 'father',
        'aliases': ['Dad', 'Father'],
    },
    'father': {
        'name': 'Ian Atkinson',
        'relation_to_dan': 'father',
        'aliases': ['Dad', 'Father'],
    },
    'mum': {
        'name': 'Elizabeth Atkinson',
        'relation_to_dan': 'mother',
        'aliases': ['Mum', 'Mother'],
    },
    'mother': {
        'name': 'Elizabeth Atkinson',
        'relation_to_dan': 'mother',
        'aliases': ['Mum', 'Mother'],
    },
}
PERSON_PAGE_BLOCKLIST = {
    'blue dot',
    'one',
    'otherwise',
    'very',
    'sex',
    'loss',
    'everyone',
    'intellectually',
    'reason co',
    'regardless',
    'pixel watch',
    'visit spain',
    'tinder',
    'croatia',
}
PERSON_RELATION_TO_DAN_LABELS = {
    'father': 'father',
    'mother': 'mother',
    'brother': 'brother',
    'sister': 'sister',
    'uncle': 'uncle',
    'aunt': 'aunt',
    'grandfather': 'grandfather',
    'grandmother': 'grandmother',
    'cousin': 'cousin',
    'family member': 'family member',
    'romantic partner': 'romantic relationship',
    'mentor and friend': 'mentor and friend',
    'mentor': 'mentor',
    'work friend': 'work friend',
    'colleague': 'colleague',
    'friend': 'friend',
    'community connection': 'community connection',
    'acquaintance': 'acquaintance',
    'significant relationship': 'significant relationship',
}
_ai_context_sync_state = {
    'signature': None,
    'lastSyncedAt': '',
    'lastError': '',
}


def _docs_dir():
    for candidate in _docs_candidates:
        if candidate and os.path.isdir(candidate):
            return candidate
    return None


def _life_files(docs_dir):
    return {
        'inventoryWorkbook': os.path.join(docs_dir, 'Akibwa - Inventory.xlsx'),
        'financeWorkbook': os.path.join(docs_dir, 'Akibwa - Finances & Inventory.xlsx'),
        'soulDoc': os.path.join(docs_dir, 'Akibwa - Soul Doc.docx'),
        'journalDoc': os.path.join(docs_dir, 'Akibwa Journal - A Generous Slice.docx'),
        'creativeDoc': os.path.join(docs_dir, 'Akibwa - Creative Writing.docx'),
    }


def _safe_float(value):
    if value in (None, ''):
        return 0.0
    try:
        return float(value)
    except (TypeError, ValueError):
        return 0.0


def _try_load_workbook(path):
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise RuntimeError(
            'openpyxl is required for the life dashboard. '
            'Install it with `python3 -m pip install openpyxl`.'
        ) from exc
    return load_workbook(path, read_only=True, data_only=True)


def _try_import_docx():
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            'python-docx is required to append journal entries. '
            'Install it with `python3 -m pip install python-docx`.'
        ) from exc
    return Document


def _journal_now():
    return datetime.now(JOURNAL_TZ)


def _collapse_whitespace(value):
    return re.sub(r'\s+', ' ', value or '').strip()


def _split_content_paragraphs(value):
    normalized = (value or '').replace('\r\n', '\n').replace('\r', '\n').strip()
    if not normalized:
        return []
    chunks = re.split(r'\n\s*\n+', normalized)
    paragraphs = []
    for chunk in chunks:
        cleaned = chunk.strip()
        if not cleaned:
            continue
        lines = [line.strip() for line in cleaned.split('\n') if line.strip()]
        merged = ' '.join(lines).strip()
        if merged:
            paragraphs.append(merged)
    return paragraphs


def _save_docx_atomic(document, target_path):
    target = Path(target_path)
    temp_path = None
    try:
        with tempfile.NamedTemporaryFile(dir=target.parent, suffix='.docx', delete=False) as handle:
            temp_path = Path(handle.name)
        document.save(temp_path)
        os.replace(temp_path, target)
    finally:
        if temp_path and temp_path.exists():
            try:
                temp_path.unlink()
            except FileNotFoundError:
                pass


def _write_json_atomic(target_path, payload):
    target = Path(target_path)
    temp_path = None
    target.parent.mkdir(parents=True, exist_ok=True)
    try:
        with tempfile.NamedTemporaryFile(dir=target.parent, suffix='.json', delete=False, mode='w', encoding='utf-8') as handle:
            temp_path = Path(handle.name)
            json.dump(payload, handle, ensure_ascii=False, indent=2)
        os.replace(temp_path, target)
    finally:
        if temp_path and temp_path.exists():
            try:
                temp_path.unlink()
            except FileNotFoundError:
                pass


def _write_text_atomic(target_path, content):
    target = Path(target_path)
    temp_path = None
    target.parent.mkdir(parents=True, exist_ok=True)
    suffix = target.suffix or '.tmp'
    try:
        with tempfile.NamedTemporaryFile(dir=target.parent, suffix=suffix, delete=False, mode='w', encoding='utf-8') as handle:
            temp_path = Path(handle.name)
            handle.write(content)
        os.replace(temp_path, target)
    finally:
        if temp_path and temp_path.exists():
            try:
                temp_path.unlink()
            except FileNotFoundError:
                pass


def _life_history_path(docs_dir):
    return _resolve_state_path(docs_dir, 'life-history.json', '.dakibwa-life-history.json')


def _journal_capture_dir(docs_dir):
    return os.path.join(docs_dir, 'Journal Captures')


def _raw_dir(docs_dir):
    return os.path.join(docs_dir, 'Raw')


def _state_dir(docs_dir):
    return os.path.join(docs_dir, 'State')


def _extracted_dir(docs_dir):
    return os.path.join(docs_dir, 'Extracted')


def _resolve_state_path(docs_dir, filename, legacy_filename=''):
    canonical = os.path.join(_state_dir(docs_dir), filename)
    if os.path.exists(canonical):
        return canonical

    if legacy_filename:
        legacy = os.path.join(docs_dir, legacy_filename)
        if os.path.exists(legacy):
            os.makedirs(_state_dir(docs_dir), exist_ok=True)
            try:
                shutil.copy2(legacy, canonical)
            except OSError:
                return legacy
            if os.path.exists(canonical):
                return canonical
            return legacy

    return canonical


def _market_cache_path(docs_dir):
    return _resolve_state_path(docs_dir, 'market-cache.json', '.dakibwa-market-cache.json')


def _load_market_cache(docs_dir):
    path = _market_cache_path(docs_dir)
    try:
        with open(path, encoding='utf-8') as handle:
            payload = json.load(handle)
            if isinstance(payload, dict):
                raw_symbols = payload.get('symbols') if isinstance(payload.get('symbols'), dict) else {}
                normalized_symbols = {}
                for symbol, entry in raw_symbols.items():
                    if not isinstance(entry, dict):
                        continue
                    raw_points = entry.get('points') if isinstance(entry.get('points'), list) else []
                    points = []
                    for point in raw_points:
                        if not isinstance(point, dict):
                            continue
                        date = _collapse_whitespace(point.get('date', ''))
                        close = point.get('close')
                        if not date:
                            continue
                        try:
                            close_value = float(close)
                        except (TypeError, ValueError):
                            continue
                        points.append({
                            'date': date,
                            'close': close_value,
                        })
                    normalized_symbols[symbol] = {
                        'symbol': _collapse_whitespace(entry.get('symbol', '')) or symbol,
                        'currency': _collapse_whitespace(entry.get('currency', '')),
                        'fetchedAt': _collapse_whitespace(entry.get('fetchedAt', '')),
                        'points': points,
                    }
                return {
                    'date': _collapse_whitespace(payload.get('date', '')),
                    'updatedAt': _collapse_whitespace(payload.get('updatedAt', '')),
                    'symbols': normalized_symbols,
                }
    except FileNotFoundError:
        pass
    except json.JSONDecodeError:
        pass

    return {
        'date': '',
        'updatedAt': '',
        'symbols': {},
    }


def _market_chart_url(host, symbol, range_name='1y'):
    quoted = urllib.parse.quote(symbol, safe='')
    return f'{host}/v8/finance/chart/{quoted}?range={range_name}&interval=1d&includePrePost=false&events=div%2Csplits'


def _parse_market_chart_points(result):
    indicators = result.get('indicators') or {}
    timestamps = result.get('timestamp') or []
    quote = (indicators.get('quote') or [{}])[0] or {}
    adjclose = (indicators.get('adjclose') or [{}])[0] or {}
    close_values = adjclose.get('adjclose') or quote.get('close') or []
    fallback_values = quote.get('close') or []

    points = []
    for index, timestamp in enumerate(timestamps):
        price = None
        if index < len(close_values):
            price = close_values[index]
        if price in (None, '') and index < len(fallback_values):
            price = fallback_values[index]
        if price in (None, ''):
            continue
        try:
            close = float(price)
        except (TypeError, ValueError):
            continue
        date = datetime.fromtimestamp(int(timestamp), timezone.utc).date().isoformat()
        points.append({
            'date': date,
            'close': close,
        })
    return points


def _fetch_market_chart(symbol, range_name='1y'):
    last_error = None
    for host in MARKET_CHART_HOSTS:
        url = _market_chart_url(host, symbol, range_name=range_name)
        result = subprocess.run(
            [
                'curl',
                '-sS',
                '-L',
                '-H',
                f'User-Agent: {MARKET_DATA_USER_AGENT}',
                url,
            ],
            capture_output=True,
            text=True,
            timeout=20,
        )
        if result.returncode != 0:
            last_error = RuntimeError(result.stderr.strip() or f'curl failed for {symbol}')
            continue

        try:
            payload = json.loads(result.stdout)
        except json.JSONDecodeError:
            last_error = RuntimeError(f'Invalid market data returned for {symbol}.')
            continue

        chart = payload.get('chart') if isinstance(payload, dict) else None
        if not isinstance(chart, dict):
            last_error = RuntimeError(f'Unexpected market data shape for {symbol}.')
            continue

        error = chart.get('error')
        if error:
            description = error.get('description') if isinstance(error, dict) else str(error)
            last_error = RuntimeError(description or f'Market data failed for {symbol}.')
            continue

        result_items = chart.get('result') if isinstance(chart.get('result'), list) else []
        if not result_items:
            last_error = RuntimeError(f'No market history returned for {symbol}.')
            continue

        first = result_items[0]
        points = _parse_market_chart_points(first)
        if not points:
            last_error = RuntimeError(f'No usable market price points returned for {symbol}.')
            continue

        meta = first.get('meta') if isinstance(first.get('meta'), dict) else {}
        return {
            'symbol': symbol,
            'currency': _collapse_whitespace(meta.get('currency', '')),
            'fetchedAt': _journal_now().isoformat(timespec='seconds'),
            'points': points,
        }

    raise last_error or RuntimeError(f'Unable to fetch market data for {symbol}.')


def _refresh_market_cache(docs_dir, required_symbols):
    symbols_needed = [symbol for symbol in dict.fromkeys(required_symbols) if symbol]
    cache = _load_market_cache(docs_dir)
    if not symbols_needed:
        return cache

    today = _journal_now().date().isoformat()
    existing_symbols = dict(cache.get('symbols') or {})
    cache_is_fresh = cache.get('date') == today and all(
        existing_symbols.get(symbol, {}).get('points')
        for symbol in symbols_needed
    )
    if cache_is_fresh:
        return cache

    updated_symbols = dict(existing_symbols)
    fetched_any = False
    errors = []
    for symbol in symbols_needed:
        if cache.get('date') == today and updated_symbols.get(symbol, {}).get('points'):
            continue
        try:
            updated_symbols[symbol] = _fetch_market_chart(symbol)
            fetched_any = True
        except Exception as exc:
            errors.append(f'{symbol}: {exc}')

    if fetched_any:
        next_cache = {
            'date': today,
            'updatedAt': _journal_now().isoformat(timespec='seconds'),
            'symbols': updated_symbols,
        }
        _write_json_atomic(_market_cache_path(docs_dir), next_cache)
        if errors:
            next_cache['errors'] = errors
        return next_cache

    cache['errors'] = errors
    return cache


def _latest_point_on_or_before(points, target_date):
    latest = None
    for point in points:
        point_date = point.get('date', '')
        if not point_date or point_date > target_date:
            break
        latest = point
    return latest


def _price_to_gbp(price, currency, fx_rate=None):
    if price in (None, ''):
        return None
    value = float(price)
    if currency == 'USD':
        if not fx_rate:
            return None
        return value / float(fx_rate)
    if currency == 'GBp':
        return value / 100.0
    return value


def _parse_holding_quantity(notes):
    match = SHARES_RE.search(str(notes or ''))
    if not match:
        return None
    try:
        return float(match.group('quantity'))
    except (TypeError, ValueError):
        return None


def _parse_holding_as_of_date(notes):
    match = AS_OF_DATE_RE.search(str(notes or ''))
    if not match:
        return ''
    try:
        return datetime.strptime(match.group('date'), '%d %b %Y').date().isoformat()
    except ValueError:
        return ''


def _format_quantity(quantity):
    if quantity in (None, ''):
        return ''
    text = f'{float(quantity):.6f}'.rstrip('0').rstrip('.')
    return text or '0'


def _value_on_or_before(history, target_date):
    value = None
    for point in history:
        point_date = point.get('date', '')
        if not point_date or point_date > target_date:
            break
        value = point.get('valueGbp')
    return value


def _build_market_performance(docs_dir, money):
    holdings = []
    for account_label, rows in (
        ('ISA', money.get('isaHoldings') or []),
        ('Taxable', money.get('taxableHoldings') or []),
    ):
        for row in rows:
            item = str(row.get('item') or '')
            spec = MARKET_SYMBOLS.get(item)
            if not spec:
                continue
            holdings.append({
                'item': item,
                'account': account_label,
                'symbol': spec['symbol'],
                'label': spec['label'],
                'currency': spec['currency'],
                'notes': str(row.get('notes') or ''),
                'workbookValueGbp': round(_safe_float(row.get('value')), 2),
                'quantity': _parse_holding_quantity(row.get('notes', '')),
                'quantityInferred': False,
                'anchorDate': _parse_holding_as_of_date(row.get('notes', '')),
            })

    required_symbols = [holding['symbol'] for holding in holdings]
    if any(holding['currency'] == 'USD' for holding in holdings):
        required_symbols.append(MARKET_FX_SYMBOL)

    cache = _refresh_market_cache(docs_dir, required_symbols)
    symbol_cache = cache.get('symbols') or {}
    fx_points = list(symbol_cache.get(MARKET_FX_SYMBOL, {}).get('points') or [])

    for holding in holdings:
        if holding['quantity'] is not None:
            continue
        if not holding['anchorDate']:
            continue

        symbol_points = list(symbol_cache.get(holding['symbol'], {}).get('points') or [])
        anchor_point = _latest_point_on_or_before(symbol_points, holding['anchorDate'])
        if not anchor_point:
            continue

        fx_point = None
        if holding['currency'] == 'USD':
            fx_point = _latest_point_on_or_before(fx_points, holding['anchorDate'])
            if not fx_point:
                continue

        anchor_price_gbp = _price_to_gbp(
            anchor_point.get('close'),
            holding['currency'],
            fx_point.get('close') if fx_point else None,
        )
        if not anchor_price_gbp:
            continue

        holding['quantity'] = holding['workbookValueGbp'] / anchor_price_gbp
        holding['quantityInferred'] = True

    today = _journal_now().date().isoformat()
    all_dates = sorted({
        point.get('date')
        for holding in holdings
        for point in (symbol_cache.get(holding['symbol'], {}).get('points') or [])
        if point.get('date') and point.get('date') < today
    })

    positions = []
    for holding in holdings:
        quantity = holding.get('quantity')
        if not quantity:
            continue

        symbol_entry = symbol_cache.get(holding['symbol']) or {}
        symbol_points = list(symbol_entry.get('points') or [])
        if not symbol_points:
            continue

        price_by_date = {point['date']: point['close'] for point in symbol_points if point.get('date')}
        fx_by_date = {point['date']: point['close'] for point in fx_points if point.get('date')}

        last_price = None
        last_fx = None
        history = []
        for date in all_dates:
            if date in price_by_date:
                last_price = price_by_date[date]
            if holding['currency'] == 'USD' and date in fx_by_date:
                last_fx = fx_by_date[date]

            if last_price is None:
                continue

            price_gbp = _price_to_gbp(last_price, holding['currency'], last_fx)
            if price_gbp is None:
                continue

            history.append({
                'date': date,
                'valueGbp': round(quantity * price_gbp, 2),
                'priceGbp': round(price_gbp, 4),
            })

        if not history:
            continue

        latest = history[-1]
        previous = history[-2] if len(history) > 1 else None
        month_cutoff = (datetime.strptime(latest['date'], '%Y-%m-%d').date() - timedelta(days=30)).isoformat()
        month_value = _value_on_or_before(history, month_cutoff)
        latest_value = latest['valueGbp']
        previous_value = previous['valueGbp'] if previous else latest_value
        day_change = round(latest_value - previous_value, 2)
        day_pct = round((day_change / previous_value) * 100, 2) if previous_value else 0.0
        month_change = round(latest_value - month_value, 2) if month_value is not None else 0.0
        month_pct = round((month_change / month_value) * 100, 2) if month_value else 0.0

        latest_price_native = _latest_point_on_or_before(symbol_points, latest['date'])
        positions.append({
            'item': holding['item'],
            'label': holding['label'],
            'account': holding['account'],
            'symbol': holding['symbol'],
            'currency': holding['currency'],
            'quantity': round(quantity, 6),
            'quantityDisplay': _format_quantity(quantity),
            'quantityInferred': holding['quantityInferred'],
            'anchorDate': holding['anchorDate'],
            'workbookValueGbp': holding['workbookValueGbp'],
            'latestValueGbp': latest_value,
            'previousValueGbp': previous_value,
            'dayChangeGbp': day_change,
            'dayChangePct': day_pct,
            'monthChangeGbp': month_change,
            'monthChangePct': month_pct,
            'latestPriceNative': round(_safe_float((latest_price_native or {}).get('close')), 4),
            'latestPriceGbp': latest['priceGbp'],
            'notes': holding['notes'],
            'history': history,
        })

    positions.sort(key=lambda item: item.get('latestValueGbp', 0), reverse=True)

    portfolio_series = []
    for date in all_dates:
        total_market = 0.0
        count = 0
        for position in positions:
            value = _value_on_or_before(position.get('history') or [], date)
            if value is None:
                continue
            total_market += float(value)
            count += 1
        if count:
            portfolio_series.append({
                'date': date,
                'marketValueGbp': round(total_market, 2),
            })

    isa_cash = next(
        (_safe_float(item.get('value')) for item in (money.get('isaHoldings') or []) if 'cash' in str(item.get('item') or '').lower()),
        0.0,
    )
    for point in portfolio_series:
        point['totalWithCashGbp'] = round(point['marketValueGbp'] + isa_cash, 2)

    latest_series_point = portfolio_series[-1] if portfolio_series else None
    previous_series_point = portfolio_series[-2] if len(portfolio_series) > 1 else None
    latest_market = _safe_float((latest_series_point or {}).get('marketValueGbp'))
    previous_market = _safe_float((previous_series_point or {}).get('marketValueGbp'))
    latest_total = _safe_float((latest_series_point or {}).get('totalWithCashGbp'))
    previous_total = _safe_float((previous_series_point or {}).get('totalWithCashGbp'))

    market_day_change = round(latest_market - previous_market, 2) if latest_series_point else 0.0
    market_day_pct = round((market_day_change / previous_market) * 100, 2) if previous_market else 0.0
    total_day_change = round(latest_total - previous_total, 2) if latest_series_point else 0.0

    return {
        'ok': bool(portfolio_series and positions),
        'asOfDate': (latest_series_point or {}).get('date', ''),
        'marketValueGbp': round(latest_market, 2),
        'cashReserveGbp': round(isa_cash, 2),
        'totalWithCashGbp': round(latest_total, 2),
        'marketDayChangeGbp': market_day_change,
        'marketDayChangePct': market_day_pct,
        'totalDayChangeGbp': total_day_change,
        'positions': positions,
        'series': portfolio_series,
        'pointCount': len(portfolio_series),
        'cacheUpdatedAt': _collapse_whitespace(cache.get('updatedAt', '')),
        'cachePath': _market_cache_path(docs_dir),
        'priceSource': {
            'provider': 'Yahoo Finance chart data',
            'fxSymbol': MARKET_FX_SYMBOL,
        },
        'rangeNote': 'Current-position market value over time. This is not a cost-basis return chart.',
        'warnings': list(cache.get('errors') or []),
    }


def _memory_store_path(docs_dir):
    return _resolve_state_path(docs_dir, 'memory.json', '.dakibwa-memory.json')


def _scratchpad_store_path(docs_dir):
    return _resolve_state_path(docs_dir, 'scratchpad.json', '.dakibwa-ai-scratchpad.json')


def _load_life_history(docs_dir):
    path = _life_history_path(docs_dir)
    try:
        with open(path, encoding='utf-8') as handle:
            payload = json.load(handle)
            if isinstance(payload, dict):
                return payload
    except FileNotFoundError:
        pass
    except json.JSONDecodeError:
        pass
    return {'portfolio': []}


def _load_memory_store(docs_dir):
    path = _memory_store_path(docs_dir)
    try:
        with open(path, encoding='utf-8') as handle:
            payload = json.load(handle)
            if isinstance(payload, dict):
                entries = payload.get('entries')
                if isinstance(entries, list):
                    return {
                        'version': int(payload.get('version', 1) or 1),
                        'updatedAt': str(payload.get('updatedAt') or ''),
                        'entries': [entry for entry in entries if isinstance(entry, dict)],
                    }
    except FileNotFoundError:
        pass
    except json.JSONDecodeError:
        pass
    return {
        'version': 1,
        'updatedAt': '',
        'entries': [],
    }


def _normalize_string_list(values):
    if not isinstance(values, list):
        return []

    normalized = []
    seen = set()
    for value in values:
        cleaned = _collapse_whitespace(value)
        if not cleaned or cleaned in seen:
            continue
        seen.add(cleaned)
        normalized.append(cleaned)
    return normalized[:24]


def _normalize_recent_requests(values):
    if not isinstance(values, list):
        return []

    normalized = []
    for value in values:
        if not isinstance(value, dict):
            continue
        request = _collapse_whitespace(value.get('request', ''))
        if not request:
            continue
        normalized.append({
            'at': _collapse_whitespace(value.get('at', '')),
            'request': request,
            'status': _collapse_whitespace(value.get('status', '')) or 'noted',
            'note': _collapse_whitespace(value.get('note', '')),
        })
    return normalized[:20]


def _load_scratchpad_store(docs_dir):
    path = _scratchpad_store_path(docs_dir)
    try:
        with open(path, encoding='utf-8') as handle:
            payload = json.load(handle)
            if isinstance(payload, dict):
                return {
                    'version': int(payload.get('version', 1) or 1),
                    'updatedAt': _collapse_whitespace(payload.get('updatedAt', '')),
                    'currentIntent': _normalize_string_list(payload.get('currentIntent', [])),
                    'workingOn': _normalize_string_list(payload.get('workingOn', [])),
                    'constraints': _normalize_string_list(payload.get('constraints', [])),
                    'preferences': _normalize_string_list(payload.get('preferences', [])),
                    'recentRequests': _normalize_recent_requests(payload.get('recentRequests', [])),
                }
    except FileNotFoundError:
        pass
    except json.JSONDecodeError:
        pass

    return {
        'version': 1,
        'updatedAt': '',
        'currentIntent': [],
        'workingOn': [],
        'constraints': [],
        'preferences': [],
        'recentRequests': [],
    }


def _build_scratchpad_snapshot(docs_dir):
    store = _load_scratchpad_store(docs_dir)
    return {
        'storagePath': _scratchpad_store_path(docs_dir),
        'updatedAt': store.get('updatedAt', ''),
        'currentIntent': list(store.get('currentIntent', [])),
        'workingOn': list(store.get('workingOn', [])),
        'constraints': list(store.get('constraints', [])),
        'preferences': list(store.get('preferences', [])),
        'recentRequests': list(store.get('recentRequests', [])),
    }


def _profile_ingest_paths(docs_dir):
    raw_dir = Path(docs_dir) / 'Raw'
    if not raw_dir.exists():
        return []
    return sorted(raw_dir.glob('profile-ingest-*.json'))


def _normalize_profile_book(entry):
    if isinstance(entry, dict):
        title = _collapse_whitespace(entry.get('title', ''))
        author = _collapse_whitespace(entry.get('author', ''))
    else:
        raw = _collapse_whitespace(entry)
        if not raw:
            return None
        if ' - ' in raw:
            title, author = [part.strip() for part in raw.split(' - ', 1)]
        else:
            title, author = raw, ''
    if not title:
        return None
    return {
        'title': title,
        'author': author,
    }


def _normalize_profile_quote(entry):
    if isinstance(entry, dict):
        quote = _collapse_whitespace(entry.get('quote', ''))
        attribution = _collapse_whitespace(entry.get('attribution', ''))
    else:
        raw = _collapse_whitespace(entry)
        if not raw:
            return None
        if ' - ' in raw:
            quote, attribution = [part.strip() for part in raw.rsplit(' - ', 1)]
        else:
            quote, attribution = raw, ''
    if not quote:
        return None
    return {
        'quote': quote,
        'attribution': attribution,
    }


def _normalize_profile_film(entry):
    if isinstance(entry, dict):
        title = _collapse_whitespace(entry.get('title', ''))
        director = _collapse_whitespace(entry.get('director', ''))
    else:
        raw = _collapse_whitespace(entry)
        if not raw:
            return None
        if ' - ' in raw:
            title, director = [part.strip().strip('"') for part in raw.split(' - ', 1)]
        else:
            title, director = raw.strip('"'), ''
    if not title:
        return None
    return {
        'title': title,
        'director': director,
    }


def _normalize_profile_series(entry):
    title = _collapse_whitespace(entry.get('title', '')) if isinstance(entry, dict) else _collapse_whitespace(entry)
    if not title:
        return None
    return {'title': title}


def _normalize_profile_influence(entry):
    if isinstance(entry, dict):
        name = _collapse_whitespace(entry.get('name', ''))
        kind = _collapse_whitespace(entry.get('kind', '')) or 'public'
        relationship = _collapse_whitespace(entry.get('relationship', ''))
        why = _collapse_whitespace(entry.get('whyInfluential', ''))
        notes = _collapse_whitespace(entry.get('notes', ''))
    else:
        name = _collapse_whitespace(entry)
        kind = 'public'
        relationship = ''
        why = ''
        notes = ''
    if not name:
        return None
    return {
        'name': name,
        'kind': kind,
        'relationship': relationship,
        'whyInfluential': why,
        'notes': notes,
    }


def _normalize_intimate_history_entry(entry):
    if isinstance(entry, dict):
        name = _collapse_whitespace(entry.get('name', ''))
        notes = _collapse_whitespace(entry.get('notes', ''))
    else:
        raw = _collapse_whitespace(entry)
        if not raw:
            return None
        if '(' in raw and raw.endswith(')'):
            name, detail = raw.rsplit('(', 1)
            name = _collapse_whitespace(name)
            notes = _collapse_whitespace(detail[:-1])
        else:
            name = raw
            notes = ''
    if not name:
        return None
    return {
        'name': name,
        'notes': notes,
    }


def _normalize_profile_research(entry):
    if not isinstance(entry, dict):
        return None
    name = _collapse_whitespace(entry.get('name', ''))
    summary = _collapse_whitespace(entry.get('summary', ''))
    sources = []
    for source in entry.get('sources', []) if isinstance(entry.get('sources', []), list) else []:
        if not isinstance(source, dict):
            continue
        url = _collapse_whitespace(source.get('url', ''))
        label = _collapse_whitespace(source.get('label', ''))
        if not url:
            continue
        sources.append({
            'label': label or url,
            'url': url,
        })
    if not name:
        return None
    return {
        'name': name,
        'summary': summary,
        'sources': sources,
    }


def _load_profile_ingests(docs_dir):
    books = []
    quotes = []
    television = []
    films = []
    public_influences = []
    personal_influences = []
    intimate_history = []
    research_notes = []
    source_files = []
    updated_at = ''

    seen_books = set()
    seen_quotes = set()
    seen_tv = set()
    seen_films = set()
    seen_public = set()
    seen_personal = set()
    seen_intimate = set()
    seen_research = set()

    for path in _profile_ingest_paths(docs_dir):
        try:
            payload = json.loads(path.read_text(encoding='utf-8'))
        except (OSError, json.JSONDecodeError):
            continue
        if not isinstance(payload, dict):
            continue

        source_files.append(_relative_docs_path(docs_dir, str(path)))
        updated_candidate = _collapse_whitespace(payload.get('createdAt', '')) or datetime.fromtimestamp(path.stat().st_mtime).isoformat(timespec='seconds')
        if updated_candidate > updated_at:
            updated_at = updated_candidate

        for raw_book in payload.get('booksRead', []) if isinstance(payload.get('booksRead', []), list) else []:
            book = _normalize_profile_book(raw_book)
            if not book:
                continue
            key = f"{book['title'].lower()}::{book['author'].lower()}"
            if key in seen_books:
                continue
            seen_books.add(key)
            books.append(book)

        for raw_quote in payload.get('favoriteQuotes', []) if isinstance(payload.get('favoriteQuotes', []), list) else []:
            quote = _normalize_profile_quote(raw_quote)
            if not quote:
                continue
            key = quote['quote'].lower()
            if key in seen_quotes:
                continue
            seen_quotes.add(key)
            quotes.append(quote)

        for raw_show in payload.get('favoriteTelevision', []) if isinstance(payload.get('favoriteTelevision', []), list) else []:
            show = _normalize_profile_series(raw_show)
            if not show:
                continue
            key = show['title'].lower()
            if key in seen_tv:
                continue
            seen_tv.add(key)
            television.append(show)

        for raw_film in payload.get('favoriteFilms', []) if isinstance(payload.get('favoriteFilms', []), list) else []:
            film = _normalize_profile_film(raw_film)
            if not film:
                continue
            key = film['title'].lower()
            if key in seen_films:
                continue
            seen_films.add(key)
            films.append(film)

        for raw_person in payload.get('influentialPeople', []) if isinstance(payload.get('influentialPeople', []), list) else []:
            person = _normalize_profile_influence(raw_person)
            if not person:
                continue
            key = person['name'].lower()
            target = personal_influences if person.get('kind') == 'personal' else public_influences
            seen_target = seen_personal if person.get('kind') == 'personal' else seen_public
            if key in seen_target:
                continue
            seen_target.add(key)
            target.append(person)

        for raw_partner in payload.get('intimateRelationshipHistory', []) if isinstance(payload.get('intimateRelationshipHistory', []), list) else []:
            partner = _normalize_intimate_history_entry(raw_partner)
            if not partner:
                continue
            key = partner['name'].lower()
            if key in seen_intimate:
                continue
            seen_intimate.add(key)
            intimate_history.append(partner)

        for raw_research in payload.get('researchNotes', []) if isinstance(payload.get('researchNotes', []), list) else []:
            note = _normalize_profile_research(raw_research)
            if not note:
                continue
            key = note['name'].lower()
            if key in seen_research:
                continue
            seen_research.add(key)
            research_notes.append(note)

    return {
        'updatedAt': updated_at,
        'sourceFiles': source_files,
        'booksRead': books,
        'favoriteQuotes': quotes,
        'favoriteTelevision': television,
        'favoriteFilms': films,
        'publicInfluences': public_influences,
        'personalInfluences': personal_influences,
        'intimateRelationshipHistory': intimate_history,
        'researchNotes': research_notes,
        'summary': {
            'book_count': len(books),
            'quote_count': len(quotes),
            'television_count': len(television),
            'film_count': len(films),
            'public_influence_count': len(public_influences),
            'personal_influence_count': len(personal_influences),
            'intimate_relationship_count': len(intimate_history),
            'research_note_count': len(research_notes),
        },
    }


def _running_store_path(docs_dir):
    return _resolve_state_path(docs_dir, 'running.json', '.dakibwa-running.json')


def _normalize_running_shoe_entry(entry):
    if not isinstance(entry, dict):
        return None

    name = _collapse_whitespace(entry.get('name', ''))
    if not name:
        return None

    return {
        'name': name,
        'brand': _collapse_whitespace(entry.get('brand', '')),
        'role': _collapse_whitespace(entry.get('role', '')),
        'feel': _collapse_whitespace(entry.get('feel', '')),
        'verdict': _collapse_whitespace(entry.get('verdict', '')),
        'issue': _collapse_whitespace(entry.get('issue', '')),
        'note': _collapse_whitespace(entry.get('note', '')),
        'summary': _collapse_whitespace(entry.get('summary', '')),
        'bestFor': _collapse_whitespace(entry.get('bestFor', '')),
        'price': _collapse_whitespace(entry.get('price', '')),
        'weight': _collapse_whitespace(entry.get('weight', '')),
        'officialUrl': _collapse_whitespace(entry.get('officialUrl', '')),
        'isRecommended': bool(entry.get('isRecommended')),
        'why': _normalize_string_list(entry.get('why', [])),
        'tradeoffs': _normalize_string_list(entry.get('tradeoffs', [])),
        'colorways': _normalize_string_list(entry.get('colorways', [])),
        'scores': {
            'style': _collapse_whitespace((entry.get('scores') or {}).get('style', '')),
            'daily': _collapse_whitespace((entry.get('scores') or {}).get('daily', '')),
            'run': _collapse_whitespace((entry.get('scores') or {}).get('run', '')),
        },
    }


def _load_running_store(docs_dir):
    path = _running_store_path(docs_dir)
    default = {
        'version': 1,
        'updatedAt': '',
        'goal': '',
        'statusLine': '',
        'currentShoes': [],
        'requirements': [],
        'wardrobeAnchors': [],
        'paletteDirection': {
            'bestShort': '',
            'best': '',
            'note': '',
            'avoid': [],
        },
        'recommendation': {
            'pick': '',
            'summary': '',
            'reasonShort': '',
            'primaryColorway': '',
            'safeColorway': '',
            'secondaryPick': '',
            'buyIf': '',
            'colorNote': '',
        },
        'shortlist': [],
    }
    try:
        with open(path, encoding='utf-8') as handle:
            payload = json.load(handle)
            if not isinstance(payload, dict):
                return default

            shortlist = []
            for raw_entry in payload.get('shortlist', []):
                normalized = _normalize_running_shoe_entry(raw_entry)
                if normalized:
                    shortlist.append(normalized)

            current_shoes = []
            for raw_entry in payload.get('currentShoes', []):
                normalized = _normalize_running_shoe_entry(raw_entry)
                if normalized:
                    current_shoes.append(normalized)

            palette = payload.get('paletteDirection') if isinstance(payload.get('paletteDirection'), dict) else {}
            recommendation = payload.get('recommendation') if isinstance(payload.get('recommendation'), dict) else {}

            return {
                'version': int(payload.get('version', 1) or 1),
                'updatedAt': _collapse_whitespace(payload.get('updatedAt', '')),
                'goal': _collapse_whitespace(payload.get('goal', '')),
                'statusLine': _collapse_whitespace(payload.get('statusLine', '')),
                'currentShoes': current_shoes,
                'requirements': _normalize_string_list(payload.get('requirements', [])),
                'wardrobeAnchors': _normalize_string_list(payload.get('wardrobeAnchors', [])),
                'paletteDirection': {
                    'bestShort': _collapse_whitespace(palette.get('bestShort', '')),
                    'best': _collapse_whitespace(palette.get('best', '')),
                    'note': _collapse_whitespace(palette.get('note', '')),
                    'avoid': _normalize_string_list(palette.get('avoid', [])),
                },
                'recommendation': {
                    'pick': _collapse_whitespace(recommendation.get('pick', '')),
                    'summary': _collapse_whitespace(recommendation.get('summary', '')),
                    'reasonShort': _collapse_whitespace(recommendation.get('reasonShort', '')),
                    'primaryColorway': _collapse_whitespace(recommendation.get('primaryColorway', '')),
                    'safeColorway': _collapse_whitespace(recommendation.get('safeColorway', '')),
                    'secondaryPick': _collapse_whitespace(recommendation.get('secondaryPick', '')),
                    'buyIf': _collapse_whitespace(recommendation.get('buyIf', '')),
                    'colorNote': _collapse_whitespace(recommendation.get('colorNote', '')),
                },
                'shortlist': shortlist,
            }
    except FileNotFoundError:
        pass
    except json.JSONDecodeError:
        pass

    return default


def _build_running_snapshot(docs_dir):
    store = _load_running_store(docs_dir)
    store['storagePath'] = _running_store_path(docs_dir)
    return store


def _ai_context_dir(docs_dir):
    return os.path.join(docs_dir, 'AI Context')


def _private_context_dir(docs_dir):
    return os.path.join(_state_dir(docs_dir), 'Private Context')


def _json_markdown_block(payload):
    return '```json\n' + json.dumps(payload, ensure_ascii=False, indent=2) + '\n```'


def _markdown_scalar(value):
    if value is None:
        return 'null'
    if isinstance(value, bool):
        return 'true' if value else 'false'
    if isinstance(value, (int, float)):
        return str(value)
    return json.dumps(str(value), ensure_ascii=False)


def _markdown_frontmatter(payload):
    lines = ['---']
    for key, value in payload.items():
        if isinstance(value, list):
            lines.append(f'{key}:')
            for item in value:
                lines.append(f'  - {_markdown_scalar(item)}')
        else:
            lines.append(f'{key}: {_markdown_scalar(value)}')
    lines.append('---')
    return '\n'.join(lines)


def _relative_docs_path(docs_dir, path):
    if not path:
        return ''
    try:
        relative = os.path.relpath(path, docs_dir)
    except ValueError:
        relative = os.path.basename(path)
    return relative.replace(os.sep, '/')


def _relative_markdown_link(from_rel_path, to_rel_path, label=''):
    origin = os.path.dirname(from_rel_path) or '.'
    relative = os.path.relpath(to_rel_path, start=origin).replace(os.sep, '/')
    link_label = label or os.path.basename(to_rel_path)
    return f'[{link_label}]({relative})'


def _extracted_rel_path_for_article(article_rel_path):
    normalized = (article_rel_path or '').replace(os.sep, '/')
    root, filename = os.path.split(normalized)
    stem, _ = os.path.splitext(filename)
    if not root or not stem:
        return ''
    return f"{root.lower()}/{stem}.json"


def _ai_tab_metadata(snapshot, tab, title, kind='domain'):
    return {
        'kind': kind,
        'tab': tab,
        'title': title,
        'generated_at': snapshot.get('generatedAt', ''),
        'source_updated_at': snapshot.get('sourceUpdatedAt', ''),
        'deployment_mode': snapshot.get('deploymentMode', 'local'),
    }


def _knowledge_source_catalog(docs_dir):
    files = _life_files(docs_dir)
    records = [
        {
            'key': 'inventory_workbook',
            'title': 'Inventory workbook',
            'kind': 'workbook',
            'local_path': _relative_docs_path(docs_dir, files['inventoryWorkbook']),
            'purpose': 'Primary onebag and owned-item source for tracked weights, values, and categories.',
        },
        {
            'key': 'finance_workbook',
            'title': 'Finances & inventory workbook',
            'kind': 'workbook',
            'local_path': _relative_docs_path(docs_dir, files['financeWorkbook']),
            'purpose': 'Money, liabilities, income, employment history, and purchase watchlist source.',
        },
        {
            'key': 'soul_doc',
            'title': 'Soul doc',
            'kind': 'document',
            'local_path': _relative_docs_path(docs_dir, files['soulDoc']),
            'purpose': 'Values, north star, becoming, and non-negotiables source.',
        },
        {
            'key': 'journal_doc',
            'title': 'Journal doc',
            'kind': 'document',
            'local_path': _relative_docs_path(docs_dir, files['journalDoc']),
            'purpose': 'Recent journal entries and recurring themes source.',
        },
        {
            'key': 'creative_doc',
            'title': 'Creative writing doc',
            'kind': 'document',
            'local_path': _relative_docs_path(docs_dir, files['creativeDoc']),
            'purpose': 'Creative shelf titles, excerpts, and writing fragments source.',
        },
        {
            'key': 'profile_ingests',
            'title': 'Profile ingests',
            'kind': 'directory',
            'local_path': _relative_docs_path(docs_dir, os.path.join(docs_dir, 'Raw', 'profile-ingest-*.json')),
            'purpose': 'Structured personal taste, quote, book, media, and influence ingests captured from conversation.',
        },
        {
            'key': 'running_state',
            'title': 'Running state',
            'kind': 'state',
            'local_path': _relative_docs_path(docs_dir, _running_store_path(docs_dir)),
            'purpose': 'Private running-shoe decision state maintained from conversations and product research.',
        },
        {
            'key': 'memory_state',
            'title': 'Memory state',
            'kind': 'state',
            'local_path': _relative_docs_path(docs_dir, _memory_store_path(docs_dir)),
            'purpose': 'Durable machine memory entries that capture facts, projects, preferences, and intent.',
        },
        {
            'key': 'scratchpad_state',
            'title': 'Scratchpad state',
            'kind': 'state',
            'local_path': _relative_docs_path(docs_dir, _scratchpad_store_path(docs_dir)),
            'purpose': 'Short-horizon conversation-derived working context and recent implemented requests.',
        },
        {
            'key': 'market_cache',
            'title': 'Market cache',
            'kind': 'state',
            'local_path': _relative_docs_path(docs_dir, _market_cache_path(docs_dir)),
            'purpose': 'Cached Yahoo Finance closes and FX data used for the investments trend line.',
        },
        {
            'key': 'life_history',
            'title': 'Life history',
            'kind': 'state',
            'local_path': _relative_docs_path(docs_dir, _life_history_path(docs_dir)),
            'purpose': 'Private longitudinal snapshots of investments, cash, and onebag weight.',
        },
        {
            'key': 'journal_captures',
            'title': 'Journal captures',
            'kind': 'directory',
            'local_path': _relative_docs_path(docs_dir, _journal_capture_dir(docs_dir)),
            'purpose': 'Timestamped note-file captures created from the journal inbox flow.',
        },
    ]
    return {record['key']: record for record in records}


def _source_subset(source_catalog, source_keys):
    subset = []
    for key in source_keys:
        record = source_catalog.get(key)
        if record:
            subset.append(record)
    return subset


def _decorate_knowledge_markdown(
    snapshot,
    rel_path,
    title,
    kind,
    source_keys,
    related_paths,
    backlink_paths,
    body,
    source_catalog,
    extracted_rel_path='',
):
    header_payload = {
        'title': title,
        'kind': kind,
        'generated_at': snapshot.get('generatedAt', ''),
        'source_updated_at': snapshot.get('sourceUpdatedAt', ''),
        'deployment_mode': snapshot.get('deploymentMode', 'local'),
        'source_keys': source_keys,
        'related_files': related_paths,
        'backlinks': backlink_paths,
    }
    if extracted_rel_path:
        header_payload['extracted_record'] = extracted_rel_path
    header = _markdown_frontmatter(header_payload)

    lines = [
        header,
        '',
        body.strip(),
        '',
        '## source_documents',
        _json_markdown_block(_source_subset(source_catalog, source_keys)),
        '',
    ]

    if extracted_rel_path:
        lines.extend([
            '## extracted_record',
            f"- {_relative_markdown_link(rel_path, extracted_rel_path)}",
            '',
        ])

    if related_paths:
        lines.extend([
            '## related_files',
            '\n'.join(f'- {_relative_markdown_link(rel_path, target)}' for target in related_paths),
            '',
        ])

    if backlink_paths:
        lines.extend([
            '## backlinks',
            '\n'.join(f'- {_relative_markdown_link(rel_path, target)}' for target in backlink_paths),
            '',
        ])

    return '\n'.join(lines)


def _fashion_group_for_category(category):
    value = _collapse_whitespace(category).lower()
    if value in ('t-shirt', 'tank top', 'polo shirt', 'base layer', 'overshirt'):
        return 'Tops'
    if value in ('trousers', 'shorts', 'swim shorts', 'jeans', 'underwear'):
        return 'Bottoms'
    if value in ('rain jacket', 'down jacket'):
        return 'Outerwear'
    if value in ('trail shoes', 'running shoes', 'socks'):
        return 'Footwear'
    if value in ('belt', 'watch', 'glasses'):
        return 'Accessories'
    return ''


def _build_fashion_context(onebag):
    fashion_items = []
    grouped = defaultdict(list)
    totals = defaultdict(lambda: {'count': 0, 'value_gbp': 0.0, 'weight_g': 0.0})

    for item in onebag.get('items', []):
        group = _fashion_group_for_category(item.get('category', ''))
        if not group:
            continue
        payload = {
            'group': group,
            'category': item.get('category', ''),
            'item': item.get('item', ''),
            'price_gbp': round(float(item.get('priceGbp', 0) or 0), 2),
            'weight_g': round(float(item.get('weightG', 0) or 0), 1),
            'status': item.get('status', ''),
        }
        fashion_items.append(payload)
        grouped[group].append(payload)
        totals[group]['count'] += 1
        totals[group]['value_gbp'] += payload['price_gbp']
        totals[group]['weight_g'] += payload['weight_g']

    group_order = ['Tops', 'Bottoms', 'Outerwear', 'Footwear', 'Accessories']
    categories = []
    for group in group_order:
        total = totals.get(group)
        if not total or not total['count']:
            continue
        categories.append({
            'group': group,
            'count': total['count'],
            'value_gbp': round(total['value_gbp'], 2),
            'weight_g': round(total['weight_g'], 1),
        })

    fashion_items.sort(key=lambda item: (group_order.index(item['group']), item['item']))
    highlights = sorted(fashion_items, key=lambda item: item['price_gbp'], reverse=True)[:10]

    return {
        'summary': {
            'item_count': len(fashion_items),
            'group_count': len(categories),
            'total_value_gbp': round(sum(item['price_gbp'] for item in fashion_items), 2),
            'total_weight_g': round(sum(item['weight_g'] for item in fashion_items), 1),
        },
        'groups': categories,
        'items': fashion_items,
        'highlights': highlights,
    }


def _overview_markdown(snapshot, file_map):
    profile = snapshot.get('profile', {})
    work = snapshot.get('work', {})
    onebag = snapshot.get('onebag', {})
    money = snapshot.get('money', {})
    running = snapshot.get('running', {})
    market = snapshot.get('marketPerformance', {})
    people = snapshot.get('people', {})
    profile_library = snapshot.get('profileLibrary', {})
    signals = snapshot.get('signals', [])

    summary = {
        'full_name': profile.get('fullName', ''),
        'current_location': profile.get('currentLiving', ''),
        'current_role': work.get('jobTitle') or work.get('linkedinTitle', ''),
        'employer': work.get('employer', ''),
        'onebag_weight_kg': onebag.get('trackedWeightKg', 0),
        'running_recommendation': (running.get('recommendation') or {}).get('pick', ''),
        'people_count': (people.get('summary') or {}).get('person_count', 0),
        'books_count': (profile_library.get('summary') or {}).get('book_count', 0),
        'cash_total_gbp': money.get('cashTotal', 0),
        'investments_total_gbp': money.get('investmentsTotal', 0),
        'market_value_gbp': market.get('marketValueGbp', 0),
    }

    lines = [
        '# Overview',
        '',
        'AI-oriented snapshot of the current Akibwa Life surface.',
        '',
        '## metadata',
        _json_markdown_block(_ai_tab_metadata(snapshot, 'overview', 'Overview')),
        '',
        '## summary',
        _json_markdown_block(summary),
        '',
        '## signals',
        _json_markdown_block(signals),
        '',
        '## tab_files',
        _json_markdown_block(file_map),
        '',
    ]
    return '\n'.join(lines)


def _onebag_markdown(snapshot):
    onebag = snapshot.get('onebag', {})
    summary = {
        'tracked_item_count': onebag.get('trackedItemCount', 0),
        'tracked_value_gbp': onebag.get('trackedValueGbp', 0),
        'tracked_weight_g': onebag.get('trackedWeightG', 0),
        'tracked_weight_kg': onebag.get('trackedWeightKg', 0),
        'target_weight_kg': onebag.get('targetWeightKg'),
        'weight_gap_kg': onebag.get('weightGapKg'),
        'backpack': onebag.get('backpack'),
    }

    lines = [
        '# OneBag',
        '',
        '## metadata',
        _json_markdown_block(_ai_tab_metadata(snapshot, 'onebag', 'OneBag')),
        '',
        '## summary',
        _json_markdown_block(summary),
        '',
        '## category_breakdown',
        _json_markdown_block(onebag.get('categoryBreakdown', [])),
        '',
        '## heaviest_items',
        _json_markdown_block(onebag.get('heaviestItems', [])),
        '',
        '## most_expensive_items',
        _json_markdown_block(onebag.get('mostExpensiveItems', [])),
        '',
        '## on_order',
        _json_markdown_block(onebag.get('onOrder', [])),
        '',
        '## watchlist',
        _json_markdown_block(onebag.get('watchlist', [])),
        '',
    ]
    return '\n'.join(lines)


def _fashion_markdown(snapshot):
    fashion = _build_fashion_context(snapshot.get('onebag', {}))
    lines = [
        '# Fashion',
        '',
        '## metadata',
        _json_markdown_block(_ai_tab_metadata(snapshot, 'fashion', 'Fashion')),
        '',
        '## summary',
        _json_markdown_block(fashion.get('summary', {})),
        '',
        '## groups',
        _json_markdown_block(fashion.get('groups', [])),
        '',
        '## highlights',
        _json_markdown_block(fashion.get('highlights', [])),
        '',
        '## items',
        _json_markdown_block(fashion.get('items', [])),
        '',
    ]
    return '\n'.join(lines)


def _running_markdown(snapshot):
    running = snapshot.get('running', {})
    summary = {
        'goal': running.get('goal', ''),
        'status_line': running.get('statusLine', ''),
        'recommendation': running.get('recommendation', {}),
        'palette_direction': running.get('paletteDirection', {}),
    }

    lines = [
        '# Running',
        '',
        '## metadata',
        _json_markdown_block(_ai_tab_metadata(snapshot, 'running', 'Running')),
        '',
        '## summary',
        _json_markdown_block(summary),
        '',
        '## current_shoes',
        _json_markdown_block(running.get('currentShoes', [])),
        '',
        '## requirements',
        _json_markdown_block(running.get('requirements', [])),
        '',
        '## wardrobe_anchors',
        _json_markdown_block(running.get('wardrobeAnchors', [])),
        '',
        '## shortlist',
        _json_markdown_block(running.get('shortlist', [])),
        '',
    ]
    return '\n'.join(lines)


def _finances_markdown(snapshot):
    money = snapshot.get('money', {})
    summary = {
        'cash_total_gbp': money.get('cashTotal', 0),
        'isa_total_gbp': money.get('isaTotal', 0),
        'taxable_invest_total_gbp': money.get('taxableInvestTotal', 0),
        'investments_total_gbp': money.get('investmentsTotal', 0),
        'monthly_outgoings_total_gbp': money.get('monthlyOutgoingsTotal', 0),
        'liabilities_gbp_total': money.get('liabilitiesGbpTotal', 0),
        'medical_bill_usd': money.get('medicalBillUsd', 0),
        'medical_bill_notes': money.get('medicalBillNotes', ''),
    }

    lines = [
        '# Finances',
        '',
        '## metadata',
        _json_markdown_block(_ai_tab_metadata(snapshot, 'finances', 'Finances')),
        '',
        '## summary',
        _json_markdown_block(summary),
        '',
        '## cash',
        _json_markdown_block(money.get('cash', [])),
        '',
        '## liabilities',
        _json_markdown_block(money.get('liabilities', [])),
        '',
        '## monthly_outgoings',
        _json_markdown_block(money.get('monthlyOutgoings', [])),
        '',
        '## isa_holdings',
        _json_markdown_block(money.get('isaHoldings', [])),
        '',
        '## taxable_holdings',
        _json_markdown_block(money.get('taxableHoldings', [])),
        '',
    ]
    return '\n'.join(lines)


def _investments_markdown(snapshot):
    market = snapshot.get('marketPerformance', {})
    positions = []
    for position in market.get('positions', []):
        positions.append({
            'symbol': position.get('symbol', ''),
            'label': position.get('label', ''),
            'account': position.get('account', ''),
            'quantity_display': position.get('quantityDisplay', ''),
            'latest_value_gbp': position.get('latestValueGbp', 0),
            'latest_price_native': position.get('latestPriceNative', 0),
            'currency': position.get('currency', ''),
            'day_change_gbp': position.get('dayChangeGbp', 0),
            'day_change_pct': position.get('dayChangePct', 0),
            'month_change_gbp': position.get('monthChangeGbp', 0),
            'month_change_pct': position.get('monthChangePct', 0),
            'quantity_inferred': position.get('quantityInferred', False),
            'anchor_date': position.get('anchorDate', ''),
        })

    summary = {
        'ok': market.get('ok', False),
        'as_of_date': market.get('asOfDate', ''),
        'market_value_gbp': market.get('marketValueGbp', 0),
        'cash_reserve_gbp': market.get('cashReserveGbp', 0),
        'total_with_cash_gbp': market.get('totalWithCashGbp', 0),
        'market_day_change_gbp': market.get('marketDayChangeGbp', 0),
        'market_day_change_pct': market.get('marketDayChangePct', 0),
        'range_note': market.get('rangeNote', ''),
        'price_source': market.get('priceSource', {}),
    }

    lines = [
        '# Investments',
        '',
        '## metadata',
        _json_markdown_block(_ai_tab_metadata(snapshot, 'investments', 'Investments')),
        '',
        '## summary',
        _json_markdown_block(summary),
        '',
        '## positions',
        _json_markdown_block(positions),
        '',
        '## recent_private_history',
        _json_markdown_block(snapshot.get('portfolioHistory', [])[-24:]),
        '',
        '## warnings',
        _json_markdown_block(market.get('warnings', [])),
        '',
    ]
    return '\n'.join(lines)


def _creative_markdown(snapshot):
    creative = snapshot.get('creative', {})
    entries = []
    for entry in creative.get('entries', []):
        paragraphs = entry.get('paragraphs', []) or []
        entries.append({
            'title': entry.get('title', ''),
            'excerpt': entry.get('excerpt', ''),
            'paragraph_count': len(paragraphs),
            'preview_paragraphs': paragraphs[:3],
        })

    lines = [
        '# Creative',
        '',
        '## metadata',
        _json_markdown_block(_ai_tab_metadata(snapshot, 'creative', 'Creative')),
        '',
        '## titles',
        _json_markdown_block(creative.get('titles', [])),
        '',
        '## entries',
        _json_markdown_block(entries),
        '',
        '## sample_lines',
        _json_markdown_block(creative.get('sampleLines', [])),
        '',
    ]
    return '\n'.join(lines)


def _journal_markdown(snapshot):
    journal = snapshot.get('journal', {})
    entries = []
    for entry in journal.get('latestEntries', []):
        entries.append({
            'date': entry.get('date', ''),
            'time': entry.get('time', ''),
            'title': entry.get('title', ''),
            'text': entry.get('text', ''),
        })

    lines = [
        '# Journal',
        '',
        '## metadata',
        _json_markdown_block(_ai_tab_metadata(snapshot, 'journal', 'Journal')),
        '',
        '## themes',
        _json_markdown_block(journal.get('themes', [])),
        '',
        '## people_summary',
        _json_markdown_block((journal.get('people') or {}).get('summary', {})),
        '',
        '## latest_entries',
        _json_markdown_block(entries),
        '',
    ]
    return '\n'.join(lines)


def _people_markdown(snapshot):
    people = snapshot.get('people', {})
    lines = [
        '# People',
        '',
        '## metadata',
        _json_markdown_block(_ai_tab_metadata(snapshot, 'people', 'People')),
        '',
        '## summary',
        _json_markdown_block(people.get('summary', {})),
        '',
        '## featured',
        _json_markdown_block(people.get('featured', [])),
        '',
        '## directory',
        _json_markdown_block(people.get('directory', [])),
        '',
    ]
    return '\n'.join(lines)


def _profile_markdown(snapshot):
    profile_library = snapshot.get('profileLibrary', {})
    lines = [
        '# Profile',
        '',
        '## metadata',
        _json_markdown_block(_ai_tab_metadata(snapshot, 'profile', 'Profile')),
        '',
        '## summary',
        _json_markdown_block(profile_library.get('summary', {})),
        '',
        '## books_read',
        _json_markdown_block(profile_library.get('booksRead', [])),
        '',
        '## favorite_quotes',
        _json_markdown_block(profile_library.get('favoriteQuotes', [])),
        '',
        '## public_influences',
        _json_markdown_block(profile_library.get('publicInfluences', [])),
        '',
        '## personal_influences',
        _json_markdown_block(profile_library.get('personalInfluences', [])),
        '',
        '## favorite_television',
        _json_markdown_block(profile_library.get('favoriteTelevision', [])),
        '',
        '## favorite_films',
        _json_markdown_block(profile_library.get('favoriteFilms', [])),
        '',
        '## research_notes',
        _json_markdown_block(profile_library.get('researchNotes', [])),
        '',
    ]
    return '\n'.join(lines)


def _person_markdown(snapshot, rel_path, person, source_catalog, extracted_rel_path=''):
    bundle = _person_context_bundle(snapshot, person, rel_path)
    source_keys = bundle['source_keys']
    personal_match = bundle['personal_match']
    research_match = bundle['research_match']
    evidence = bundle['evidence']
    role_notes = bundle['role_notes']
    related_paths = bundle['related_paths']

    body = [
        f"# {person.get('name', 'Unknown person')}",
        '',
        person.get('articleLead', person.get('summary', '')),
        '',
        '## role_in_life',
        '\n\n'.join(dict.fromkeys(note for note in role_notes if note)),
        '',
        '## relationship_to_dan',
        _bullet_markdown(person.get('relationSummary', [])),
        '',
        '## known_facts',
        _bullet_markdown(person.get('facts', [])),
        '',
        '## current_context',
        _bullet_markdown(person.get('contextNotes', [])),
        '',
        '## source_notes',
        _bullet_markdown([
            f"{item.get('date', '')}: {item.get('snippet', '')}".strip(': ')
            for item in evidence[:8]
        ]),
        '',
    ]
    if personal_match:
        body.extend([
            '## profile_context',
            _bullet_markdown([
                f"Relationship signal: {_person_relation_label(personal_match.get('relationship'))}." if personal_match.get('relationship') else '',
                personal_match.get('whyInfluential', ''),
                personal_match.get('notes', ''),
            ]),
            '',
        ])
    if research_match:
        body.extend([
            '## public_research',
            _bullet_markdown([
                research_match.get('summary', ''),
                *[
                    f"{item.get('label', 'Source')}: {item.get('url', '')}"
                    for item in research_match.get('sources', []) or []
                ],
            ]),
            '',
        ])
    body.extend([
        '## related_pages',
        _bullet_markdown([
            _relative_markdown_link(rel_path, path, os.path.splitext(os.path.basename(path))[0].replace('-', ' ').title())
            for path in related_paths
        ]),
        '',
    ])
    return _decorate_knowledge_markdown(
        snapshot,
        rel_path,
        person.get('name', 'Unknown person'),
        'person',
        source_keys,
        related_paths,
        ['Indexes/people-index.md'],
        '\n'.join(body),
        source_catalog,
        extracted_rel_path=extracted_rel_path,
    )


def _employment_markdown(snapshot):
    work = snapshot.get('work', {})
    current_role = ''
    current_employer = work.get('employer', '')
    for item in work.get('history', []):
        if item.get('isCurrent'):
            current_role = item.get('role', '')
            current_employer = item.get('employer', current_employer)
            break

    summary = {
        'current_role': current_role or work.get('jobTitle') or work.get('linkedinTitle', ''),
        'current_employer': current_employer,
        'office_address': work.get('officeAddress', ''),
        'working_pattern': work.get('workingPattern', ''),
        'annual_holiday': work.get('annualHoliday', ''),
        'monthly_gross_salary': work.get('monthlyGrossSalary', 0),
        'monthly_net_pay': work.get('monthlyNetPay', 0),
        'annual_gross_salary': work.get('annualGrossSalary', 0),
        'annual_bonus_gross': work.get('annualBonusGross', 0),
        'total_annual_comp_gross': work.get('totalAnnualCompGross', 0),
        'employer_pension': work.get('employerPension', 0),
        'start_date_contract': work.get('startDateContract', ''),
        'start_date_linkedin': work.get('StartDateLinkedIn', work.get('startDateLinkedIn', '')),
    }

    lines = [
        '# Employment',
        '',
        '## metadata',
        _json_markdown_block(_ai_tab_metadata(snapshot, 'employment', 'Employment')),
        '',
        '## summary',
        _json_markdown_block(summary),
        '',
        '## recent_payslips',
        _json_markdown_block(work.get('recentPayslips', [])),
        '',
        '## history',
        _json_markdown_block(work.get('history', [])),
        '',
    ]
    return '\n'.join(lines)


def _soul_markdown(snapshot):
    soul = snapshot.get('soul', {})
    summary = {
        'title': soul.get('title', ''),
        'north_star': soul.get('northStar', ''),
        'question': soul.get('question', ''),
    }

    lines = [
        '# Soul',
        '',
        '## metadata',
        _json_markdown_block(_ai_tab_metadata(snapshot, 'soul', 'Soul')),
        '',
        '## summary',
        _json_markdown_block(summary),
        '',
        '## opening',
        _json_markdown_block(soul.get('opening', [])),
        '',
        '## values',
        _json_markdown_block(soul.get('values', [])),
        '',
        '## becoming',
        _json_markdown_block(soul.get('becoming', [])),
        '',
        '## non_negotiables',
        _json_markdown_block(soul.get('nonNegotiables', [])),
        '',
    ]
    return '\n'.join(lines)


def _bullet_markdown(items):
    cleaned = [_collapse_whitespace(item) for item in items if _collapse_whitespace(item)]
    if not cleaned:
        return '- None yet'
    return '\n'.join(f'- {item}' for item in cleaned)


def _find_known_person_record(name, directory):
    target = _collapse_whitespace(name).lower()
    if not target:
        return None

    exact = [person for person in directory if _collapse_whitespace(person.get('name', '')).lower() == target]
    if exact:
        return exact[0]

    first_token = target.split(' ', 1)[0]
    if not first_token:
        return None

    first_matches = [
        person
        for person in directory
        if _collapse_whitespace(person.get('name', '')).lower() == first_token
    ]
    if len(first_matches) == 1:
        return first_matches[0]

    return None


def _build_private_context_markdown(snapshot):
    docs_dir = snapshot.get('documentsDir', '')
    if not docs_dir:
        return {'directory': '', 'files': []}

    target_dir = _private_context_dir(docs_dir)
    intimate_history = snapshot.get('profileLibrary', {}).get('intimateRelationshipHistory', []) or []
    journal_people = snapshot.get('people', {}).get('directory', []) or []

    if os.path.exists(target_dir):
        shutil.rmtree(target_dir)
    os.makedirs(target_dir, exist_ok=True)

    content_map = {}
    index_entries = []

    for entry in intimate_history:
        name = _collapse_whitespace(entry.get('name', ''))
        if not name:
            continue
        slug = _slugify(name)
        rel_path = f'People/{slug}.md'
        known_person = _find_known_person_record(name, journal_people)
        linked_people_path = ''
        if known_person:
            linked_people_path = f"AI Context/People/{known_person.get('slug', _slugify(known_person.get('name', name)))}.md"

        body = [
            f'# {name}',
            '',
            '## private_relationship_record',
            _json_markdown_block({
                'name': name,
                'notes': entry.get('notes', ''),
                'source': 'Raw/profile-ingest-*-intimate-history.json',
                'visibility': 'private-agent-only',
            }),
            '',
        ]

        if known_person:
            body.extend([
                '## linked_people_layer',
                _json_markdown_block({
                    'name': known_person.get('name', ''),
                    'primary_relationship': known_person.get('primaryRelationship', ''),
                    'mentions': known_person.get('mentions', 0),
                    'summary': known_person.get('summary', ''),
                    'facts': known_person.get('facts', [])[:3],
                    'ai_context_path': linked_people_path,
                }),
                '',
            ])

        notes = _collapse_whitespace(entry.get('notes', ''))
        body.extend([
            '## current_notes',
            _bullet_markdown([notes] if notes else []),
            '',
            '## handling',
            _bullet_markdown([
                'Keep this local-only unless the user explicitly asks for it to appear in a visible UI or hosted export.',
                'Treat notes here as user-supplied private relationship context, not public research.',
                'When useful, reconcile names here with richer people pages or journal evidence rather than duplicating conflicting claims.',
            ]),
            '',
        ])

        content_map[rel_path] = '\n'.join(body)
        index_entries.append({
            'name': name,
            'path': rel_path,
            'linked_people_page': linked_people_path,
            'has_journal_record': bool(known_person),
            'notes': notes,
        })

    content_map['README.md'] = '\n'.join([
        '# Private Context',
        '',
        'Highly sensitive local-only agent context lives here.',
        'Do not surface this material in the visible site or hosted export unless the user explicitly asks for that.',
        '',
        '## scope',
        _bullet_markdown([
            'Private relationship and people notes that are too sensitive for the normal compiled wiki.',
            'Context that future local agents may need in order to reason well about Dan’s life without making it part of the visible product surface.',
        ]),
        '',
        '## files',
        _json_markdown_block(index_entries),
        '',
    ])

    content_map['relationship-history.md'] = '\n'.join([
        '# Relationship History',
        '',
        'Local-only intimate relationship register compiled from direct user-provided context.',
        '',
        '## summary',
        _json_markdown_block({
            'entry_count': len(index_entries),
            'visibility': 'private-agent-only',
            'linked_journal_records': sum(1 for item in index_entries if item.get('has_journal_record')),
        }),
        '',
        '## entries',
        _json_markdown_block(index_entries),
        '',
        '## next_moves',
        _bullet_markdown([
            'Use this to enrich local reasoning about relationships and personal history.',
            'Prefer linking these records to richer people pages when the journal or later sources provide evidence.',
            'Keep future ingests incremental so new private notes update these files rather than splintering into separate ad hoc records.',
        ]),
        '',
    ])

    for relative_path, content in content_map.items():
        _write_text_atomic(os.path.join(target_dir, relative_path), content)

    return {
        'directory': target_dir,
        'files': sorted(relative_path for relative_path in content_map),
    }


def _build_concept_specs(snapshot, scratchpad, memory, fashion):
    running = snapshot.get('running', {})
    money = snapshot.get('money', {})
    onebag = snapshot.get('onebag', {})
    market = snapshot.get('marketPerformance', {})
    work = snapshot.get('work', {})
    soul = snapshot.get('soul', {})
    journal = snapshot.get('journal', {})
    creative = snapshot.get('creative', {})
    profile_library = snapshot.get('profileLibrary', {})

    current_role = work.get('jobTitle') or work.get('linkedinTitle') or ''
    current_employer = work.get('employer', '')
    work_history = work.get('history', []) or []
    top_groups = [group['group'] for group in fashion.get('groups', [])[:3]]
    creative_titles = creative.get('titles', []) or []
    books = profile_library.get('booksRead', []) or []
    quotes = profile_library.get('favoriteQuotes', []) or []
    films = profile_library.get('favoriteFilms', []) or []
    television = profile_library.get('favoriteTelevision', []) or []
    public_influences = profile_library.get('publicInfluences', []) or []
    personal_influences = profile_library.get('personalInfluences', []) or []
    research_notes = profile_library.get('researchNotes', []) or []

    return {
        'capsule-wardrobe': {
            'title': 'Capsule wardrobe',
            'summary': 'The wardrobe project is about making the onebag clothing system coherent enough that travel, work, and everyday wear all pull from the same compact set of pieces.',
            'source_keys': ['inventory_workbook', 'finance_workbook', 'running_state', 'memory_state', 'scratchpad_state'],
            'related_domains': ['onebag', 'fashion', 'running'],
            'highlights': [
                f"Tracked carry is {onebag.get('weightGapKg', 0):+.2f} kg versus the current target." if onebag.get('weightGapKg') is not None else 'No onebag weight target is currently set.',
                f"Top wardrobe groups right now: {', '.join(top_groups)}." if top_groups else 'Wardrobe groups have not been surfaced yet.',
                f"Known wardrobe anchor: {(running.get('wardrobeAnchors') or [''])[0]}." if running.get('wardrobeAnchors') else '',
                f"Preferred palette direction: {(running.get('paletteDirection') or {}).get('bestShort', '')}." if (running.get('paletteDirection') or {}).get('bestShort') else '',
            ],
            'structured_state': {
                'fashion_summary': fashion.get('summary', {}),
                'wardrobe_anchors': running.get('wardrobeAnchors', []),
                'onebag_target_weight_kg': onebag.get('targetWeightKg'),
                'onebag_weight_gap_kg': onebag.get('weightGapKg'),
            },
            'next_questions': [
                'Which missing warm-weather pieces would improve the capsule most without pushing the bag further over target?',
                'Which existing pieces should count as the permanent capsule backbone versus optional seasonal rotation?',
                'Should the everyday running shoe deliberately act as part of the capsule rather than a separate specialist shoe?',
            ],
        },
        'one-shoe-strategy': {
            'title': 'One-shoe strategy',
            'summary': 'The active running-shoe decision is really a lifestyle-system choice: find a single shoe that works with the wardrobe, at work, and on quicker runs without feeling like a compromise everywhere.',
            'source_keys': ['running_state', 'inventory_workbook', 'memory_state', 'scratchpad_state'],
            'related_domains': ['running', 'fashion', 'onebag'],
            'highlights': [
                f"Current leading answer: {(running.get('recommendation') or {}).get('pick', 'No pick yet')}.",
                f"Current fallback: {(running.get('recommendation') or {}).get('secondaryPick', 'No fallback yet')}.",
                f"Current shoes in rotation: {', '.join(item.get('name', '') for item in running.get('currentShoes', []))}." if running.get('currentShoes') else 'No current shoes have been saved.',
                f"Color direction: {(running.get('paletteDirection') or {}).get('best', '')}." if (running.get('paletteDirection') or {}).get('best') else '',
            ],
            'structured_state': {
                'recommendation': running.get('recommendation', {}),
                'current_shoes': running.get('currentShoes', []),
                'shortlist': [item.get('name', '') for item in running.get('shortlist', [])],
                'requirements': running.get('requirements', []),
            },
            'next_questions': [
                'Which contender still feels best in person after style, comfort, and run-feel are all weighted together?',
                'Does the eventual shoe need a second seasonal counterpart, or can it truly be the single default shoe?',
                'Which colorway will age best with denim, olive, and off-white summer trousers?',
            ],
        },
        'financial-position': {
            'title': 'Financial position',
            'summary': 'The money picture currently balances decent invested assets and steady income against recurring outgoings, a live medical-bill issue, and a still-young historical record.',
            'source_keys': ['finance_workbook', 'market_cache', 'life_history'],
            'related_domains': ['finances', 'investments', 'employment'],
            'highlights': [
                f"Cash tracked: £{money.get('cashTotal', 0):,.2f}.",
                f"Investments tracked: £{money.get('investmentsTotal', 0):,.2f}.",
                f"Outstanding medical bill: ${money.get('medicalBillUsd', 0):,.2f}." if money.get('medicalBillUsd') else 'No medical bill recorded.',
                f"Latest live market value: £{market.get('marketValueGbp', 0):,.2f} as of {market.get('asOfDate', 'unknown date')}." if market.get('ok') else 'Live market tracking is still incomplete.',
            ],
            'structured_state': {
                'cash_total_gbp': money.get('cashTotal', 0),
                'monthly_outgoings_total_gbp': money.get('monthlyOutgoingsTotal', 0),
                'liabilities_gbp_total': money.get('liabilitiesGbpTotal', 0),
                'medical_bill_usd': money.get('medicalBillUsd', 0),
                'market_summary': {
                    'ok': market.get('ok', False),
                    'market_value_gbp': market.get('marketValueGbp', 0),
                    'total_with_cash_gbp': market.get('totalWithCashGbp', 0),
                    'point_count': market.get('pointCount', 0),
                },
                'private_history_points': len(snapshot.get('portfolioHistory', []) or []),
            },
            'next_questions': [
                'What is the right target cash reserve once the medical-bill uncertainty is resolved?',
                'Should the investment view start separating strategic long-term holdings from tactical or inferred positions?',
                'What timeline and process should clear the remaining medical-bill follow-up?',
            ],
        },
        'career-arc': {
            'title': 'Career arc',
            'summary': 'The employment story is moving toward a clearer career narrative: analytical depth, BI leadership, and a current senior role, but with some missing compensation history for older chapters.',
            'source_keys': ['finance_workbook', 'memory_state', 'scratchpad_state'],
            'related_domains': ['employment'],
            'highlights': [
                f"Current chapter: {current_role or 'Current role'} at {current_employer or 'current employer'}.",
                f"Tracked work history depth: {len(work_history)} roles.",
                f"Current total annual gross comp: £{work.get('totalAnnualCompGross', 0):,.2f}.",
                'Past-role salary history has not been captured yet.' if work_history else 'Employment history has not been surfaced yet.',
            ],
            'structured_state': {
                'current_role': current_role,
                'current_employer': current_employer,
                'working_pattern': work.get('workingPattern', ''),
                'history': work_history,
                'current_comp_gross': work.get('totalAnnualCompGross', 0),
            },
            'next_questions': [
                'What were the salary and compensation steps across each previous role?',
                'How should the career arc be framed: technical specialist, strategic BI lead, or broader operating builder?',
                'Which work themes from the journal belong in the longer-term career narrative?',
            ],
        },
        'north-star': {
            'title': 'North star',
            'summary': 'The soul layer is the orienting frame beneath the rest of Akibwa: values, becoming, and the questions that keep resurfacing across journal entries and current plans.',
            'source_keys': ['soul_doc', 'journal_doc', 'memory_state', 'scratchpad_state'],
            'related_domains': ['soul', 'journal', 'overview'],
            'highlights': [
                soul.get('northStar', ''),
                f"Recurring journal themes: {', '.join(journal.get('themes', [])[:4])}." if journal.get('themes') else '',
                soul.get('question', ''),
            ],
            'structured_state': {
                'north_star': soul.get('northStar', ''),
                'question': soul.get('question', ''),
                'values': [item.get('name', '') for item in soul.get('values', [])],
                'journal_themes': journal.get('themes', []),
            },
            'next_questions': [
                'Which current commitments genuinely compound toward the north star, and which ones are just noise?',
                'How should journal capture surface the deepest recurring questions rather than only recent notes?',
                'Which relationships, practices, and environments matter most for the next season?',
            ],
        },
        'creative-practice': {
            'title': 'Creative practice',
            'summary': 'The creative layer is less about publishing volume and more about keeping hold of fragments, titles, and recurring lines that should stay available for future synthesis.',
            'source_keys': ['creative_doc', 'journal_doc', 'scratchpad_state'],
            'related_domains': ['creative', 'journal'],
            'highlights': [
                f"Creative shelf size: {len(creative_titles)} surfaced pieces.",
                f"Current sample line: {(creative.get('sampleLines') or [''])[0]}" if creative.get('sampleLines') else '',
                'The shelf now works as a piece-by-piece browsing surface in the dashboard.',
            ],
            'structured_state': {
                'titles': creative_titles,
                'sample_lines': creative.get('sampleLines', []),
                'entry_count': len(creative.get('entries', []) or []),
            },
            'next_questions': [
                'Which pieces are seeds for larger essays, and which are better kept as fragments?',
                'What themes recur across the creative doc and the journal strongly enough to deserve their own concept pages?',
                'Which outputs should be filed back into the wiki as polished drafts or slide-ready briefings?',
            ],
        },
        'media-canon': {
            'title': 'Media canon',
            'summary': 'The books, films, and television named so far sketch a cultural canon that leans philosophical, comic, psychologically intense, and aesthetically literary rather than purely mainstream.',
            'source_keys': ['profile_ingests'],
            'related_domains': ['profile', 'creative', 'overview'],
            'highlights': [
                f"Books captured: {len(books)}.",
                f"Television favourites captured: {len(television)}.",
                f"Film favourites captured: {len(films)}.",
                f"Representative books: {', '.join(item.get('title', '') for item in books[:4])}." if books else '',
                f"Representative films: {', '.join(item.get('title', '') for item in films[:4])}." if films else '',
            ],
            'structured_state': {
                'books': books,
                'favorite_television': television,
                'favorite_films': films,
            },
            'next_questions': [
                'Which repeated directors, writers, or tonal signatures deserve their own pages?',
                'What creative themes link the favourite books, films, and television most strongly?',
                'Which references should feed landing-page, writing, or design tasks by default?',
            ],
        },
        'quote-bank': {
            'title': 'Quote bank',
            'summary': 'The stored quote taste mixes Blake, Aurelius, Tolstoy, Ram Dass, Thoreau, Alan Partridge, and film dialogue - serious spiritual guidance held alongside deadpan comedy.',
            'source_keys': ['profile_ingests', 'soul_doc'],
            'related_domains': ['profile', 'soul', 'creative', 'overview'],
            'highlights': [
                f"Quote count captured: {len(quotes)}.",
                f"Repeated quote voices: {', '.join(item.get('attribution', '') for item in quotes[:5])}." if quotes else '',
                f"Opening line: {(quotes[0] or {}).get('quote', '')}" if quotes else '',
            ],
            'structured_state': {
                'quotes': quotes,
            },
            'next_questions': [
                'Which quotes are principles versus tonal signals versus pure comedy?',
                'Which quotations should surface when summarising Dan’s worldview or writing voice?',
                'What themes recur across the quote bank strongly enough to become concept pages of their own?',
            ],
        },
        'influence-map': {
            'title': 'Influence map',
            'summary': 'The influence layer spans public thinkers and artists alongside a smaller set of personally known people who seem to have shaped Dan’s direction more directly.',
            'source_keys': ['profile_ingests', 'journal_doc', 'memory_state'],
            'related_domains': ['profile', 'people', 'soul', 'overview'],
            'highlights': [
                f"Public influences captured: {len(public_influences)}.",
                f"Personal influences captured: {len(personal_influences)}.",
                f"Personal names currently named: {', '.join(item.get('name', '') for item in personal_influences)}." if personal_influences else '',
                f"Public note count: {len(research_notes)}." if research_notes else 'No public research notes filed yet.',
            ],
            'structured_state': {
                'public_influences': public_influences,
                'personal_influences': personal_influences,
                'research_notes': research_notes,
            },
            'next_questions': [
                'Which influences are intellectual inputs versus models of character, style, or temperament?',
                'Which personally known people deserve dedicated relationship pages or richer fact sheets next?',
                'What tensions exist between the spiritual, analytical, comic, and aesthetic influence clusters?',
            ],
        },
    }


def _concept_markdown(snapshot, rel_path, slug, spec, domain_paths, source_catalog, extracted_rel_path=''):
    related_paths = [domain_paths[domain] for domain in spec.get('related_domains', []) if domain in domain_paths]
    lines = [
        f"# {spec['title']}",
        '',
        spec.get('summary', ''),
        '',
        '## highlights',
        _bullet_markdown(spec.get('highlights', [])),
        '',
        '## structured_state',
        _json_markdown_block(spec.get('structured_state', {})),
        '',
        '## next_questions',
        _bullet_markdown(spec.get('next_questions', [])),
        '',
    ]
    return _decorate_knowledge_markdown(
        snapshot,
        rel_path,
        spec['title'],
        'concept',
        spec.get('source_keys', []),
        related_paths,
        [],
        '\n'.join(lines),
        source_catalog,
        extracted_rel_path=extracted_rel_path,
    )


def _journal_keyword_matches(journal_path, keywords, limit=8):
    paragraphs = _read_docx_paragraphs(journal_path)
    if not paragraphs:
        return []

    lowered = [_collapse_whitespace(keyword).lower() for keyword in keywords if _collapse_whitespace(keyword)]
    records = _journal_sentence_records(paragraphs)
    matches = []
    seen = set()
    for record in records:
        sentence = _collapse_whitespace(record.get('sentence', ''))
        if not sentence:
            continue
        haystack = sentence.lower()
        if not any(keyword in haystack for keyword in lowered):
            continue
        key = (record.get('date', ''), sentence)
        if key in seen:
            continue
        seen.add(key)
        matches.append({
            'date': record.get('date', ''),
            'title': record.get('title', ''),
            'snippet': sentence,
        })
        if len(matches) >= limit:
            break
    return matches


def _person_page_path(name, people_directory, person_paths):
    target = _collapse_whitespace(name).lower()
    for person in people_directory:
        if _collapse_whitespace(person.get('name', '')).lower() == target:
            slug = person.get('slug', '')
            if slug and slug in person_paths:
                return person_paths[slug]
    return ''


def _build_supporting_article_specs(snapshot, domain_paths, concept_paths, person_paths, people_directory):
    docs_dir = snapshot.get('documentsDir', '')
    if not docs_dir:
        return {}

    files = _life_files(docs_dir)
    journal_path = files.get('journalDoc', '')
    if not journal_path or not os.path.exists(journal_path):
        return {}

    full_name = _collapse_whitespace(snapshot.get('profile', {}).get('fullName', '')) or 'Daniel James Atkinson'
    work_history = snapshot.get('work', {}).get('history', []) or []

    leeds_roles = [item for item in work_history if 'leeds' in _collapse_whitespace(item.get('location', '')).lower()]
    leeds_matches = _journal_keyword_matches(journal_path, ['Leeds', 'Lovell Park', 'Leeds Pride', 'Leeds Festival'], limit=8)
    london_matches = _journal_keyword_matches(journal_path, ['London', 'Canary Wharf', 'hostels', 'London marathon'], limit=8)
    austria_matches = _journal_keyword_matches(journal_path, ['Austria', 'Austrian', 'Austrian Alps'], limit=8)

    people_links = {
        'ian': _person_page_path('Ian Atkinson', people_directory, person_paths),
        'elizabeth': _person_page_path('Elizabeth Atkinson', people_directory, person_paths),
        'william': _person_page_path('William Atkinson', people_directory, person_paths),
        'anton': _person_page_path('Anton Metro', people_directory, person_paths),
        'ben': _person_page_path('Ben', people_directory, person_paths),
        'izzy': _person_page_path('Izzy', people_directory, person_paths),
    }

    leeds_related = [
        path for path in [
            domain_paths.get('people', ''),
            domain_paths.get('employment', ''),
            concept_paths.get('career-arc', ''),
            people_links['ian'],
            people_links['elizabeth'],
            people_links['william'],
            people_links['izzy'],
            'Places/london.md',
        ]
        if path
    ]
    leeds_body = '\n'.join([
        '# Leeds',
        '',
        f"Leeds is a recurring city in {full_name}'s life. It appears in the documents as a return point from London, a long-running work base, and a setting for community memories such as Leeds Pride, Leeds Festival, and office life.",
        '',
        '## role_in_life',
        f"In April 2019 the journal records a decision to leave work in London and move back to Leeds in order to save money and prepare for a run. Later entries treat Leeds as both a practical base and a place of cultural memory. It appears as a city of work, training, and imagined domestic stability rather than only as a place passed through.",
        '',
        '## documented_facts',
        _bullet_markdown([
            f"Hometown recorded in the workbook: {snapshot.get('profile', {}).get('hometown', '')}.",
            'The journal records a decision in 2019 to leave London and move back to Leeds to save money and prepare for a run.',
            'Leeds appears in the journal as the setting for Leeds Pride and multiple Leeds Festival memories.',
            'The work history places multiple BI roles in Leeds, including Sky Betting & Gaming, Leeds Building Society, and the current hybrid National Wealth Fund chapter.',
            'A journal memory places a Lloyds Banking Group office moment at Lovell Park in Leeds.',
        ]),
        '',
        '## work_connection',
        _bullet_markdown([
            f"{item.get('display', '')}." for item in leeds_roles
        ]),
        '',
        '## related_pages',
        _bullet_markdown([
            _relative_markdown_link('Places/leeds.md', concept_paths['career-arc'], 'Career arc') if concept_paths.get('career-arc') else '',
            _relative_markdown_link('Places/leeds.md', people_links['ian'], 'Ian Atkinson') if people_links['ian'] else '',
            _relative_markdown_link('Places/leeds.md', people_links['elizabeth'], 'Elizabeth Atkinson') if people_links['elizabeth'] else '',
            _relative_markdown_link('Places/leeds.md', people_links['william'], 'William Atkinson') if people_links['william'] else '',
            _relative_markdown_link('Places/leeds.md', people_links['izzy'], 'Izzy') if people_links['izzy'] else '',
            _relative_markdown_link('Places/leeds.md', 'Places/london.md', 'London'),
        ]),
        '',
        '## source_notes',
        _bullet_markdown([
            f"{item.get('date', '')}: {item.get('snippet', '')}".strip(': ')
            for item in leeds_matches[:6]
        ]),
        '',
    ])

    london_related = [
        path for path in [
            domain_paths.get('people', ''),
            concept_paths.get('career-arc', ''),
            people_links['anton'],
            'Places/leeds.md',
        ]
        if path
    ]
    london_body = '\n'.join([
        '# London',
        '',
        f"London appears in {full_name}'s documents as a period of work, optimization, and experimentation that still failed to feel like home. The journal records both commitment to staying in the city and a later decision to leave it for Leeds.",
        '',
        '## role_in_life',
        'The early 2019 journal treats London as a city of self-improvement, work pressure, dates, running, and temporary hostel living. It is described as a place where he could optimise but not fully belong. Later reflection frames London as a phase that was useful but insufficient.',
        '',
        '## documented_facts',
        _bullet_markdown([
            'A journal entry from April 2019 says London did not feel like home despite being a place of optimisation.',
            'Another entry from April 2019 says he intended to remain firmly in London and organise more dates.',
            'A later entry records the decision to leave his job in London and move back to Leeds.',
            'The journal also records living in hostels in London for more than five weeks and later seeing that period as something he conquered.',
            'An early relationship-defining conversation with Anton Metro took place at The Gun in Canary Wharf.',
        ]),
        '',
        '## related_pages',
        _bullet_markdown([
            _relative_markdown_link('Places/london.md', people_links['anton'], 'Anton Metro') if people_links['anton'] else '',
            _relative_markdown_link('Places/london.md', 'Places/leeds.md', 'Leeds'),
            _relative_markdown_link('Places/london.md', concept_paths['career-arc'], 'Career arc') if concept_paths.get('career-arc') else '',
        ]),
        '',
        '## source_notes',
        _bullet_markdown([
            f"{item.get('date', '')}: {item.get('snippet', '')}".strip(': ')
            for item in london_matches[:6]
        ]),
        '',
    ])

    austria_related = [
        path for path in [
            domain_paths.get('journal', ''),
            domain_paths.get('running', ''),
            concept_paths.get('north-star', ''),
        ]
        if path
    ]
    austria_body = '\n'.join([
        '# Austria leg of the 2019 run',
        '',
        f"The Austria leg of the 2019 run appears in the journal as a spiritual and physical high point. {full_name} records crossing into Austria on 20 October 2019 and describes the country as sacred, mountainous, and physically demanding in a way that felt welcome rather than punishing.",
        '',
        '## significance',
        'Austria is not described simply as another country on the route. The journal frames the crossing as ceremonial. A waterfall marks the closing of one chapter, and the mountains are treated as a setting that intensifies both physical effort and spiritual feeling.',
        '',
        '## timeline',
        '| Date | Event |',
        '|------|-------|',
        '| 19/10/2019 | In a hotel room 20km away from Austria after a difficult logistical day. |',
        '| 20/10/2019 | Crossed the Austrian border and described entering a sacred land surrounded by mountains. |',
        '| 20/10/2019 | Described walking under a small waterfall as the closing of another chapter in the journey. |',
        '| Day 30 | Recorded climbing over the Austrian Alps to a peak height of 2500m. |',
        '',
        '## documented_facts',
        _bullet_markdown([
            'The journal records Austria as country number 3 on the route.',
            'The crossing is described in explicitly spiritual terms rather than only athletic ones.',
            'The route in Austria included steep mountains, streams, scrambles, ridge trails, and a climb to 2500m.',
            'One entry says he would be out of Austria in two days with very low mileage left to go.',
        ]),
        '',
        '## related_pages',
        _bullet_markdown([
            _relative_markdown_link('Events/austria-leg-of-the-2019-run.md', domain_paths['running'], 'Running') if domain_paths.get('running') else '',
            _relative_markdown_link('Events/austria-leg-of-the-2019-run.md', domain_paths['journal'], 'Journal') if domain_paths.get('journal') else '',
            _relative_markdown_link('Events/austria-leg-of-the-2019-run.md', concept_paths['north-star'], 'North star') if concept_paths.get('north-star') else '',
        ]),
        '',
        '## source_notes',
        _bullet_markdown([
            f"{item.get('date', '')}: {item.get('snippet', '')}".strip(': ')
            for item in austria_matches[:8]
        ]),
        '',
    ])

    return {
        'Places/leeds.md': {
            'title': 'Leeds',
            'kind': 'place',
            'summary': 'Leeds functions as a recurring city of return, work, memory, and future-planning in Daniel’s documents.',
            'source_keys': ['journal_doc', 'finance_workbook'],
            'related_paths': leeds_related,
            'backlink_paths': ['Indexes/article-index.md'],
            'structured_state': {
                'hometown': snapshot.get('profile', {}).get('hometown', ''),
                'work_roles': leeds_roles,
                'related_people': [name for name, path in [('Ian Atkinson', people_links['ian']), ('Elizabeth Atkinson', people_links['elizabeth']), ('William Atkinson', people_links['william']), ('Izzy', people_links['izzy'])] if path],
            },
            'source_notes': leeds_matches[:8],
            'body': leeds_body,
        },
        'Places/london.md': {
            'title': 'London',
            'kind': 'place',
            'summary': 'London appears as a formative but ultimately insufficient phase of work, optimisation, dates, and hostel living.',
            'source_keys': ['journal_doc'],
            'related_paths': london_related,
            'backlink_paths': ['Indexes/article-index.md'],
            'structured_state': {
                'related_people': [name for name, path in [('Anton Metro', people_links['anton'])] if path],
                'themes': ['work', 'optimisation', 'dates', 'hostel living', 'transition back to Leeds'],
            },
            'source_notes': london_matches[:8],
            'body': london_body,
        },
        'Events/austria-leg-of-the-2019-run.md': {
            'title': 'Austria leg of the 2019 run',
            'kind': 'event',
            'summary': 'A journal-defined high point in the 2019 run, marked by spiritual language, mountain terrain, and a ceremonial border crossing.',
            'source_keys': ['journal_doc'],
            'related_paths': austria_related,
            'backlink_paths': ['Indexes/article-index.md'],
            'structured_state': {
                'timeline': [
                    {'date': '2019-10-19', 'event': 'In a hotel room 20km away from Austria after a difficult logistical day.'},
                    {'date': '2019-10-20', 'event': 'Crossed the Austrian border and described entering a sacred land surrounded by mountains.'},
                    {'date': '2019-10-20', 'event': 'Described walking under a small waterfall as the closing of another chapter in the journey.'},
                    {'date': 'Day 30', 'event': 'Recorded climbing over the Austrian Alps to a peak height of 2500m.'},
                ],
                'themes': ['spiritual high point', 'mountains', 'ceremonial border crossing', 'physical hardship'],
            },
            'source_notes': austria_matches[:10],
            'body': austria_body,
        },
    }


def _article_markdown(snapshot, rel_path, spec, source_catalog, extracted_rel_path=''):
    return _decorate_knowledge_markdown(
        snapshot,
        rel_path,
        spec['title'],
        spec.get('kind', 'article'),
        spec.get('source_keys', []),
        spec.get('related_paths', []),
        spec.get('backlink_paths', []),
        spec.get('body', ''),
        source_catalog,
        extracted_rel_path=extracted_rel_path,
    )


def _build_health_checks(snapshot, scratchpad):
    findings = []
    onebag = snapshot.get('onebag', {})
    money = snapshot.get('money', {})
    work = snapshot.get('work', {})
    market = snapshot.get('marketPerformance', {})
    history = snapshot.get('portfolioHistory', []) or []
    journal_capture = snapshot.get('journalCapture', {})

    if onebag.get('weightGapKg') is not None and onebag.get('weightGapKg', 0) > 0:
        findings.append({
            'severity': 'warning',
            'issue': 'Onebag is still over target weight.',
            'detail': f"Tracked setup is {onebag['weightGapKg']:.2f} kg above the current {onebag.get('targetWeightKg', 0):.1f} kg target.",
        })

    if money.get('medicalBillUsd'):
        findings.append({
            'severity': 'warning',
            'issue': 'Medical bill remains unresolved.',
            'detail': f"Outstanding note remains at ${money['medicalBillUsd']:,.2f} and still deserves explicit follow-up tracking.",
        })

    if len(history) < 14:
        findings.append({
            'severity': 'info',
            'issue': 'Investment history is still shallow.',
            'detail': f"Private life-history snapshots only contain {len(history)} points so far, so trends are still early.",
        })

    if not journal_capture.get('canServerTranscribe'):
        findings.append({
            'severity': 'info',
            'issue': 'AI transcription is not enabled locally.',
            'detail': 'Voice-note transcription still depends on a local OpenAI key and SDK install.',
        })

    if work.get('history'):
        findings.append({
            'severity': 'info',
            'issue': 'Previous salary history is missing.',
            'detail': 'Employment history has role timeline data, but older compensation points are not yet recorded.',
        })

    if not scratchpad.get('workingOn'):
        findings.append({
            'severity': 'info',
            'issue': 'Scratchpad focus is thin.',
            'detail': 'Current working-on items are missing, which will make future sessions less coherent.',
        })

    recommendations = [
        'Add previous salary information for past roles so the career arc becomes financially legible.',
        'Keep filing durable outputs back into Outputs/ so useful work does not disappear into chat history.',
        'Start putting new external research or clipped sources into Raw/ when Akibwa expands beyond the current personal docs.',
    ]

    if market.get('ok'):
        recommendations.insert(1, 'Let the investment history keep compounding so the current-position line becomes more decision-useful over time.')

    return findings, recommendations[:4]


def _write_support_layer_files(docs_dir, source_catalog):
    raw_dir = _raw_dir(docs_dir)
    state_dir = _state_dir(docs_dir)
    os.makedirs(raw_dir, exist_ok=True)
    os.makedirs(state_dir, exist_ok=True)

    raw_records = [record for record in source_catalog.values() if record.get('kind') in {'workbook', 'document', 'directory'}]
    state_records = [record for record in source_catalog.values() if record.get('kind') == 'state']

    _write_text_atomic(
        os.path.join(raw_dir, 'README.md'),
        '\n'.join([
            '# Raw',
            '',
            'Immutable or externally sourced material should land here first.',
            'For now, the canonical personal source documents still live at the root of Akibwa Documents, and this folder acts as the intake area for future clipped articles, PDFs, images, and other research assets.',
            '',
            '## monitored_sources',
            _json_markdown_block(raw_records),
            '',
        ]),
    )
    _write_text_atomic(
        os.path.join(raw_dir, 'source-manifest.md'),
        '\n'.join([
            '# Raw Source Manifest',
            '',
            'Catalog of the currently monitored raw sources that feed the Akibwa knowledge base.',
            '',
            _json_markdown_block(raw_records),
            '',
        ]),
    )
    _write_text_atomic(
        os.path.join(state_dir, 'README.md'),
        '\n'.join([
            '# State',
            '',
            'Machine-managed state lives here. These files are designed for agents and scripts rather than direct manual editing.',
            'Highly sensitive local-only markdown that should not be shipped to hosted exports belongs in `State/Private Context/`.',
            '',
            '## state_files',
            _json_markdown_block(state_records),
            '',
        ]),
    )


def _cleanup_legacy_ai_context_files(target_dir):
    legacy = [
        'overview.md',
        'onebag.md',
        'fashion.md',
        'running.md',
        'people.md',
        'finances.md',
        'investments.md',
        'creative.md',
        'journal.md',
        'employment.md',
        'soul.md',
    ]
    for filename in legacy:
        path = os.path.join(target_dir, filename)
        if os.path.exists(path):
            try:
                os.remove(path)
            except OSError:
                pass


def _prune_generated_ai_context_files(target_dir, keep_paths):
    normalized_keep = {path.replace(os.sep, '/') for path in keep_paths}
    for root, dirs, files in os.walk(target_dir, topdown=False):
        for filename in files:
            full_path = os.path.join(root, filename)
            relative_path = os.path.relpath(full_path, start=target_dir).replace(os.sep, '/')
            if relative_path in normalized_keep:
                continue
            try:
                os.remove(full_path)
            except OSError:
                pass
        for dirname in dirs:
            full_path = os.path.join(root, dirname)
            try:
                if not os.listdir(full_path):
                    os.rmdir(full_path)
            except OSError:
                pass


def _prune_generated_layer_files(target_dir, keep_paths):
    normalized_keep = {path.replace(os.sep, '/') for path in keep_paths}
    for root, dirs, files in os.walk(target_dir, topdown=False):
        for filename in files:
            full_path = os.path.join(root, filename)
            relative_path = os.path.relpath(full_path, start=target_dir).replace(os.sep, '/')
            if relative_path in normalized_keep:
                continue
            try:
                os.remove(full_path)
            except OSError:
                pass
        for dirname in dirs:
            full_path = os.path.join(root, dirname)
            try:
                if not os.listdir(full_path):
                    os.rmdir(full_path)
            except OSError:
                pass


def _person_context_bundle(snapshot, person, rel_path):
    profile_library = snapshot.get('profileLibrary', {})
    personal_match = next(
        (item for item in profile_library.get('personalInfluences', []) if _collapse_whitespace(item.get('name', '')).lower() == person.get('name', '').lower()),
        None,
    )
    research_match = next(
        (item for item in profile_library.get('researchNotes', []) if _collapse_whitespace(item.get('name', '')).lower() == person.get('name', '').lower()),
        None,
    )
    source_keys = ['journal_doc']
    if personal_match or research_match:
        source_keys.append('profile_ingests')

    evidence = person.get('evidence', []) or []
    evidence_text = ' '.join(item.get('snippet', '') for item in evidence[:8]).lower()
    role_notes = []
    if person.get('relationToDan'):
        role_notes.append(
            f"{person.get('name')} is recorded in Akibwa as Daniel James Atkinson's {PERSON_RELATION_TO_DAN_LABELS.get(person.get('relationToDan'), person.get('relationToDan'))}."
        )
    else:
        role_notes.append(
            f"{person.get('name')} is treated as a recurring figure in Daniel James Atkinson's documents."
        )
    role_notes.append(
        f"The current compiled record spans {person.get('mentions', 0)} mention{'s' if person.get('mentions', 0) != 1 else ''}"
        + (
            f" from {person.get('firstMentionDate')} to {person.get('lastMentionDate')}."
            if person.get('firstMentionDate') and person.get('lastMentionDate') and person.get('firstMentionDate') != person.get('lastMentionDate')
            else f" on {person.get('firstMentionDate')}."
            if person.get('firstMentionDate')
            else '.'
        )
    )
    if any(keyword in evidence_text for keyword in ('honest', 'chat', 'conversation', 'opened up', 'speak to')):
        role_notes.append('Entries repeatedly frame the relationship through conversation, honesty, or emotional openness.')
    if any(keyword in evidence_text for keyword in ('walk', 'cycling', 'cycle', 'birthday', 'coach', 'day out')):
        role_notes.append('Shared movement, day trips, or ordinary time together recur in the recorded context.')
    if any(keyword in evidence_text for keyword in ('drink', 'beer', 'pint', 'alcohol')):
        role_notes.append('Drinking and family habits appear as one of the recurring pressures or themes around this person.')
    if any(keyword in evidence_text for keyword in ('work', 'office', 'job', 'manager', 'career')):
        role_notes.append('Work and career context are part of the documented relationship.')
    if any(keyword in evidence_text for keyword in ('dating', 'kiss', 'love', 'romantic', 'relationship', 'sex')):
        role_notes.append('The source material places this person inside a romantic or intimate chapter of life.')
    if personal_match and personal_match.get('whyInfluential'):
        role_notes.append(personal_match.get('whyInfluential'))
    if research_match and research_match.get('summary'):
        role_notes.append(research_match.get('summary'))

    related_paths = ['Domains/people.md', 'Domains/journal.md']
    if 'leeds' in evidence_text:
        related_paths.append('Places/leeds.md')
    if 'london' in evidence_text:
        related_paths.append('Places/london.md')
    if 'austria' in evidence_text:
        related_paths.append('Events/austria-leg-of-the-2019-run.md')
    family_cross_links = {
        'father': ['People/elizabeth-atkinson.md', 'People/william-atkinson.md', 'Places/leeds.md'],
        'mother': ['People/ian-atkinson.md', 'People/william-atkinson.md', 'Places/leeds.md'],
        'brother': ['People/ian-atkinson.md', 'People/elizabeth-atkinson.md', 'Places/leeds.md'],
    }
    for path in family_cross_links.get(person.get('relationToDan'), []):
        if path != rel_path:
            related_paths.append(path)
    related_paths = list(dict.fromkeys(path for path in related_paths if path != rel_path))

    return {
        'source_keys': source_keys,
        'related_paths': related_paths,
        'personal_match': personal_match,
        'research_match': research_match,
        'evidence': evidence,
        'role_notes': [note for note in dict.fromkeys(role_notes) if note],
        'evidence_text': evidence_text,
    }


def _build_domain_extract_payload(snapshot, slug, title, article_path, source_keys, related_paths):
    profile = snapshot.get('profile', {})
    profile_library = snapshot.get('profileLibrary', {})
    onebag = snapshot.get('onebag', {})
    money = snapshot.get('money', {})
    market = snapshot.get('marketPerformance', {})
    journal = snapshot.get('journal', {})
    work = snapshot.get('work', {})
    soul = snapshot.get('soul', {})
    running = snapshot.get('running', {})
    fashion = _build_fashion_context(onebag)

    summary_map = {
        'overview': 'Top-level orientation page over the current Akibwa knowledge base.',
        'profile': 'Declared reading, quote, influence, and media profile.',
        'onebag': 'Tracked onebag carry system and current loadout.',
        'fashion': 'Wearable wardrobe subset extracted from the tracked inventory.',
        'running': 'Current one-shoe strategy and running-related decision record.',
        'people': 'Directory and entrypoint into known people articles.',
        'finances': 'Current liquid, outgoing, and liability position.',
        'investments': 'Held positions and market-value tracking.',
        'creative': 'Writing shelf and compiled creative fragments.',
        'journal': 'Journal themes and recent entries.',
        'employment': 'Current role, compensation, and career timeline.',
        'soul': 'North star, values, and becoming layer.',
    }

    structured_state = {
        'overview': {
            'full_name': profile.get('fullName', ''),
            'current_location': profile.get('currentLiving', ''),
            'current_role': work.get('jobTitle') or work.get('linkedinTitle', ''),
            'employer': work.get('employer', ''),
            'people_count': (snapshot.get('people', {}).get('summary') or {}).get('person_count', 0),
            'source_updated_at': snapshot.get('sourceUpdatedAt', ''),
        },
        'profile': {
            'summary': profile_library.get('summary', {}),
            'books': profile_library.get('booksRead', []),
            'favorite_quotes': profile_library.get('favoriteQuotes', []),
            'public_influences': profile_library.get('publicInfluences', []),
            'personal_influences': profile_library.get('personalInfluences', []),
            'favorite_television': profile_library.get('favoriteTelevision', []),
            'favorite_films': profile_library.get('favoriteFilms', []),
        },
        'onebag': {
            'summary': {
                'tracked_weight_kg': onebag.get('trackedWeightKg', 0),
                'target_weight_kg': onebag.get('targetWeightKg'),
                'weight_gap_kg': onebag.get('weightGapKg'),
                'tracked_item_count': onebag.get('trackedItemCount', 0),
                'tracked_value_gbp': onebag.get('trackedValueGbp', 0),
                'backpack': onebag.get('backpack', {}),
            },
            'category_breakdown': onebag.get('categoryBreakdown', []),
            'heaviest_items': onebag.get('heaviestItems', []),
            'most_expensive_items': onebag.get('mostExpensiveItems', []),
        },
        'fashion': {
            'summary': fashion.get('summary', {}),
            'groups': fashion.get('groups', []),
            'items': fashion.get('items', []),
            'highlights': fashion.get('highlights', []),
        },
        'running': {
            'goal': running.get('goal', ''),
            'status_line': running.get('statusLine', ''),
            'requirements': running.get('requirements', []),
            'current_shoes': running.get('currentShoes', []),
            'shortlist': running.get('shortlist', []),
            'palette_direction': running.get('paletteDirection', {}),
            'recommendation': running.get('recommendation', {}),
        },
        'people': {
            'summary': snapshot.get('people', {}).get('summary', {}),
            'directory': snapshot.get('people', {}).get('directory', []),
        },
        'finances': {
            'summary': {
                'cash_total_gbp': money.get('cashTotal', 0),
                'investments_total_gbp': money.get('investmentsTotal', 0),
                'monthly_outgoings_total_gbp': money.get('monthlyOutgoingsTotal', 0),
                'liabilities_gbp_total': money.get('liabilitiesGbpTotal', 0),
                'medical_bill_usd': money.get('medicalBillUsd', 0),
            },
            'liabilities': money.get('liabilities', []),
            'monthly_outgoings': money.get('monthlyOutgoings', []),
        },
        'investments': {
            'summary': {
                'ok': market.get('ok', False),
                'market_value_gbp': market.get('marketValueGbp', 0),
                'total_with_cash_gbp': market.get('totalWithCashGbp', 0),
                'as_of_date': market.get('asOfDate', ''),
                'point_count': market.get('pointCount', 0),
            },
            'positions': market.get('positions', []),
            'warnings': market.get('warnings', []),
        },
        'creative': {
            'titles': snapshot.get('creative', {}).get('titles', []),
            'entries': snapshot.get('creative', {}).get('entries', []),
            'sample_lines': snapshot.get('creative', {}).get('sampleLines', []),
        },
        'journal': {
            'themes': journal.get('themes', []),
            'latest_entries': journal.get('latestEntries', []),
        },
        'employment': {
            'summary': {
                'current_role': work.get('jobTitle') or work.get('linkedinTitle', ''),
                'employer': work.get('employer', ''),
                'office_address': work.get('officeAddress', ''),
                'working_pattern': work.get('workingPattern', ''),
                'annual_gross_salary': work.get('annualGrossSalary', 0),
                'total_annual_comp_gross': work.get('totalAnnualCompGross', 0),
            },
            'recent_payslips': work.get('recentPayslips', []),
            'history': work.get('history', []),
        },
        'soul': {
            'title': soul.get('title', ''),
            'north_star': soul.get('northStar', ''),
            'question': soul.get('question', ''),
            'values': soul.get('values', []),
            'becoming': soul.get('becoming', []),
            'non_negotiables': soul.get('nonNegotiables', []),
        },
    }

    return {
        'title': title,
        'kind': 'domain',
        'slug': slug,
        'article_path': article_path,
        'summary': summary_map.get(slug, f'{title} domain page.'),
        'source_keys': source_keys,
        'related_article_paths': related_paths,
        'structured_state': structured_state.get(slug, {}),
    }


def _build_extracted_layer(snapshot, domain_titles, domain_paths, domain_sources, domain_concepts, concepts, concept_paths, people_directory, person_paths, article_specs):
    docs_dir = snapshot.get('documentsDir', '')
    if not docs_dir:
        return {
            'directory': '',
            'files': [],
            'entrypoints': [],
        }

    target_dir = _extracted_dir(docs_dir)
    os.makedirs(target_dir, exist_ok=True)

    content_map = {}
    catalogs = defaultdict(list)
    page_records = {}

    for slug, title in domain_titles.items():
        article_path = domain_paths[slug]
        extracted_rel_path = _extracted_rel_path_for_article(article_path)
        payload = _build_domain_extract_payload(
            snapshot,
            slug,
            title,
            article_path,
            domain_sources.get(slug, []),
            [concept_paths[key] for key in domain_concepts.get(slug, []) if key in concept_paths],
        )
        content_map[extracted_rel_path] = payload
        page_records[article_path] = payload
        catalogs['domains'].append({
            'path': extracted_rel_path,
            'title': title,
            'article_path': article_path,
            'source_keys': payload.get('source_keys', []),
        })

    for slug, spec in concepts.items():
        article_path = concept_paths[slug]
        extracted_rel_path = _extracted_rel_path_for_article(article_path)
        payload = {
            'title': spec['title'],
            'kind': 'concept',
            'slug': slug,
            'article_path': article_path,
            'summary': spec.get('summary', ''),
            'source_keys': spec.get('source_keys', []),
            'related_article_paths': [domain_paths[domain] for domain in spec.get('related_domains', []) if domain in domain_paths],
            'highlights': spec.get('highlights', []),
            'structured_state': spec.get('structured_state', {}),
            'next_questions': spec.get('next_questions', []),
        }
        content_map[extracted_rel_path] = payload
        page_records[article_path] = payload
        catalogs['concepts'].append({
            'path': extracted_rel_path,
            'title': spec['title'],
            'article_path': article_path,
            'source_keys': spec.get('source_keys', []),
        })

    for person in people_directory:
        slug = person.get('slug', _slugify(person.get('name', 'person')))
        article_path = person_paths.get(slug)
        if not article_path:
            continue
        extracted_rel_path = _extracted_rel_path_for_article(article_path)
        bundle = _person_context_bundle(snapshot, person, article_path)
        payload = {
            'title': person.get('name', 'Unknown person'),
            'kind': 'person',
            'slug': slug,
            'article_path': article_path,
            'summary': person.get('articleLead', person.get('summary', '')),
            'source_keys': bundle['source_keys'],
            'related_article_paths': bundle['related_paths'],
            'structured_state': {
                'relation_to_dan': person.get('relationToDan', ''),
                'primary_relationship': person.get('primaryRelationship', ''),
                'mentions': person.get('mentions', 0),
                'first_mention_date': person.get('firstMentionDate', ''),
                'last_mention_date': person.get('lastMentionDate', ''),
                'aliases': person.get('aliases', []),
                'facts': person.get('facts', []),
                'relation_summary': person.get('relationSummary', []),
                'context_notes': person.get('contextNotes', []),
                'role_notes': bundle['role_notes'],
                'evidence': bundle['evidence'][:12],
                'profile_context': bundle['personal_match'] or {},
                'research_context': bundle['research_match'] or {},
            },
        }
        content_map[extracted_rel_path] = payload
        page_records[article_path] = payload
        catalogs['people'].append({
            'path': extracted_rel_path,
            'title': person.get('name', 'Unknown person'),
            'article_path': article_path,
            'source_keys': bundle['source_keys'],
            'primary_relationship': person.get('primaryRelationship', ''),
            'mentions': person.get('mentions', 0),
        })

    for article_path, spec in article_specs.items():
        extracted_rel_path = _extracted_rel_path_for_article(article_path)
        payload = {
            'title': spec.get('title', ''),
            'kind': spec.get('kind', 'article'),
            'slug': os.path.splitext(os.path.basename(article_path))[0],
            'article_path': article_path,
            'summary': spec.get('summary', ''),
            'source_keys': spec.get('source_keys', []),
            'related_article_paths': spec.get('related_paths', []),
            'structured_state': spec.get('structured_state', {}),
            'source_notes': spec.get('source_notes', []),
        }
        content_map[extracted_rel_path] = payload
        page_records[article_path] = payload
        group_name = article_path.split('/', 1)[0].lower()
        catalogs[group_name].append({
            'path': extracted_rel_path,
            'title': spec.get('title', ''),
            'article_path': article_path,
            'source_keys': spec.get('source_keys', []),
            'kind': spec.get('kind', 'article'),
        })

    content_map['index.json'] = {
        'generated_at': snapshot.get('generatedAt', ''),
        'source_updated_at': snapshot.get('sourceUpdatedAt', ''),
        'counts': {key: len(value) for key, value in catalogs.items()},
        'entrypoints': ['README.md', 'index.json', 'domains/index.json', 'people/index.json'],
        'catalogs': {
            key: [item.get('path', '') for item in value]
            for key, value in catalogs.items()
        },
    }
    for key, items in catalogs.items():
        content_map[f'{key}/index.json'] = {
            'kind': key,
            'count': len(items),
            'records': items,
        }

    readme_lines = [
        '# Extracted',
        '',
        'Machine-readable normalized records live here. Each article page in Akibwa can point back to one of these extracted JSON files.',
        '',
        '## purpose',
        '- Keep a stable intermediate layer between raw sources and compiled wiki articles.',
        '- Make full recompilation with stronger future models easier without losing article-level provenance.',
        '- Let pages such as people, places, events, and domain articles read from reusable extracted records instead of ad hoc parsing.',
        '',
        '## entry_points',
        '- `index.json`',
    ]
    for key in sorted(catalogs):
        readme_lines.append(f"- `{key}/index.json`")
    readme_lines.extend([
        '',
        '## counts',
        _json_markdown_block({key: len(value) for key, value in catalogs.items()}),
        '',
    ])
    _write_text_atomic(os.path.join(target_dir, 'README.md'), '\n'.join(readme_lines))

    _prune_generated_layer_files(target_dir, list(content_map.keys()) + ['README.md'])
    for relative_path, payload in content_map.items():
        _write_json_atomic(os.path.join(target_dir, relative_path), payload)

    generated_files = sorted(
        _relative_docs_path(docs_dir, os.path.join(target_dir, relative_path))
        for relative_path in list(content_map.keys()) + ['README.md']
    )
    return {
        'directory': target_dir,
        'files': generated_files,
        'entrypoints': ['README.md', 'index.json', 'domains/index.json', 'people/index.json'],
        'records': page_records,
        'catalogs': {key: value for key, value in catalogs.items()},
    }


def _ai_context_site_group(path):
    if not path:
        return 'root'

    mapping = {
        'Domains': 'domains',
        'Concepts': 'concepts',
        'People': 'people',
        'Places': 'places',
        'Events': 'events',
        'Indexes': 'indexes',
        'Outputs': 'outputs',
    }
    head = path.split('/', 1)[0]
    if head in mapping:
        return mapping[head]
    if path in ('README.md', 'index.md'):
        return 'root'
    if path in ('schema.md', 'log.md'):
        return 'meta'
    return 'root'


def _site_link_record(path, title, note=''):
    return {
        'path': path,
        'title': title,
        'note': note,
    }


def _build_ai_context_site_pages(extracted_records, extra_pages):
    pages = []
    for article_path, record in (extracted_records or {}).items():
        if not article_path or not isinstance(record, dict):
            continue

        kind = record.get('kind', 'article')
        slug = record.get('slug', os.path.splitext(os.path.basename(article_path))[0])
        pages.append({
            'path': article_path,
            'group': _ai_context_site_group(article_path),
            'kind': kind,
            'slug': slug,
            'title': record.get('title', article_path),
            'summary': record.get('summary', ''),
            'source_keys': record.get('source_keys', []),
            'related_paths': record.get('related_article_paths', []),
            'structured_state': record.get('structured_state', {}),
            'highlights': record.get('highlights', []),
            'next_questions': record.get('next_questions', []),
            'source_notes': record.get('source_notes', []),
            'section': slug if kind == 'domain' else '',
            'person_slug': slug if kind == 'person' else '',
        })

    for page in extra_pages:
        if not isinstance(page, dict) or not page.get('path'):
            continue
        payload = {
            'group': _ai_context_site_group(page.get('path', '')),
            'kind': page.get('kind', 'index'),
            'slug': page.get('slug', os.path.splitext(os.path.basename(page.get('path', '')))[0]),
            'title': page.get('title', page.get('path', '')),
            'summary': page.get('summary', ''),
            'source_keys': page.get('source_keys', []),
            'related_paths': page.get('related_paths', []),
            'structured_state': page.get('structured_state', {}),
            'highlights': page.get('highlights', []),
            'next_questions': page.get('next_questions', []),
            'source_notes': page.get('source_notes', []),
            'section': '',
            'person_slug': '',
            **page,
        }
        pages.append(payload)

    known_paths = {page.get('path', '') for page in pages if page.get('path')}
    incoming = defaultdict(set)
    for page in pages:
        related = []
        for target in page.get('related_paths', []) or []:
            if not target or target in related:
                continue
            related.append(target)
            if target in known_paths:
                incoming[target].add(page.get('path', ''))
        page['related_paths'] = related

    order = {
        'root': 0,
        'domains': 1,
        'concepts': 2,
        'people': 3,
        'places': 4,
        'events': 5,
        'indexes': 6,
        'outputs': 7,
        'meta': 8,
    }

    for page in pages:
        page['incoming_paths'] = sorted(path for path in incoming.get(page.get('path', ''), set()) if path)

    pages.sort(key=lambda page: (order.get(page.get('group', ''), 99), str(page.get('title', '')).lower()))
    return pages


def _world_map_add_node(nodes, node_id, title, kind, cluster, path='', summary='', status='', source_keys=None):
    if not node_id:
        return

    record = nodes.get(node_id)
    if not record:
        record = {
            'id': node_id,
            'title': title,
            'kind': kind,
            'cluster': cluster,
            'path': path,
            'summary': summary,
            'status': status,
            'source_keys': [],
        }
        nodes[node_id] = record

    if title and not record.get('title'):
        record['title'] = title
    if path and not record.get('path'):
        record['path'] = path
    if summary and not record.get('summary'):
        record['summary'] = summary
    if status and not record.get('status'):
        record['status'] = status
    for key in source_keys or []:
        if key and key not in record['source_keys']:
            record['source_keys'].append(key)


def _world_map_add_edge(edges, seen, source, relation, target, note='', source_keys=None):
    if not source or not relation or not target:
        return
    key = (source, relation, target, note)
    if key in seen:
        return
    seen.add(key)
    edges.append({
        'from': source,
        'relation': relation,
        'to': target,
        'note': note,
        'source_keys': [key for key in (source_keys or []) if key],
    })


def _build_world_map(snapshot, domain_paths, concept_paths, person_paths, people_directory, article_specs):
    profile = snapshot.get('profile', {})
    profile_library = snapshot.get('profileLibrary', {})
    work = snapshot.get('work', {})
    market = snapshot.get('marketPerformance', {})
    money = snapshot.get('money', {})
    soul = snapshot.get('soul', {})
    running = snapshot.get('running', {})
    creative = snapshot.get('creative', {})
    full_name = _collapse_whitespace(profile.get('fullName', '')) or 'Daniel James Atkinson'
    public_influences = profile_library.get('publicInfluences', []) or []
    personal_influences = profile_library.get('personalInfluences', []) or []
    person_slug_by_name = {
        _collapse_whitespace(person.get('name', '')).lower(): person.get('slug', _slugify(person.get('name', '')))
        for person in people_directory
        if _collapse_whitespace(person.get('name', ''))
    }

    nodes = {}
    edges = []
    seen_edges = set()

    _world_map_add_node(
        nodes,
        'self:daniel-james-atkinson',
        full_name,
        'self',
        'self',
        path=domain_paths.get('overview', ''),
        summary='Central subject of the Akibwa knowledge base.',
        source_keys=['inventory_workbook', 'journal_doc', 'memory_state'],
    )

    current_living = _collapse_whitespace(profile.get('currentLiving', ''))
    current_residence_place = ''
    if current_living:
        residence_label = current_living.split('(')[0].strip().rstrip(',')
        residence_note = ''
        if '(' in current_living and ')' in current_living:
            residence_note = current_living[current_living.find('(') + 1:current_living.rfind(')')].strip()
        residence_id = f"residence:{_slugify(residence_label)}"
        _world_map_add_node(
            nodes,
            residence_id,
            residence_label,
            'residence',
            'residences',
            summary='Current recorded home base from the overview layer.',
            status='current',
            source_keys=['inventory_workbook', 'finance_workbook'],
        )
        _world_map_add_edge(
            edges,
            seen_edges,
            full_name,
            'currently lives at',
            residence_label,
            residence_note,
            source_keys=['inventory_workbook', 'finance_workbook'],
        )
        residence_parts = [part.strip() for part in residence_label.split(',') if part.strip()]
        broader_place = residence_parts[-1] if residence_parts else ''
        broader_place = re.sub(r'\b[A-Z]{1,2}\d[\dA-Z]?\s*\d[A-Z]{2}\b', '', broader_place, flags=re.I).strip(' ,')
        if broader_place:
            current_residence_place = broader_place
            _world_map_add_node(
                nodes,
                f"place:{_slugify(broader_place)}",
                broader_place,
                'place',
                'places',
                summary='Broader place extracted from the recorded current residence.',
                source_keys=['inventory_workbook', 'finance_workbook'],
            )
            _world_map_add_edge(
                edges,
                seen_edges,
                residence_label,
                'sits within',
                broader_place,
                'Derived from the current residence string in the overview layer.',
                source_keys=['inventory_workbook', 'finance_workbook'],
            )

    hometown = _collapse_whitespace(profile.get('hometown', ''))
    if hometown:
        hometown_id = f"place:{_slugify(hometown)}"
        _world_map_add_node(
            nodes,
            hometown_id,
            hometown,
            'place',
            'places',
            summary='Recorded hometown from workbook sources.',
            source_keys=['finance_workbook'],
        )
        _world_map_add_edge(
            edges,
            seen_edges,
            full_name,
            'comes from',
            hometown,
            'Recorded as hometown in the workbook layer.',
            source_keys=['finance_workbook'],
        )

    important_domain_nodes = [
        ('people', 'people'),
        ('journal', 'journal'),
        ('employment', 'work'),
        ('investments', 'money'),
        ('finances', 'money'),
        ('soul', 'philosophy'),
        ('profile', 'identity'),
        ('creative', 'philosophy'),
        ('running', 'systems'),
    ]
    for slug, cluster in important_domain_nodes:
        path = domain_paths.get(slug, '')
        if not path:
            continue
        _world_map_add_node(
            nodes,
            f"page:{path}",
            slug.replace('-', ' ').title(),
            'domain',
            cluster,
            path=path,
            summary=f'{slug.title()} domain page.',
            source_keys=[],
        )

    important_concepts = [
        ('north-star', 'philosophy'),
        ('career-arc', 'work'),
        ('financial-position', 'money'),
        ('influence-map', 'relationships'),
        ('media-canon', 'philosophy'),
        ('quote-bank', 'philosophy'),
        ('creative-practice', 'philosophy'),
        ('capsule-wardrobe', 'systems'),
        ('one-shoe-strategy', 'systems'),
    ]
    for slug, cluster in important_concepts:
        path = concept_paths.get(slug, '')
        if not path:
            continue
        _world_map_add_node(
            nodes,
            f"page:{path}",
            slug.replace('-', ' ').title(),
            'concept',
            cluster,
            path=path,
            summary='Concept page in the compiled wiki.',
            source_keys=[],
        )

    _world_map_add_edge(
        edges,
        seen_edges,
        full_name,
        'orients around',
        'North star',
        soul.get('question', ''),
        source_keys=['soul_doc', 'journal_doc'],
    )
    _world_map_add_edge(
        edges,
        seen_edges,
        full_name,
        'understands work through',
        'Career arc',
        'Employment history is synthesised into a longer career narrative.',
        source_keys=['finance_workbook'],
    )
    _world_map_add_edge(
        edges,
        seen_edges,
        full_name,
        'understands money through',
        'Financial position',
        'Cash, liabilities, and investments are treated as one picture.',
        source_keys=['finance_workbook', 'market_cache', 'life_history'],
    )
    _world_map_add_edge(
        edges,
        seen_edges,
        full_name,
        'draws language from',
        'Quote bank',
        'Stored quotations act as worldview and tonal reference.',
        source_keys=['profile_ingests', 'soul_doc'],
    )
    _world_map_add_edge(
        edges,
        seen_edges,
        full_name,
        'is shaped by',
        'Influence map',
        'Public and personal influences sit together here.',
        source_keys=['profile_ingests', 'journal_doc', 'memory_state'],
    )
    _world_map_add_edge(
        edges,
        seen_edges,
        full_name,
        'keeps taste in',
        'Media canon',
        'Books, films, and television shape the wider reference layer.',
        source_keys=['profile_ingests'],
    )
    _world_map_add_edge(
        edges,
        seen_edges,
        full_name,
        'organises clothing through',
        'Capsule wardrobe',
        'Wardrobe, onebag, and shoe choices are linked.',
        source_keys=['inventory_workbook', 'running_state', 'memory_state'],
    )
    _world_map_add_edge(
        edges,
        seen_edges,
        full_name,
        'tests trade-offs through',
        'One-shoe strategy',
        running.get('recommendation', {}).get('summary', ''),
        source_keys=['running_state', 'inventory_workbook'],
    )

    for rel_path in ('Places/leeds.md', 'Places/london.md', 'Events/austria-leg-of-the-2019-run.md'):
        spec = article_specs.get(rel_path, {})
        if not spec:
            continue
        title = spec.get('title', rel_path)
        cluster = 'events' if rel_path.startswith('Events/') else 'places'
        kind = spec.get('kind', 'article')
        _world_map_add_node(
            nodes,
            f"page:{rel_path}",
            title,
            kind,
            cluster,
            path=rel_path,
            summary=spec.get('summary', ''),
            source_keys=spec.get('source_keys', []),
        )

    if 'Places/leeds.md' in article_specs:
        _world_map_add_edge(
            edges,
            seen_edges,
            full_name,
            'returns to and works through',
            'Leeds',
            'Leeds acts as a recurring city of return, work, and imagined stability.',
            source_keys=['journal_doc', 'finance_workbook'],
        )
    if 'Places/london.md' in article_specs:
        _world_map_add_edge(
            edges,
            seen_edges,
            full_name,
            'passed through as a life phase',
            'London',
            'Work, optimisation, dates, and hostel living; useful but not home.',
            source_keys=['journal_doc'],
        )
        _world_map_add_node(
            nodes,
            'residence:london-hostels',
            'London hostels',
            'residence',
            'residences',
            summary='Temporary hostel living documented in the London phase.',
            source_keys=['journal_doc'],
        )
        _world_map_add_edge(
            edges,
            seen_edges,
            full_name,
            'lived temporarily in',
            'London hostels',
            'Journal-documented hostel period in London.',
            source_keys=['journal_doc'],
        )
        _world_map_add_edge(
            edges,
            seen_edges,
            'London hostels',
            'sit within',
            'London',
            'Temporary urban base during the London phase.',
            source_keys=['journal_doc'],
        )
    if 'Events/austria-leg-of-the-2019-run.md' in article_specs:
        _world_map_add_node(
            nodes,
            'place:austria',
            'Austria',
            'place',
            'places',
            summary='Country-level setting for the strongest currently absorbed event page.',
            source_keys=['journal_doc'],
        )
        _world_map_add_edge(
            edges,
            seen_edges,
            full_name,
            'experienced as a defining event',
            'Austria leg of the 2019 run',
            'A spiritual and physical high point in the journal.',
            source_keys=['journal_doc'],
        )
        _world_map_add_edge(
            edges,
            seen_edges,
            'Austria leg of the 2019 run',
            'illuminates',
            'North star',
            'The event carries themes of ceremony, hardship, and meaning.',
            source_keys=['journal_doc', 'soul_doc'],
        )

    for person in people_directory:
        name = person.get('name', '')
        slug = person.get('slug', _slugify(name))
        if not name or not slug:
            continue
        path = person_paths.get(slug, '')
        _world_map_add_node(
            nodes,
            f"person:{slug}",
            name,
            'person',
            'people',
            path=path,
            summary=person.get('articleLead', person.get('summary', '')),
            source_keys=['journal_doc'],
        )
        relation = person.get('relationToDan') or person.get('primaryRelationship') or 'knows'
        note_parts = [
            person.get('articleLead', ''),
            f"Mentions: {person.get('mentions', 0)}." if person.get('mentions') else '',
        ]
        _world_map_add_edge(
            edges,
            seen_edges,
            full_name,
            relation.replace('-', ' '),
            name,
            ' '.join(part for part in note_parts if part).strip(),
            source_keys=['journal_doc', 'profile_ingests' if person.get('profileContext') else ''],
        )

    for influence in public_influences:
        name = _collapse_whitespace(influence.get('name', ''))
        if not name:
            continue
        _world_map_add_node(
            nodes,
            f"influence-public:{_slugify(name)}",
            name,
            'influence',
            'philosophy',
            summary='Named public influence from the profile-ingest layer.',
            source_keys=['profile_ingests'],
        )
        _world_map_add_edge(
            edges,
            seen_edges,
            full_name,
            'is influenced by',
            name,
            influence.get('notes', '') or 'Named as a public influence in the profile layer.',
            source_keys=['profile_ingests'],
        )

    for influence in personal_influences:
        name = _collapse_whitespace(influence.get('name', ''))
        if not name:
            continue
        matched_slug = person_slug_by_name.get(name.lower(), '')
        if not matched_slug:
            _world_map_add_node(
                nodes,
                f"influence-personal:{_slugify(name)}",
                name,
                'influence',
                'people',
                summary='Personally known influence named in the profile-ingest layer.',
                source_keys=['profile_ingests'],
            )
        _world_map_add_edge(
            edges,
            seen_edges,
            full_name,
            'is shaped personally by',
            name,
            influence.get('whyInfluential', '') or influence.get('notes', '') or influence.get('relationship', ''),
            source_keys=['profile_ingests'],
        )

    for item in work.get('history', []) or []:
        role = _collapse_whitespace(item.get('role', ''))
        employer = _collapse_whitespace(item.get('employer', ''))
        location = _collapse_whitespace(item.get('location', ''))
        if not role:
            continue
        job_id = f"job:{_slugify(role)}-{_slugify(employer or item.get('range', ''))}-{_slugify(item.get('startIso', ''))}"
        job_title = item.get('display', '') or role
        _world_map_add_node(
            nodes,
            job_id,
            job_title,
            'job',
            'work',
            summary=item.get('range', ''),
            status='current' if item.get('isCurrent') else 'past',
            source_keys=['finance_workbook'],
        )
        _world_map_add_edge(
            edges,
            seen_edges,
            full_name,
            'works as' if item.get('isCurrent') else 'worked as',
            job_title,
            item.get('range', ''),
            source_keys=['finance_workbook'],
        )
        if employer:
            org_id = f"org:{_slugify(employer)}"
            _world_map_add_node(
                nodes,
                org_id,
                employer,
                'organization',
                'work',
                summary='Employer organisation from employment history.',
                source_keys=['finance_workbook'],
            )
            _world_map_add_edge(
                edges,
                seen_edges,
                job_title,
                'at',
                employer,
                item.get('range', ''),
                source_keys=['finance_workbook'],
            )
        if location:
            known_place_title = 'Leeds' if 'Leeds' in location and 'Places/leeds.md' in article_specs else 'London' if 'London' in location and 'Places/london.md' in article_specs else location
            _world_map_add_node(
                nodes,
                f"place:{_slugify(known_place_title)}",
                known_place_title,
                'place',
                'places',
                path='Places/leeds.md' if known_place_title == 'Leeds' and 'Places/leeds.md' in article_specs else 'Places/london.md' if known_place_title == 'London' and 'Places/london.md' in article_specs else '',
                summary='Work location surfaced from employment history.',
                source_keys=['finance_workbook'],
            )
            _world_map_add_edge(
                edges,
                seen_edges,
                job_title,
                'based in',
                known_place_title,
                location,
                source_keys=['finance_workbook'],
            )

    current_summary = work.get('summary', {}) if isinstance(work.get('summary', {}), dict) else {}
    if current_summary.get('office_address'):
        _world_map_add_edge(
            edges,
            seen_edges,
            work.get('jobTitle') or work.get('linkedinTitle') or 'Current role',
            'works from',
            current_summary.get('office_address'),
            work.get('workingPattern', ''),
            source_keys=['finance_workbook'],
        )

    for position in (market.get('positions', []) or []):
        label = position.get('label') or position.get('item') or position.get('symbol')
        if not label:
            continue
        holding_id = f"holding:{_slugify(position.get('symbol') or label)}"
        _world_map_add_node(
            nodes,
            holding_id,
            label,
            'holding',
            'money',
            summary=position.get('account', ''),
            source_keys=['finance_workbook', 'market_cache', 'life_history'],
        )
        _world_map_add_edge(
            edges,
            seen_edges,
            'Financial position',
            'currently includes',
            label,
            f"{position.get('account', '')} holding worth about £{float(position.get('latestValueGbp', 0) or 0):,.2f}.",
            source_keys=['finance_workbook', 'market_cache', 'life_history'],
        )

    graph_nodes = sorted(nodes.values(), key=lambda item: (item.get('cluster', ''), item.get('kind', ''), item.get('title', '').lower()))
    cluster_counts = Counter(node.get('cluster', 'other') for node in graph_nodes)

    hub_paths = [
        domain_paths.get('overview', ''),
        domain_paths.get('people', ''),
        domain_paths.get('journal', ''),
        domain_paths.get('employment', ''),
        domain_paths.get('investments', ''),
        domain_paths.get('soul', ''),
        concept_paths.get('north-star', ''),
        concept_paths.get('career-arc', ''),
        concept_paths.get('financial-position', ''),
        concept_paths.get('influence-map', ''),
        'Places/leeds.md' if 'Places/leeds.md' in article_specs else '',
        'Places/london.md' if 'Places/london.md' in article_specs else '',
        'Events/austria-leg-of-the-2019-run.md' if 'Events/austria-leg-of-the-2019-run.md' in article_specs else '',
    ]
    hub_pages = [
        _site_link_record(path, os.path.splitext(os.path.basename(path))[0].replace('-', ' ').title(), '')
        for path in hub_paths
        if path
    ]

    # Improve labels for known page-backed hubs.
    for item in hub_pages:
        if item['path'] == domain_paths.get('overview', ''):
            item['title'] = 'Overview'
            item['note'] = 'Top-level orientation page'
        elif item['path'] == domain_paths.get('people', ''):
            item['title'] = 'People'
            item['note'] = 'Canonical relationship layer'
        elif item['path'] == domain_paths.get('journal', ''):
            item['title'] = 'Journal'
            item['note'] = 'Live pattern and evidence stream'
        elif item['path'] == domain_paths.get('employment', ''):
            item['title'] = 'Employment'
            item['note'] = 'Work chapters and compensation'
        elif item['path'] == domain_paths.get('investments', ''):
            item['title'] = 'Investments'
            item['note'] = 'Held positions and market history'
        elif item['path'] == domain_paths.get('soul', ''):
            item['title'] = 'Soul'
            item['note'] = 'Values, becoming, and north star'
        elif item['path'] == concept_paths.get('north-star', ''):
            item['title'] = 'North star'
            item['note'] = 'Underlying orienting frame'
        elif item['path'] == concept_paths.get('career-arc', ''):
            item['title'] = 'Career arc'
            item['note'] = 'Longer work narrative'
        elif item['path'] == concept_paths.get('financial-position', ''):
            item['title'] = 'Financial position'
            item['note'] = 'Money picture across cash, liabilities, and holdings'
        elif item['path'] == concept_paths.get('influence-map', ''):
            item['title'] = 'Influence map'
            item['note'] = 'Public and personal shapers'
        elif item['path'] == 'Places/leeds.md':
            item['title'] = 'Leeds'
            item['note'] = 'Recurring city of return and work'
        elif item['path'] == 'Places/london.md':
            item['title'] = 'London'
            item['note'] = 'Formative but insufficient optimisation phase'
        elif item['path'] == 'Events/austria-leg-of-the-2019-run.md':
            item['title'] = 'Austria leg of the 2019 run'
            item['note'] = 'Journal-defined spiritual high point'

    relationship_threads = [
        {
            'from': full_name,
            'relation': (person.get('relationToDan') or person.get('primaryRelationship') or 'knows').replace('-', ' '),
            'to': person.get('name', ''),
            'note': ' '.join([
                person.get('articleLead', ''),
                f"Mentions: {person.get('mentions', 0)}." if person.get('mentions') else '',
            ]).strip(),
        }
        for person in people_directory
        if person.get('name')
    ]

    work_threads = []
    for item in work.get('history', []) or []:
        role = item.get('display', '') or item.get('role', '')
        if not role:
            continue
        work_threads.append({
            'from': full_name,
            'relation': 'works as' if item.get('isCurrent') else 'worked as',
            'to': role,
            'note': item.get('range', ''),
        })
        if item.get('employer'):
            work_threads.append({
                'from': role,
                'relation': 'at',
                'to': item.get('employer', ''),
                'note': item.get('range', ''),
            })
        if item.get('location'):
            work_threads.append({
                'from': role,
                'relation': 'based in',
                'to': item.get('location', ''),
                'note': item.get('range', ''),
            })

    money_threads = [
        {
            'from': full_name,
            'relation': 'understands money through',
            'to': 'Financial position',
            'note': 'Concept page over cash, liabilities, and portfolio.',
        }
    ]
    for position in (market.get('positions', []) or []):
        label = position.get('label') or position.get('item') or position.get('symbol')
        if not label:
            continue
        money_threads.append({
            'from': 'Financial position',
            'relation': 'currently includes',
            'to': label,
            'note': f"{position.get('account', '')} holding worth about £{float(position.get('latestValueGbp', 0) or 0):,.2f}.",
        })

    residence_threads = []
    if current_living:
        residence_threads.append({
            'from': full_name,
            'relation': 'currently lives at',
            'to': current_living.split('(')[0].strip().rstrip(','),
            'note': current_living[current_living.find('(') + 1:current_living.rfind(')')].strip() if '(' in current_living and ')' in current_living else '',
        })
        if current_residence_place:
            residence_threads.append({
                'from': current_living.split('(')[0].strip().rstrip(','),
                'relation': 'sits within',
                'to': current_residence_place,
                'note': 'Broader place parsed from the current residence record.',
            })
    if hometown:
        residence_threads.append({
            'from': full_name,
            'relation': 'comes from',
            'to': hometown,
            'note': 'Recorded as hometown in the workbook layer.',
        })
    if 'Places/london.md' in article_specs:
        residence_threads.append({
            'from': full_name,
            'relation': 'lived temporarily in',
            'to': 'London hostels',
            'note': 'Journal-documented hostel period in London.',
        })

    place_threads = []
    if 'Places/leeds.md' in article_specs:
        place_threads.append({
            'from': full_name,
            'relation': 'returns to and works through',
            'to': 'Leeds',
            'note': 'City of return, work, memory, and future-planning.',
        })
    if 'Places/london.md' in article_specs:
        place_threads.append({
            'from': full_name,
            'relation': 'passed through as a life phase',
            'to': 'London',
            'note': 'Formative work and optimisation phase that never fully felt like home.',
        })
    if 'Events/austria-leg-of-the-2019-run.md' in article_specs:
        place_threads.append({
            'from': 'Austria leg of the 2019 run',
            'relation': 'sits within',
            'to': 'Austria',
            'note': 'Country-level setting is present, but a dedicated Austria place page does not exist yet.',
        })

    philosophy_threads = [
        {
            'from': full_name,
            'relation': 'orients around',
            'to': 'North star',
            'note': soul.get('question', ''),
        },
        {
            'from': full_name,
            'relation': 'draws language from',
            'to': 'Quote bank',
            'note': 'Quotes act as tonal and philosophical scaffolding.',
        },
        {
            'from': full_name,
            'relation': 'keeps taste in',
            'to': 'Media canon',
            'note': 'Books, films, and television act as reference context.',
        },
        {
            'from': full_name,
            'relation': 'is shaped by',
            'to': 'Influence map',
            'note': 'Public thinkers and personal influences sit together here.',
        },
    ]
    if creative.get('entries'):
        philosophy_threads.append({
            'from': full_name,
            'relation': 'develops thought through',
            'to': 'Creative practice',
            'note': 'Creative fragments and essays surface recurring ideas.',
        })

    influence_threads = [
        {
            'from': full_name,
            'relation': 'is influenced by',
            'to': _collapse_whitespace(item.get('name', '')),
            'note': item.get('notes', '') or 'Named as a public influence in the profile layer.',
        }
        for item in public_influences
        if _collapse_whitespace(item.get('name', ''))
    ]
    influence_threads.extend([
        {
            'from': full_name,
            'relation': 'is shaped personally by',
            'to': _collapse_whitespace(item.get('name', '')),
            'note': item.get('whyInfluential', '') or item.get('notes', '') or item.get('relationship', ''),
        }
        for item in personal_influences
        if _collapse_whitespace(item.get('name', ''))
    ])

    event_threads = []
    if 'Events/austria-leg-of-the-2019-run.md' in article_specs:
        event_threads.extend([
            {
                'from': full_name,
                'relation': 'experienced as a defining event',
                'to': 'Austria leg of the 2019 run',
                'note': 'Spiritual and physical high point in the run journal.',
            },
            {
                'from': 'Austria leg of the 2019 run',
                'relation': 'illuminates',
                'to': 'North star',
                'note': 'Ceremony, sacredness, mountains, hardship, and meaning.',
            },
        ])

    page_candidates = [
        {
            'title': 'Residence history',
            'kind': 'timeline',
            'reason': 'Current home and London hostel period are visible, but the full house/flat timeline is not yet compiled.',
            'status': 'needed',
        },
        {
            'title': 'Employer pages',
            'kind': 'organization',
            'reason': 'National Wealth Fund, Leeds Building Society, Sky Betting & Gaming, Vanquis Bank, and Lloyds Banking Group are central but do not yet have dedicated pages.',
            'status': 'needed',
        },
        {
            'title': 'Investment conviction pages',
            'kind': 'holding',
            'reason': 'Current positions exist as portfolio holdings, but not yet as pages explaining why they matter.',
            'status': 'needed',
        },
        {
            'title': 'Home and place pages below city level',
            'kind': 'place',
            'reason': 'Castle Bank/Castle Bolton and other lived-in locations need a proper place/residence layer once evidence is assembled.',
            'status': 'needed',
        },
    ]

    missing_nodes = [
        'Residence history is incomplete beyond the current Castle Bank location and the London hostel period.',
        'Employer and organisation pages do not yet exist for the major work institutions in the career timeline.',
        'Investment pages do not yet capture conviction, thesis, or personal meaning for each holding.',
        'Below-city place pages such as homes, offices, and repeated venues have not been normalised yet.',
    ]

    return {
        'summary': {
            'node_count': len(graph_nodes),
            'edge_count': len(edges),
            'cluster_counts': dict(cluster_counts),
            'current_role': work.get('jobTitle') or work.get('linkedinTitle', ''),
            'current_employer': work.get('employer', ''),
            'current_residence': current_living,
            'current_portfolio_value_gbp': round(float(market.get('marketValueGbp', 0) or 0), 2),
            'current_cash_gbp': round(float(money.get('cashTotal', 0) or 0), 2),
        },
        'hub_pages': hub_pages,
        'residence_threads': residence_threads,
        'place_threads': place_threads,
        'relationship_threads': relationship_threads,
        'work_threads': work_threads,
        'money_threads': money_threads,
        'philosophy_threads': philosophy_threads,
        'influence_threads': influence_threads,
        'event_threads': event_threads,
        'page_candidates': page_candidates,
        'missing_nodes': missing_nodes,
        'nodes': graph_nodes,
        'edges': edges,
    }


def _build_ai_context_markdown(snapshot):
    docs_dir = snapshot.get('documentsDir', '')
    if not docs_dir:
        return {
            'directory': '',
            'files': [],
        }

    scratchpad = _build_scratchpad_snapshot(docs_dir)
    memory = _build_memory_snapshot(docs_dir)
    source_catalog = _knowledge_source_catalog(docs_dir)
    fashion = _build_fashion_context(snapshot.get('onebag', {}))
    findings, recommendations = _build_health_checks(snapshot, scratchpad)
    concepts = _build_concept_specs(snapshot, scratchpad, memory, fashion)

    domain_titles = {
        'overview': 'Overview',
        'profile': 'Profile',
        'onebag': 'OneBag',
        'fashion': 'Fashion',
        'running': 'Running',
        'people': 'People',
        'finances': 'Finances',
        'investments': 'Investments',
        'creative': 'Creative',
        'journal': 'Journal',
        'employment': 'Employment',
        'soul': 'Soul',
    }
    domain_paths = {slug: f'Domains/{slug}.md' for slug in domain_titles}
    concept_paths = {slug: f'Concepts/{slug}.md' for slug in concepts}
    people_directory = snapshot.get('people', {}).get('directory', []) or []
    person_paths = {person.get('slug', _slugify(person.get('name', 'person'))): f"People/{person.get('slug', _slugify(person.get('name', 'person')))}.md" for person in people_directory}

    domain_sources = {
        'overview': ['inventory_workbook', 'finance_workbook', 'journal_doc', 'soul_doc', 'profile_ingests', 'memory_state', 'scratchpad_state'],
        'profile': ['profile_ingests'],
        'onebag': ['inventory_workbook', 'finance_workbook'],
        'fashion': ['inventory_workbook', 'memory_state', 'running_state'],
        'running': ['running_state', 'memory_state', 'scratchpad_state'],
        'people': ['journal_doc'],
        'finances': ['finance_workbook'],
        'investments': ['finance_workbook', 'market_cache', 'life_history'],
        'creative': ['creative_doc', 'journal_doc'],
        'journal': ['journal_doc', 'journal_captures'],
        'employment': ['finance_workbook'],
        'soul': ['soul_doc', 'journal_doc'],
    }
    domain_concepts = {
        'overview': ['capsule-wardrobe', 'one-shoe-strategy', 'financial-position', 'career-arc', 'north-star', 'media-canon', 'quote-bank', 'influence-map'],
        'profile': ['media-canon', 'quote-bank', 'influence-map'],
        'onebag': ['capsule-wardrobe', 'one-shoe-strategy'],
        'fashion': ['capsule-wardrobe', 'one-shoe-strategy'],
        'running': ['one-shoe-strategy', 'capsule-wardrobe'],
        'people': ['north-star', 'creative-practice', 'influence-map'],
        'finances': ['financial-position'],
        'investments': ['financial-position'],
        'creative': ['creative-practice', 'media-canon', 'quote-bank'],
        'journal': ['north-star', 'creative-practice', 'influence-map'],
        'employment': ['career-arc'],
        'soul': ['north-star', 'quote-bank', 'influence-map'],
    }

    overview_map = {
        domain_paths[slug]: f'{title} domain'
        for slug, title in domain_titles.items()
    }
    overview_map.update({
        concept_paths[slug]: f"{spec['title']} concept"
        for slug, spec in concepts.items()
    })
    overview_map.update({
        'index.md': 'Primary wiki index',
        'schema.md': 'Wiki maintenance rules',
        'log.md': 'Chronological wiki log',
        'Indexes/current-focus.md': 'Current focus index',
        'Indexes/world-map.md': 'World map index',
        'Indexes/health-check.md': 'Health-check index',
        'Indexes/source-map.md': 'Source map',
        'Indexes/people-index.md': 'People index',
        'Outputs/current-brief.md': 'Current brief output',
    })
    overview_map.update({
        path: 'Person page'
        for path in person_paths.values()
    })

    domain_bodies = {
        'overview': _overview_markdown(snapshot, overview_map),
        'profile': _profile_markdown(snapshot),
        'onebag': _onebag_markdown(snapshot),
        'fashion': _fashion_markdown(snapshot),
        'running': _running_markdown(snapshot),
        'people': _people_markdown(snapshot),
        'finances': _finances_markdown(snapshot),
        'investments': _investments_markdown(snapshot),
        'creative': _creative_markdown(snapshot),
        'journal': _journal_markdown(snapshot),
        'employment': _employment_markdown(snapshot),
        'soul': _soul_markdown(snapshot),
    }

    domain_backlinks = defaultdict(list)
    for concept_slug, spec in concepts.items():
        for domain_slug in spec.get('related_domains', []):
            if domain_slug in domain_paths:
                domain_backlinks[domain_slug].append(concept_paths[concept_slug])

    article_specs = _build_supporting_article_specs(snapshot, domain_paths, concept_paths, person_paths, people_directory)
    world_map = _build_world_map(snapshot, domain_paths, concept_paths, person_paths, people_directory, article_specs)
    extracted = _build_extracted_layer(
        snapshot,
        domain_titles,
        domain_paths,
        domain_sources,
        domain_concepts,
        concepts,
        concept_paths,
        people_directory,
        person_paths,
        article_specs,
    )
    world_map_source_keys = sorted({
        key
        for collection in (world_map.get('nodes', []), world_map.get('edges', []))
        for item in collection
        for key in item.get('source_keys', [])
        if key
    })
    current_living = _collapse_whitespace(snapshot.get('profile', {}).get('currentLiving', ''))
    current_work_summary = (snapshot.get('work', {}) or {}).get('summary', {}) if isinstance((snapshot.get('work', {}) or {}).get('summary', {}), dict) else {}
    north_star_question = _collapse_whitespace((concepts.get('north-star', {}) or {}).get('structured_state', {}).get('question', ''))
    world_map_source_notes = []
    if current_living:
        world_map_source_notes.append({
            'date': '',
            'title': 'Current residence',
            'snippet': current_living,
        })
    if world_map.get('summary', {}).get('current_role'):
        work_snippet = f"{world_map['summary'].get('current_role', '')} at {world_map['summary'].get('current_employer', '')}."
        if current_work_summary.get('office_address'):
            work_snippet = f"{work_snippet} Office: {current_work_summary.get('office_address', '')}."
        if current_work_summary.get('working_pattern'):
            work_snippet = f"{work_snippet} Pattern: {current_work_summary.get('working_pattern', '')}."
        world_map_source_notes.append({
            'date': '',
            'title': 'Current work anchor',
            'snippet': work_snippet.strip(),
        })
    if snapshot.get('marketPerformance', {}).get('ok'):
        world_map_source_notes.append({
            'date': snapshot.get('marketPerformance', {}).get('asOfDate', ''),
            'title': 'Portfolio anchor',
            'snippet': (
                f"Tracked live market value £{float(snapshot.get('marketPerformance', {}).get('marketValueGbp', 0) or 0):,.2f} "
                f"across {len(snapshot.get('marketPerformance', {}).get('positions', []) or [])} current holdings."
            ),
        })
    if north_star_question:
        world_map_source_notes.append({
            'date': '',
            'title': 'North star question',
            'snippet': north_star_question,
        })
    for rel_path in ('Places/leeds.md', 'Places/london.md', 'Events/austria-leg-of-the-2019-run.md'):
        source_note = ((article_specs.get(rel_path) or {}).get('source_notes') or [None])[0]
        if source_note:
            world_map_source_notes.append({
                'date': source_note.get('date', ''),
                'title': source_note.get('title', '') or (article_specs.get(rel_path) or {}).get('title', ''),
                'snippet': source_note.get('snippet', ''),
            })

    leeds_source_notes = (article_specs.get('Places/leeds.md') or {}).get('source_notes') or []
    london_source_notes = (article_specs.get('Places/london.md') or {}).get('source_notes') or []
    austria_source_notes = (article_specs.get('Events/austria-leg-of-the-2019-run.md') or {}).get('source_notes') or []
    world_map_quote_anchors = [
        f"Leeds: \"{_truncate_text(leeds_source_notes[0].get('snippet', ''), limit=144)}\"" if leeds_source_notes else '',
        f"London: \"{_truncate_text(london_source_notes[3].get('snippet', ''), limit=144)}\"" if len(london_source_notes) > 3 else '',
        f"Austria: \"{_truncate_text(austria_source_notes[2].get('snippet', ''), limit=144)}\"" if len(austria_source_notes) > 2 else '',
        f"North star: \"{_truncate_text(north_star_question, limit=144)}\"" if north_star_question else '',
    ]
    world_map_highlights = [
        f"Current home base is recorded as {current_living}." if current_living else '',
        (
            f"Current work anchor is {world_map.get('summary', {}).get('current_role', '')} "
            f"at {world_map.get('summary', {}).get('current_employer', '')}."
        ).strip()
        if world_map.get('summary', {}).get('current_role')
        else '',
        (
            f"Live investments currently track £{float(world_map.get('summary', {}).get('current_portfolio_value_gbp', 0) or 0):,.2f} "
            f"of market value and £{float(world_map.get('summary', {}).get('current_cash_gbp', 0) or 0):,.2f} in cash."
        ),
        (
            f"The first graph pass currently contains {world_map.get('summary', {}).get('node_count', 0)} nodes "
            f"and {world_map.get('summary', {}).get('edge_count', 0)} edges."
        ),
        'Leeds, London, and the Austria leg of the 2019 run now act as typed place-event anchors rather than buried references.',
    ]
    world_map_related_paths = list(dict.fromkeys([
        'index.md',
        'Domains/overview.md',
        'Domains/people.md',
        'Domains/employment.md',
        'Domains/investments.md',
        'Domains/soul.md',
        'Concepts/north-star.md',
        'Concepts/career-arc.md',
        'Concepts/financial-position.md',
        'Concepts/influence-map.md',
        *[item.get('path', '') for item in world_map.get('hub_pages', []) if item.get('path')],
        'Indexes/source-map.md',
        'Indexes/backlinks.md',
    ]))
    world_map_body = '\n'.join([
        '# World map',
        '',
        'This page is the first-pass knowledge graph for Akibwa. It maps the strongest current links across people, places, homes, work, money, philosophy, and defining events so the wiki can expand without losing the underlying shape.',
        '',
        'It is intentionally evidence-backed and specific. The links here are compiled from the normalized workbook, journal, profile-ingest, memory, and absorbed article layers rather than from a loose summary or a raw quote dump.',
        '',
        '## current_shape',
        _bullet_markdown([
            f"Current home base: {current_living}." if current_living else '',
            (
                f"Current work anchor: {world_map.get('summary', {}).get('current_role', '')} "
                f"at {world_map.get('summary', {}).get('current_employer', '')}."
            ).strip()
            if world_map.get('summary', {}).get('current_role')
            else '',
            (
                f"Current live portfolio value: £{float(world_map.get('summary', {}).get('current_portfolio_value_gbp', 0) or 0):,.2f}"
                + (
                    f" as of {snapshot.get('marketPerformance', {}).get('asOfDate', '')}."
                    if snapshot.get('marketPerformance', {}).get('asOfDate')
                    else '.'
                )
            )
            if snapshot.get('marketPerformance', {}).get('ok')
            else '',
            (
                f"Current graph size: {world_map.get('summary', {}).get('node_count', 0)} nodes and "
                f"{world_map.get('summary', {}).get('edge_count', 0)} edges."
            ),
        ]),
        '',
        '## hub_pages',
        _bullet_markdown([
            f"{_relative_markdown_link('Indexes/world-map.md', item.get('path', ''), item.get('title', ''))} — {item.get('note', '')}".rstrip(' —')
            for item in world_map.get('hub_pages', [])
            if item.get('path')
        ]),
        '',
        '## cluster_reads',
        f"People: {len(world_map.get('relationship_threads', []) or [])} typed relationship threads currently connect Daniel to family, friends, mentors, work ties, and romantic figures.",
        '',
        'Places and homes: the graph currently anchors the recorded current home, its broader place node, the hometown record, the London hostel phase, and the page-backed Leeds and London city articles.',
        '',
        f"Work: {len(snapshot.get('work', {}).get('history', []) or [])} workbook roles currently form the career chain from Lloyds Banking Group through Vanquis Bank, Sky Betting & Gaming, Leeds Building Society, and National Wealth Fund.",
        '',
        f"Money: the graph currently links Financial position to {len(snapshot.get('marketPerformance', {}).get('positions', []) or [])} live holdings while keeping cash and liabilities inside the same frame.",
        '',
        'Philosophy: North star, Influence map, Media canon, Quote bank, and Creative practice act as the worldview spine, while named public and personal influences are now promoted into explicit relationship threads.',
        '',
        'Events: the Austria leg of the 2019 run is currently the strongest event node because it bridges travel, embodiment, hardship, and soul-language in a single absorbed page.',
        '',
        '## quote_anchors',
        _bullet_markdown(world_map_quote_anchors),
        '',
        '## next_compilation_targets',
        _bullet_markdown([
            f"{item.get('title', '')} ({item.get('kind', '')}) — {item.get('reason', '')}"
            for item in world_map.get('page_candidates', [])
        ] + (world_map.get('missing_nodes', []) or [])),
        '',
    ])

    content_map = {}
    for slug, title in domain_titles.items():
        rel_path = domain_paths[slug]
        body = domain_bodies[slug]
        related_paths = [concept_paths[key] for key in domain_concepts.get(slug, []) if key in concept_paths]
        content_map[rel_path] = _decorate_knowledge_markdown(
            snapshot,
            rel_path,
            title,
            'domain',
            domain_sources.get(slug, []),
            related_paths,
            domain_backlinks.get(slug, []),
            body,
            source_catalog,
            extracted_rel_path=_extracted_rel_path_for_article(rel_path),
        )

    for slug, spec in concepts.items():
        rel_path = concept_paths[slug]
        content_map[rel_path] = _concept_markdown(
            snapshot,
            rel_path,
            slug,
            spec,
            domain_paths,
            source_catalog,
            extracted_rel_path=_extracted_rel_path_for_article(rel_path),
        )

    for person in people_directory:
        rel_path = person_paths.get(person.get('slug', ''))
        if rel_path:
            content_map[rel_path] = _person_markdown(
                snapshot,
                rel_path,
                person,
                source_catalog,
                extracted_rel_path=_extracted_rel_path_for_article(rel_path),
            )

    article_paths = sorted(article_specs)
    for rel_path, spec in article_specs.items():
        content_map[rel_path] = _article_markdown(
            snapshot,
            rel_path,
            spec,
            source_catalog,
            extracted_rel_path=_extracted_rel_path_for_article(rel_path),
        )
    content_map['Indexes/world-map.md'] = _decorate_knowledge_markdown(
        snapshot,
        'Indexes/world-map.md',
        'World map',
        'index',
        world_map_source_keys,
        world_map_related_paths,
        ['index.md', 'Indexes/backlinks.md', 'Indexes/source-map.md'],
        world_map_body,
        source_catalog,
    )

    page_catalog = []
    for slug, title in domain_titles.items():
        page_catalog.append({
            'path': domain_paths[slug],
            'title': title,
            'kind': 'domain',
            'summary': f'{title} domain page compiled from the current Akibwa dashboard sources.',
        })
    for slug, spec in concepts.items():
        page_catalog.append({
            'path': concept_paths[slug],
            'title': spec['title'],
            'kind': 'concept',
            'summary': spec['summary'],
        })

    index_body = [
        '# Index',
        '',
        'Primary catalog for the Akibwa machine-maintained wiki. Start here before drilling into individual pages.',
        '',
        '## domains',
        '\n'.join(
            f"- {_relative_markdown_link('index.md', domain_paths[slug], domain_titles[slug])} — {domain_titles[slug]} life surface."
            for slug in domain_titles
        ),
        '',
        '## concepts',
        '\n'.join(
            f"- {_relative_markdown_link('index.md', concept_paths[slug], spec['title'])} — {spec['summary']}"
            for slug, spec in concepts.items()
        ),
        '',
        '## articles',
        '\n'.join(
            f"- {_relative_markdown_link('index.md', rel_path, spec.get('title', rel_path))} — {spec.get('summary', '')}"
            for rel_path, spec in article_specs.items()
        ) or '- No absorbed article pages have been compiled yet.',
        '',
        '## people',
        '\n'.join(
            f"- {_relative_markdown_link('index.md', path, person.get('name', 'Unknown person'))} — {person.get('primaryRelationship', 'relationship')}"
            for person, path in zip(people_directory[:24], [person_paths[p.get('slug', '')] for p in people_directory[:24]])
            if path
        ) or '- No person pages have been compiled yet.',
        '',
        '## indexes',
        '\n'.join([
            f"- {_relative_markdown_link('index.md', 'Indexes/current-focus.md', 'Current focus')}",
            f"- {_relative_markdown_link('index.md', 'Indexes/world-map.md', 'World map')}",
            f"- {_relative_markdown_link('index.md', 'Indexes/source-map.md', 'Source map')}",
            f"- {_relative_markdown_link('index.md', 'Indexes/health-check.md', 'Health check')}",
            f"- {_relative_markdown_link('index.md', 'Indexes/domain-index.md', 'Domain index')}",
            f"- {_relative_markdown_link('index.md', 'Indexes/concept-index.md', 'Concept index')}",
            f"- {_relative_markdown_link('index.md', 'Indexes/article-index.md', 'Article index')}",
            f"- {_relative_markdown_link('index.md', 'Indexes/people-index.md', 'People index')}",
            f"- {_relative_markdown_link('index.md', 'Indexes/backlinks.md', 'Backlinks')}",
            f"- {_relative_markdown_link('index.md', 'schema.md', 'Schema')}",
            f"- {_relative_markdown_link('index.md', 'log.md', 'Log')}",
        ]),
        '',
        '## extracted_layer',
        _bullet_markdown([
            f"`{path}`" for path in extracted.get('entrypoints', [])
        ]),
        '',
    ]
    content_map['index.md'] = '\n'.join(index_body)

    readme_body = [
        '# Akibwa Knowledge Base',
        '',
        'Persistent, machine-oriented wiki compiled from private Akibwa sources.',
        '',
        'Start with [index.md](index.md).',
        '',
        '## layers',
        '- `Raw/` contains or will contain immutable source material and ingest manifests.',
        '- `Extracted/` contains machine-readable normalized records for article pages.',
        '- `AI Context/` is the compiled markdown wiki owned by the LLM layer.',
        '- `State/` contains machine state used to keep the wiki and dashboard coherent over time.',
        '',
        '## current_entry_points',
        '\n'.join([
            f"- {_relative_markdown_link('README.md', 'index.md', 'index.md')}",
            f"- {_relative_markdown_link('README.md', 'Indexes/article-index.md', 'article-index.md')}",
            f"- {_relative_markdown_link('README.md', 'Indexes/current-focus.md', 'current-focus.md')}",
            f"- {_relative_markdown_link('README.md', 'Indexes/world-map.md', 'world-map.md')}",
            f"- {_relative_markdown_link('README.md', 'Indexes/health-check.md', 'health-check.md')}",
            f"- {_relative_markdown_link('README.md', 'Outputs/current-brief.md', 'current-brief.md')}",
            '- `../Extracted/index.json`',
        ]),
        '',
    ]
    content_map['README.md'] = '\n'.join(readme_body)

    schema_body = [
        '# Schema',
        '',
        'Rules for maintaining the Akibwa knowledge base as an LLM-owned wiki.',
        '',
        '## workflow',
        '- Treat raw sources as immutable. Read them, but do not rewrite them.',
        '- Treat `Extracted/` as the normalized intermediate layer between raw material and compiled wiki prose.',
        '- Treat the wiki as compiled state. Update multiple files when new information changes the synthesis.',
        '- Every meaningful answer should have two outputs: the user-facing answer and any durable wiki updates worth filing.',
        '- Start navigation from `index.md`, then drill into domain pages, concept pages, and supporting indexes.',
        '- Treat compilation as rerunnable. Better models or better prompts should be able to rebuild weak pages from the same raw inputs.',
        '- Prefer modular stages: ingest, normalize, synthesize, link, and render.',
        '',
        '## conventions',
        '- Keep pages interlinked using relative markdown links.',
        '- Give each article page an `Extracted/` record whenever possible so later recompilation can improve the prose without re-deriving everything from scratch.',
        '- Treat person pages as evidence-backed relationship records, not invented biographies.',
        '- Prefer article pages for meaningful places, events, and recurring themes rather than trapping everything inside domain summaries.',
        '- Use source maps and backlinks to preserve provenance and navigability.',
        '- Prefer adding `[TODO: ...]` gaps over inventing unsupported facts.',
        '- Update `log.md` whenever the compiled knowledge meaningfully changes.',
        '- Keep `Outputs/` for durable query products, briefs, or analyses worth reusing.',
        '- Avoid baking one model’s quirks too deeply into the stored structure. Keep enough evidence and normalized fields that later recompilation can improve article quality.',
        '',
        '## linting_targets',
        '- Contradictions between pages',
        '- Orphan or weakly linked pages',
        '- Missing concept pages for repeated themes',
        '- Stale claims superseded by newer sources',
        '- Open loops that should become explicit tracking items',
        '',
    ]
    content_map['schema.md'] = '\n'.join(schema_body)

    log_entries = []
    for item in scratchpad.get('recentRequests', []):
        raw_date = _collapse_whitespace(item.get('at', '')).split('T')[0] or snapshot.get('generatedAt', '').split('T')[0]
        log_entries.extend([
            f"## [{raw_date}] request | {item.get('request', '')}",
            '',
            f"status: {item.get('status', 'noted')}",
            f"note: {item.get('note', '') or 'No extra note recorded.'}",
            '',
        ])
    if not log_entries:
        log_entries = ['## [' + snapshot.get('generatedAt', '').split('T')[0] + '] sync | Knowledge base refresh', '', 'status: generated', 'note: The compiled wiki was regenerated from the current Akibwa sources.', '']
    content_map['log.md'] = '\n'.join(['# Log', '', 'Chronological record of meaningful Akibwa wiki updates.', ''] + log_entries)

    current_focus_body = [
        '# Current Focus',
        '',
        'Short-horizon orientation file assembled from scratchpad and durable memory.',
        '',
        '## current_intent',
        _bullet_markdown(scratchpad.get('currentIntent', [])),
        '',
        '## working_on',
        _bullet_markdown(scratchpad.get('workingOn', [])),
        '',
        '## constraints',
        _bullet_markdown(scratchpad.get('constraints', [])),
        '',
        '## preferences',
        _bullet_markdown(scratchpad.get('preferences', [])),
        '',
        '## durable_memory_brief',
        _json_markdown_block(memory.get('brief', {})),
        '',
    ]
    content_map['Indexes/current-focus.md'] = '\n'.join(current_focus_body)

    source_to_outputs = defaultdict(list)
    for slug in domain_titles:
        for source_key in domain_sources.get(slug, []):
            source_to_outputs[source_key].append(domain_paths[slug])
    for slug, spec in concepts.items():
        for source_key in spec.get('source_keys', []):
            source_to_outputs[source_key].append(concept_paths[slug])
    for path in person_paths.values():
        source_to_outputs['journal_doc'].append(path)
    for rel_path, spec in article_specs.items():
        for source_key in spec.get('source_keys', []):
            source_to_outputs[source_key].append(rel_path)
    profile_people_names = {
        _collapse_whitespace(item.get('name', '')).lower()
        for item in snapshot.get('profileLibrary', {}).get('personalInfluences', [])
    } | {
        _collapse_whitespace(item.get('name', '')).lower()
        for item in snapshot.get('profileLibrary', {}).get('researchNotes', [])
    }
    for person in people_directory:
        person_path = person_paths.get(person.get('slug', ''))
        if person_path and person.get('name', '').lower() in profile_people_names:
            source_to_outputs['profile_ingests'].append(person_path)
    source_to_outputs['scratchpad_state'].extend(['Indexes/current-focus.md', 'log.md', 'Outputs/current-brief.md'])
    source_to_outputs['memory_state'].extend(['Indexes/current-focus.md'])
    source_to_outputs['life_history'].extend(['Indexes/health-check.md'])
    source_to_outputs['market_cache'].extend(['Indexes/health-check.md'])
    for source_key in world_map_source_keys:
        source_to_outputs[source_key].append('Indexes/world-map.md')

    source_map_payload = []
    for key, record in source_catalog.items():
        source_map_payload.append({
            **record,
            'derived_files': sorted(set(source_to_outputs.get(key, []))),
        })
    content_map['Indexes/source-map.md'] = '\n'.join([
        '# Source Map',
        '',
        'Provenance map from raw/state sources into compiled wiki pages.',
        '',
        _json_markdown_block(source_map_payload),
        '',
    ])

    domain_index_payload = [
        {
            'path': domain_paths[slug],
            'title': domain_titles[slug],
            'related_concepts': [concept_paths[key] for key in domain_concepts.get(slug, []) if key in concept_paths],
            'source_keys': domain_sources.get(slug, []),
            'extracted_path': _extracted_rel_path_for_article(domain_paths[slug]),
        }
        for slug in domain_titles
    ]
    content_map['Indexes/domain-index.md'] = '\n'.join([
        '# Domain Index',
        '',
        'Catalog of domain pages, their source keys, and their concept links.',
        '',
        _json_markdown_block(domain_index_payload),
        '',
    ])

    concept_index_payload = [
        {
            'path': concept_paths[slug],
            'title': spec['title'],
            'summary': spec['summary'],
            'related_domains': [domain_paths[domain] for domain in spec.get('related_domains', []) if domain in domain_paths],
            'source_keys': spec.get('source_keys', []),
            'extracted_path': _extracted_rel_path_for_article(concept_paths[slug]),
        }
        for slug, spec in concepts.items()
    ]
    content_map['Indexes/concept-index.md'] = '\n'.join([
        '# Concept Index',
        '',
        'Catalog of synthesis pages that sit between raw data and user-facing answers.',
        '',
        _json_markdown_block(concept_index_payload),
        '',
    ])

    article_index_payload = [
        {
            'path': rel_path,
            'title': spec.get('title', ''),
            'kind': spec.get('kind', 'article'),
            'summary': spec.get('summary', ''),
            'source_keys': spec.get('source_keys', []),
            'related_paths': spec.get('related_paths', []),
            'extracted_path': _extracted_rel_path_for_article(rel_path),
        }
        for rel_path, spec in article_specs.items()
    ]
    content_map['Indexes/article-index.md'] = '\n'.join([
        '# Article Index',
        '',
        'Catalog of absorbed article pages compiled from current Akibwa documents.',
        '',
        _json_markdown_block(article_index_payload),
        '',
    ])

    people_index_payload = [
        {
            'path': person_paths.get(person.get('slug', ''), ''),
            'name': person.get('name', ''),
            'relation_to_dan': person.get('relationToDan', ''),
            'primary_relationship': person.get('primaryRelationship', ''),
            'mentions': person.get('mentions', 0),
            'aliases': person.get('aliases', []),
            'article_lead': person.get('articleLead', person.get('summary', '')),
            'extracted_path': _extracted_rel_path_for_article(person_paths.get(person.get('slug', ''), '')),
        }
        for person in people_directory
        if person_paths.get(person.get('slug', ''))
    ]
    content_map['Indexes/people-index.md'] = '\n'.join([
        '# People Index',
        '',
        'Catalog of canonical people pages compiled primarily from the journal and then enriched into stable relationship articles.',
        '',
        _json_markdown_block(people_index_payload),
        '',
    ])

    backlinks_payload = {
        'concept_to_domains': {
            concept_paths[slug]: [domain_paths[domain] for domain in spec.get('related_domains', []) if domain in domain_paths]
            for slug, spec in concepts.items()
        },
        'domain_to_concepts': {
            domain_paths[slug]: [concept_paths[key] for key in domain_concepts.get(slug, []) if key in concept_paths]
            for slug in domain_titles
        },
        'person_to_domains': {
            person_paths.get(person.get('slug', ''), ''): [domain_paths['people'], domain_paths['journal']]
            for person in people_directory
            if person_paths.get(person.get('slug', ''))
        },
        'article_to_related': {
            rel_path: spec.get('related_paths', [])
            for rel_path, spec in article_specs.items()
        },
    }
    content_map['Indexes/backlinks.md'] = '\n'.join([
        '# Backlinks',
        '',
        'Cross-reference map between domain pages and concept pages.',
        '',
        _json_markdown_block(backlinks_payload),
        '',
    ])

    content_map['Indexes/health-check.md'] = '\n'.join([
        '# Health Check',
        '',
        'Periodic lint-style findings for the current Akibwa knowledge base.',
        '',
        '## findings',
        _json_markdown_block(findings),
        '',
        '## recommended_next_moves',
        _bullet_markdown(recommendations),
        '',
    ])

    current_brief_body = [
        '# Current Brief',
        '',
        'A compact current-state output meant to be read quickly before deeper work.',
        '',
        '## current_state',
        _bullet_markdown([
            f"Primary active focus: {(scratchpad.get('workingOn') or [''])[0]}" if scratchpad.get('workingOn') else '',
            *[signal.get('text', '') for signal in snapshot.get('signals', [])[:4]],
        ]),
        '',
        '## watch_now',
        _bullet_markdown([
            f"Running recommendation: {(snapshot.get('running', {}).get('recommendation') or {}).get('pick', '')}",
            f"Current role: {(snapshot.get('work', {}).get('jobTitle') or snapshot.get('work', {}).get('linkedinTitle') or '')} at {snapshot.get('work', {}).get('employer', '')}",
            f"Portfolio value: £{snapshot.get('marketPerformance', {}).get('marketValueGbp', 0):,.2f}" if snapshot.get('marketPerformance', {}).get('ok') else '',
        ]),
        '',
        '## related_pages',
        _bullet_markdown([
            _relative_markdown_link('Outputs/current-brief.md', 'Indexes/current-focus.md', 'current-focus.md'),
            _relative_markdown_link('Outputs/current-brief.md', 'Indexes/health-check.md', 'health-check.md'),
            _relative_markdown_link('Outputs/current-brief.md', concept_paths['one-shoe-strategy'], 'One-shoe strategy'),
            _relative_markdown_link('Outputs/current-brief.md', concept_paths['capsule-wardrobe'], 'Capsule wardrobe'),
        ]),
        '',
    ]
    content_map['Outputs/README.md'] = '\n'.join([
        '# Outputs',
        '',
        'Durable generated deliverables live here. Keep useful answers, briefs, and analyses instead of letting them disappear into chat history.',
        '',
    ])
    content_map['Outputs/current-brief.md'] = '\n'.join(current_brief_body)

    site_extra_pages = [
        {
            'path': 'index.md',
            'kind': 'index',
            'title': 'Index',
            'summary': 'Primary catalog for the Akibwa machine-maintained wiki.',
            'related_paths': [
                'Domains/overview.md',
                'Indexes/world-map.md',
                'Indexes/domain-index.md',
                'Indexes/concept-index.md',
                'Indexes/article-index.md',
                'Indexes/people-index.md',
                'Outputs/current-brief.md',
            ],
            'structured_state': {
                'domains': [
                    _site_link_record(domain_paths[slug], domain_titles[slug], f'{domain_titles[slug]} domain page')
                    for slug in domain_titles
                ],
                'concepts': [
                    _site_link_record(concept_paths[slug], spec['title'], spec.get('summary', ''))
                    for slug, spec in concepts.items()
                ],
                'articles': [
                    _site_link_record(rel_path, spec.get('title', rel_path), spec.get('summary', ''))
                    for rel_path, spec in article_specs.items()
                ],
                'people': [
                    _site_link_record(
                        person_paths.get(person.get('slug', ''), ''),
                        person.get('name', 'Unknown person'),
                        person.get('primaryRelationship', 'relationship'),
                    )
                    for person in people_directory
                    if person_paths.get(person.get('slug', ''))
                ],
                'indexes': [
                    _site_link_record('Indexes/current-focus.md', 'Current focus', 'Short-horizon orientation'),
                    _site_link_record('Indexes/world-map.md', 'World map', 'First-pass typed relationship map'),
                    _site_link_record('Indexes/source-map.md', 'Source map', 'Source-to-page provenance map'),
                    _site_link_record('Indexes/health-check.md', 'Health check', 'Lint-style findings and maintenance prompts'),
                    _site_link_record('Indexes/domain-index.md', 'Domain index', 'Catalog of domain pages'),
                    _site_link_record('Indexes/concept-index.md', 'Concept index', 'Catalog of concept pages'),
                    _site_link_record('Indexes/article-index.md', 'Article index', 'Catalog of absorbed article pages'),
                    _site_link_record('Indexes/people-index.md', 'People index', 'Catalog of person pages'),
                    _site_link_record('Indexes/backlinks.md', 'Backlinks', 'Cross-reference graph'),
                ],
                'outputs': [
                    _site_link_record('Outputs/current-brief.md', 'Current brief', 'Compact current-state output'),
                ],
            },
        },
        {
            'path': 'README.md',
            'kind': 'meta',
            'title': 'Akibwa Knowledge Base',
            'summary': 'Top-level guide to the compiled knowledge base layers and entry points.',
            'related_paths': ['index.md', 'schema.md', 'Indexes/current-focus.md', 'Indexes/world-map.md', 'Outputs/current-brief.md'],
            'structured_state': {
                'layers': [
                    'Raw contains immutable source material and ingest manifests.',
                    'Extracted contains normalized machine-readable article records.',
                    'AI Context contains the compiled markdown wiki owned by the LLM layer.',
                    'State contains machine-oriented memory and operational files.',
                ],
                'current_entry_points': [
                    _site_link_record('index.md', 'index.md', 'Primary wiki index'),
                    _site_link_record('Indexes/article-index.md', 'article-index.md', 'Absorbed articles catalog'),
                    _site_link_record('Indexes/current-focus.md', 'current-focus.md', 'Short-horizon orientation'),
                    _site_link_record('Indexes/world-map.md', 'world-map.md', 'Typed relationship map'),
                    _site_link_record('Indexes/health-check.md', 'health-check.md', 'Maintenance findings'),
                    _site_link_record('Outputs/current-brief.md', 'current-brief.md', 'Reusable current-state brief'),
                ],
            },
        },
        {
            'path': 'schema.md',
            'kind': 'meta',
            'title': 'Schema',
            'summary': 'Rules for maintaining the Akibwa knowledge base as an LLM-owned wiki.',
            'related_paths': ['index.md', 'log.md', 'Indexes/source-map.md', 'Indexes/backlinks.md'],
            'structured_state': {
                'workflow': [
                    'Treat raw sources as immutable and the wiki as compiled state.',
                    'Update multiple files when new information changes the synthesis.',
                    'Start navigation from index.md, then drill into domains, concepts, and indexes.',
                    'Prefer modular stages: ingest, normalize, synthesize, link, and render.',
                ],
                'conventions': [
                    'Keep pages interlinked using relative markdown links.',
                    'Give each article page an Extracted record whenever possible.',
                    'Treat person pages as evidence-backed relationship records.',
                    'Use source maps and backlinks to preserve provenance and navigability.',
                    'Update log.md whenever the compiled knowledge meaningfully changes.',
                ],
                'linting_targets': [
                    'Contradictions between pages',
                    'Orphan or weakly linked pages',
                    'Missing concept pages for repeated themes',
                    'Stale claims superseded by newer sources',
                    'Open loops that should become explicit tracking items',
                ],
            },
        },
        {
            'path': 'log.md',
            'kind': 'meta',
            'title': 'Log',
            'summary': 'Chronological record of meaningful Akibwa wiki updates.',
            'related_paths': ['index.md', 'Indexes/current-focus.md', 'Outputs/current-brief.md'],
            'structured_state': {
                'entries': [
                    {
                        'date': _collapse_whitespace(item.get('at', '')).split('T')[0] or snapshot.get('generatedAt', '').split('T')[0],
                        'request': item.get('request', ''),
                        'status': item.get('status', 'noted'),
                        'note': item.get('note', '') or 'No extra note recorded.',
                    }
                    for item in scratchpad.get('recentRequests', [])
                ] or [
                    {
                        'date': snapshot.get('generatedAt', '').split('T')[0],
                        'request': 'Knowledge base refresh',
                        'status': 'generated',
                        'note': 'The compiled wiki was regenerated from the current Akibwa sources.',
                    }
                ],
            },
        },
        {
            'path': 'Indexes/current-focus.md',
            'kind': 'index',
            'title': 'Current focus',
            'summary': 'Short-horizon orientation file assembled from scratchpad and durable memory.',
            'related_paths': ['Domains/overview.md', 'Outputs/current-brief.md', 'Indexes/health-check.md'],
            'structured_state': {
                'current_intent': scratchpad.get('currentIntent', []),
                'working_on': scratchpad.get('workingOn', []),
                'constraints': scratchpad.get('constraints', []),
                'preferences': scratchpad.get('preferences', []),
                'durable_memory_brief': memory.get('brief', {}),
            },
        },
        {
            'path': 'Indexes/world-map.md',
            'kind': 'index',
            'title': 'World map',
            'summary': 'First-pass typed relationship map across people, places, homes, work, money, philosophy, and events.',
            'source_keys': world_map_source_keys,
            'related_paths': world_map_related_paths,
            'highlights': world_map_highlights,
            'source_notes': world_map_source_notes,
            'structured_state': {
                'summary': world_map.get('summary', {}),
                'hub_pages': world_map.get('hub_pages', []),
                'residence_threads': world_map.get('residence_threads', []),
                'place_threads': world_map.get('place_threads', []),
                'relationship_threads': world_map.get('relationship_threads', []),
                'work_threads': world_map.get('work_threads', []),
                'money_threads': world_map.get('money_threads', []),
                'philosophy_threads': world_map.get('philosophy_threads', []),
                'influence_threads': world_map.get('influence_threads', []),
                'event_threads': world_map.get('event_threads', []),
                'page_candidates': world_map.get('page_candidates', []),
                'missing_nodes': world_map.get('missing_nodes', []),
            },
        },
        {
            'path': 'Indexes/source-map.md',
            'kind': 'index',
            'title': 'Source map',
            'summary': 'Provenance map from raw and state sources into compiled wiki pages.',
            'related_paths': ['index.md', 'schema.md', 'Indexes/backlinks.md'],
            'structured_state': {
                'sources': source_map_payload,
            },
        },
        {
            'path': 'Indexes/domain-index.md',
            'kind': 'index',
            'title': 'Domain index',
            'summary': 'Catalog of domain pages, their source keys, and their concept links.',
            'related_paths': ['index.md', *[domain_paths[slug] for slug in domain_titles]],
            'structured_state': {
                'records': domain_index_payload,
            },
        },
        {
            'path': 'Indexes/concept-index.md',
            'kind': 'index',
            'title': 'Concept index',
            'summary': 'Catalog of synthesis pages that sit between raw data and user-facing answers.',
            'related_paths': ['index.md', *[concept_paths[slug] for slug in concepts]],
            'structured_state': {
                'records': concept_index_payload,
            },
        },
        {
            'path': 'Indexes/article-index.md',
            'kind': 'index',
            'title': 'Article index',
            'summary': 'Catalog of absorbed article pages compiled from current Akibwa documents.',
            'related_paths': ['index.md', *article_paths],
            'structured_state': {
                'records': article_index_payload,
            },
        },
        {
            'path': 'Indexes/people-index.md',
            'kind': 'index',
            'title': 'People index',
            'summary': 'Catalog of canonical people pages compiled from the journal and profile context.',
            'related_paths': ['index.md', *[path for path in person_paths.values() if path]],
            'structured_state': {
                'records': people_index_payload,
            },
        },
        {
            'path': 'Indexes/backlinks.md',
            'kind': 'index',
            'title': 'Backlinks',
            'summary': 'Cross-reference map between domain pages, concepts, people pages, and absorbed articles.',
            'related_paths': ['index.md', 'Indexes/domain-index.md', 'Indexes/concept-index.md', 'Indexes/article-index.md', 'Indexes/people-index.md'],
            'structured_state': backlinks_payload,
        },
        {
            'path': 'Indexes/health-check.md',
            'kind': 'index',
            'title': 'Health check',
            'summary': 'Periodic lint-style findings for the current Akibwa knowledge base.',
            'related_paths': ['Domains/overview.md', 'Domains/finances.md', 'Domains/investments.md', 'Domains/employment.md', 'Outputs/current-brief.md'],
            'structured_state': {
                'findings': findings,
                'recommended_next_moves': recommendations,
            },
        },
        {
            'path': 'Outputs/current-brief.md',
            'kind': 'output',
            'title': 'Current brief',
            'summary': 'Compact current-state output meant to be read quickly before deeper work.',
            'related_paths': [
                'Indexes/current-focus.md',
                'Indexes/health-check.md',
                concept_paths['one-shoe-strategy'],
                concept_paths['capsule-wardrobe'],
            ],
            'structured_state': {
                'current_state': [
                    f"Primary active focus: {(scratchpad.get('workingOn') or [''])[0]}" if scratchpad.get('workingOn') else '',
                    *[signal.get('text', '') for signal in snapshot.get('signals', [])[:4]],
                ],
                'watch_now': [
                    f"Running recommendation: {(snapshot.get('running', {}).get('recommendation') or {}).get('pick', '')}",
                    f"Current role: {(snapshot.get('work', {}).get('jobTitle') or snapshot.get('work', {}).get('linkedinTitle') or '')} at {snapshot.get('work', {}).get('employer', '')}",
                    f"Portfolio value: £{snapshot.get('marketPerformance', {}).get('marketValueGbp', 0):,.2f}" if snapshot.get('marketPerformance', {}).get('ok') else '',
                ],
            },
        },
    ]

    site_pages = _build_ai_context_site_pages(extracted.get('records', {}), site_extra_pages)

    target_dir = _ai_context_dir(docs_dir)
    os.makedirs(target_dir, exist_ok=True)
    _cleanup_legacy_ai_context_files(target_dir)
    _prune_generated_ai_context_files(target_dir, content_map.keys())

    for relative_path, content in content_map.items():
        _write_text_atomic(os.path.join(target_dir, relative_path), content)

    _write_support_layer_files(docs_dir, source_catalog)
    _build_private_context_markdown(snapshot)

    generated_files = sorted(
        _relative_docs_path(target_dir, os.path.join(target_dir, relative_path))
        for relative_path in content_map
    )
    return {
        'directory': target_dir,
        'files': generated_files,
        'entrypoints': ['index.md', 'schema.md', 'log.md', 'Indexes/current-focus.md', 'Indexes/world-map.md', 'Indexes/health-check.md'],
        'extracted': extracted,
        'graph': world_map,
        'sources': list(source_catalog.values()),
        'site': {
            'homePath': 'index.md',
            'pages': site_pages,
        },
    }


def _ai_context_source_paths(docs_dir):
    files = _life_files(docs_dir)
    paths = [path for path in files.values() if path and os.path.exists(path)]
    for path in (
        _running_store_path(docs_dir),
        _memory_store_path(docs_dir),
        _scratchpad_store_path(docs_dir),
        _life_history_path(docs_dir),
        _market_cache_path(docs_dir),
    ):
        if os.path.exists(path):
            paths.append(path)
    paths.extend(str(path) for path in _profile_ingest_paths(docs_dir))
    return paths


def _ai_context_signature(docs_dir):
    return {
        'date': _journal_now().date().isoformat(),
        'sources': [
            {
                'path': path,
                'mtime': round(os.path.getmtime(path), 6),
                'size': os.path.getsize(path),
            }
            for path in _ai_context_source_paths(docs_dir)
        ],
    }


def _sync_ai_context(force=False):
    docs_dir = _docs_dir()
    if not docs_dir:
        _ai_context_sync_state['lastError'] = 'Akibwa Documents folder not found.'
        return False

    signature = _ai_context_signature(docs_dir)
    readme_path = os.path.join(_ai_context_dir(docs_dir), 'README.md')
    if (
        not force
        and _ai_context_sync_state.get('signature') == signature
        and os.path.exists(readme_path)
    ):
        return False

    snapshot = _build_life_snapshot()
    if not snapshot.get('ok'):
        _ai_context_sync_state['signature'] = signature
        _ai_context_sync_state['lastError'] = snapshot.get('error', 'Unknown sync error.')
        return False

    _ai_context_sync_state['signature'] = signature
    _ai_context_sync_state['lastSyncedAt'] = snapshot.get('generatedAt', '')
    _ai_context_sync_state['lastError'] = ''
    return True


def _ai_context_sync_loop():
    import time

    while True:
        try:
            _sync_ai_context(force=False)
        except Exception as exc:
            _ai_context_sync_state['lastError'] = str(exc)
        time.sleep(AI_CONTEXT_SYNC_INTERVAL_SECONDS)


def _slugify(value):
    cleaned = re.sub(r'[^a-z0-9]+', '-', str(value or '').lower()).strip('-')
    return cleaned or 'memory'


def _truncate_text(value, limit=180):
    text = _collapse_whitespace(value)
    if len(text) <= limit:
        return text
    return text[: max(limit - 1, 1)].rstrip() + '…'


def _derive_memory_title(title, content):
    normalized_title = _collapse_whitespace(title)
    if normalized_title:
        return normalized_title[:120]

    paragraphs = _split_content_paragraphs(content)
    if not paragraphs:
        return 'Memory note'

    first = paragraphs[0]
    sentence = re.split(r'(?<=[.!?])\s+', first, maxsplit=1)[0]
    candidate = _collapse_whitespace(sentence) or _collapse_whitespace(first)
    return _truncate_text(candidate or 'Memory note', limit=96)


def _normalize_memory_category(value):
    key = _slugify(value)
    if key in MEMORY_CATEGORY_LABELS:
        return key
    return 'working-on'


def _normalize_memory_tags(value):
    raw_items = []
    if isinstance(value, str):
        raw_items = re.split(r'[,|\n]+', value)
    elif isinstance(value, (list, tuple, set)):
        raw_items = list(value)

    tags = []
    seen = set()
    for item in raw_items:
        cleaned = _collapse_whitespace(item).lower()
        if not cleaned or cleaned in seen:
            continue
        seen.add(cleaned)
        tags.append(cleaned)
    return tags[:12]


def _normalize_memory_entry(entry):
    text = '\n\n'.join(_split_content_paragraphs(entry.get('text', '')))
    if not text:
        return None

    created_at = _collapse_whitespace(entry.get('createdAt') or entry.get('created_at') or '')
    title = _derive_memory_title(entry.get('title', ''), text)
    category = _normalize_memory_category(entry.get('category', 'working-on'))
    source = _collapse_whitespace(entry.get('source', '')) or 'dashboard'
    tags = _normalize_memory_tags(entry.get('tags', []))
    entry_id = _collapse_whitespace(entry.get('id', '')) or f"memory-{_slugify(created_at or title)}"

    return {
        'id': entry_id,
        'createdAt': created_at,
        'category': category,
        'categoryLabel': MEMORY_CATEGORY_LABELS.get(category, category.title()),
        'title': title,
        'tags': tags,
        'text': text,
        'preview': _truncate_text(text, limit=180),
        'source': source,
    }


def _build_memory_line(entry):
    title = _collapse_whitespace(entry.get('title', ''))
    preview = _collapse_whitespace(entry.get('preview') or entry.get('text') or '')
    if title and preview and not preview.lower().startswith(title.lower()):
        return _truncate_text(f'{title}: {preview}', limit=220)
    return _truncate_text(preview or title, limit=220)


def _build_memory_brief(entries):
    grouped = defaultdict(list)
    for entry in entries:
        grouped[entry['category']].append(_build_memory_line(entry))

    return {
        'workingOn': grouped.get('working-on', [])[:4],
        'intent': grouped.get('intent', [])[:4],
        'constraints': grouped.get('constraint', [])[:4],
        'preferences': grouped.get('preference', [])[:4],
        'projects': grouped.get('project', [])[:4],
        'facts': grouped.get('fact', [])[:4],
    }


def _build_memory_snapshot(docs_dir):
    store = _load_memory_store(docs_dir)
    entries = []
    for raw_entry in store.get('entries', []):
        normalized = _normalize_memory_entry(raw_entry)
        if normalized:
            entries.append(normalized)

    entries.sort(key=lambda entry: entry.get('createdAt', ''), reverse=True)
    category_counts = Counter(entry['category'] for entry in entries)
    categories = [
        {
            'key': key,
            'label': MEMORY_CATEGORY_LABELS.get(key, key.title()),
            'count': count,
        }
        for key, count in category_counts.most_common()
    ]

    focus = [entry for entry in entries if entry['category'] in MEMORY_FOCUS_CATEGORIES][:5]
    if not focus:
        focus = entries[:5]

    return {
        'storagePath': _memory_store_path(docs_dir),
        'canWrite': bool(docs_dir),
        'updatedAt': store.get('updatedAt', ''),
        'entryCount': len(entries),
        'categories': categories,
        'focus': focus,
        'recent': entries[:8],
        'brief': _build_memory_brief(entries),
    }


def _create_memory_entry(docs_dir, title, content, category='working-on', tags=None, source='dashboard'):
    paragraphs = _split_content_paragraphs(content)
    if not paragraphs:
        raise RuntimeError('Memory content is empty.')

    now = _journal_now()
    normalized_text = '\n\n'.join(paragraphs)
    normalized_title = _derive_memory_title(title, normalized_text)
    normalized_category = _normalize_memory_category(category)
    normalized_tags = _normalize_memory_tags(tags)
    normalized_source = _collapse_whitespace(source) or 'dashboard'

    entry = {
        'id': f"memory-{now.strftime('%Y%m%dT%H%M%S%f')}",
        'createdAt': now.isoformat(timespec='seconds'),
        'category': normalized_category,
        'title': normalized_title,
        'tags': normalized_tags,
        'text': normalized_text,
        'source': normalized_source,
    }

    store = _load_memory_store(docs_dir)
    entries = [raw for raw in store.get('entries', []) if isinstance(raw, dict)]
    entries.append(entry)
    entries.sort(key=lambda item: str(item.get('createdAt') or ''), reverse=True)

    store['version'] = 1
    store['updatedAt'] = now.isoformat(timespec='seconds')
    store['entries'] = entries[:400]
    _write_json_atomic(_memory_store_path(docs_dir), store)

    normalized_entry = _normalize_memory_entry(entry)
    return normalized_entry or entry


def _record_life_history(docs_dir, money, onebag, work):
    history = _load_life_history(docs_dir)
    portfolio = list(history.get('portfolio') or [])
    now = _journal_now()
    today = now.strftime('%Y-%m-%d')

    snapshot = {
        'date': today,
        'timestamp': now.isoformat(timespec='seconds'),
        'investmentsTotal': round(money.get('investmentsTotal', 0), 2),
        'isaTotal': round(money.get('isaTotal', 0), 2),
        'taxableInvestTotal': round(money.get('taxableInvestTotal', 0), 2),
        'cashTotal': round(money.get('cashTotal', 0), 2),
        'trackedWeightKg': round(onebag.get('trackedWeightKg', 0), 2),
        'monthlyOutgoingsTotal': round(money.get('monthlyOutgoingsTotal', 0), 2),
        'monthlyNetPay': round(work.get('monthlyNetPay', 0), 2),
    }

    if portfolio and portfolio[-1].get('date') == today:
        portfolio[-1] = snapshot
    else:
        portfolio.append(snapshot)

    history['portfolio'] = portfolio[-180:]
    _write_json_atomic(_life_history_path(docs_dir), history)
    return history['portfolio']


def _journal_capture_capabilities(journal_doc_path, notes_dir=''):
    docx_ready = True
    docx_reason = ''
    try:
        _try_import_docx()
    except RuntimeError as exc:
        docx_ready = False
        docx_reason = str(exc)

    openai_ready = True
    openai_reason = ''
    try:
        __import__('openai')
    except ImportError:
        openai_ready = False
        openai_reason = 'Install the OpenAI Python SDK with `python3 -m pip install openai`.'

    has_api_key = bool(os.getenv('OPENAI_API_KEY'))
    can_server_transcribe = openai_ready and has_api_key and TRANSCRIBE_CLI.exists()

    transcribe_reasons = []
    if not has_api_key:
        transcribe_reasons.append('Set OPENAI_API_KEY locally to enable AI transcription.')
    if not openai_ready and openai_reason:
        transcribe_reasons.append(openai_reason)
    if not TRANSCRIBE_CLI.exists():
        transcribe_reasons.append(f'Transcribe CLI not found at {TRANSCRIBE_CLI}.')

    return {
        'canAppend': docx_ready and os.path.exists(journal_doc_path),
        'appendReason': docx_reason,
        'canCreateNoteFile': bool(notes_dir),
        'noteDirectory': notes_dir,
        'canServerTranscribe': can_server_transcribe,
        'hasOpenAIKey': has_api_key,
        'openaiInstalled': openai_ready,
        'transcribeCliPresent': TRANSCRIBE_CLI.exists(),
        'transcribeReasons': transcribe_reasons,
    }


def _append_journal_entry(journal_doc_path, title, content, source='typed'):
    if not os.path.exists(journal_doc_path):
        raise RuntimeError(f'Journal document not found: {journal_doc_path}')

    paragraphs = _split_content_paragraphs(content)
    if not paragraphs:
        raise RuntimeError('Journal entry content is empty.')

    Document = _try_import_docx()
    document = Document(journal_doc_path)

    now = _journal_now()
    normalized_title = _collapse_whitespace(title)
    if not normalized_title:
        normalized_title = 'Voice note' if source in {'browser-dictation', 'server-transcription'} else 'Jarvis capture'

    heading = f"{now.strftime('%d/%m/%Y %H:%M %Z')}: {normalized_title}"

    if document.paragraphs and document.paragraphs[-1].text.strip():
        document.add_paragraph('')
    document.add_paragraph(heading)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)

    _save_docx_atomic(document, journal_doc_path)

    return {
        'timestamp': now.isoformat(timespec='seconds'),
        'date': now.strftime('%d/%m/%Y'),
        'time': now.strftime('%H:%M %Z'),
        'title': normalized_title,
        'text': '\n\n'.join(paragraphs),
        'source': source,
    }


def _create_journal_note_file(docs_dir, title, content, source='typed'):
    notes_dir = Path(_journal_capture_dir(docs_dir))
    notes_dir.mkdir(parents=True, exist_ok=True)

    paragraphs = _split_content_paragraphs(content)
    if not paragraphs:
        raise RuntimeError('Journal entry content is empty.')

    now = _journal_now()
    normalized_title = _collapse_whitespace(title)
    if not normalized_title:
        normalized_title = 'Journal capture'

    filename = now.strftime('%Y-%m-%d %H.%M.%S') + '.txt'
    target_path = notes_dir / filename
    body = [
        normalized_title,
        f'Created: {now.strftime("%d/%m/%Y %H:%M %Z")}',
        f'Source: {source}',
        '',
        *paragraphs,
        '',
    ]
    target_path.write_text('\n'.join(body), encoding='utf-8')

    return {
        'timestamp': now.isoformat(timespec='seconds'),
        'date': now.strftime('%d/%m/%Y'),
        'time': now.strftime('%H:%M %Z'),
        'title': normalized_title,
        'text': '\n\n'.join(paragraphs),
        'source': source,
        'path': str(target_path),
        'directory': str(notes_dir),
    }


def _mime_suffix(mime_type):
    mapping = {
        'audio/webm': '.webm',
        'audio/webm;codecs=opus': '.webm',
        'audio/mp4': '.m4a',
        'audio/mpeg': '.mp3',
        'audio/mp3': '.mp3',
        'audio/wav': '.wav',
        'audio/x-wav': '.wav',
        'audio/ogg': '.ogg',
        'audio/ogg;codecs=opus': '.ogg',
    }
    return mapping.get((mime_type or '').lower(), '.webm')


def _transcribe_audio_blob(audio_base64, mime_type, language='en'):
    capabilities = _journal_capture_capabilities('', '')
    if not capabilities['canServerTranscribe']:
        reason = capabilities['transcribeReasons'][0] if capabilities['transcribeReasons'] else 'Server transcription is unavailable.'
        raise RuntimeError(reason)

    try:
        audio_bytes = base64.b64decode(audio_base64, validate=True)
    except Exception as exc:
        raise RuntimeError('Audio payload could not be decoded.') from exc

    if not audio_bytes:
        raise RuntimeError('Audio payload is empty.')
    if len(audio_bytes) > MAX_AUDIO_BYTES:
        raise RuntimeError('Audio payload exceeds the 25MB transcription limit.')

    with tempfile.TemporaryDirectory() as tmp_dir:
        audio_path = Path(tmp_dir) / f'journal-note{_mime_suffix(mime_type)}'
        audio_path.write_bytes(audio_bytes)

        command = [
            'python3',
            str(TRANSCRIBE_CLI),
            str(audio_path),
            '--stdout',
            '--response-format',
            'text',
        ]
        if language:
            command.extend(['--language', language])

        result = subprocess.run(command, capture_output=True, text=True)
        if result.returncode != 0:
            message = (result.stderr or result.stdout).strip() or 'Transcription failed.'
            raise RuntimeError(message)

        transcript = (result.stdout or '').strip()
        if not transcript:
            raise RuntimeError('Transcription returned empty text.')
        return transcript


def _read_docx_paragraphs(path):
    if not os.path.exists(path):
        return []

    with zipfile.ZipFile(path) as archive:
        xml = archive.read('word/document.xml')

    root = ET.fromstring(xml)
    paragraphs = []
    for paragraph in root.findall('.//w:body/w:p', WORD_NS):
        parts = [node.text or '' for node in paragraph.findall('.//w:t', WORD_NS)]
        text = ''.join(parts).strip()
        if text:
            paragraphs.append(text)
    return paragraphs


def _sectionise_soul_doc(paragraphs):
    sections = {'__intro__': []}
    current = '__intro__'

    for paragraph in paragraphs:
        heading = paragraph.rstrip(':').strip()
        if heading in SOUL_SECTION_HEADINGS:
            current = heading
            sections.setdefault(current, [])
            continue
        sections.setdefault(current, []).append(paragraph)

    return sections


def _parse_soul_doc(path):
    paragraphs = _read_docx_paragraphs(path)
    if not paragraphs:
        return {}

    sections = _sectionise_soul_doc(paragraphs)
    values = []
    for name in ['Love', 'Home', 'Embodiment', 'Beauty', 'Truthfulness', 'Compounding']:
        body = ' '.join(sections.get(name, []))
        if body:
            values.append({'name': name, 'text': body})

    north_star = ''
    for paragraph in paragraphs:
        if paragraph.startswith('I want a life that feels'):
            north_star = paragraph
            break

    becoming = [line for line in sections.get('I am becoming a man who', []) if line]
    if not becoming:
        becoming = [line for line in sections.get('What I am becoming', []) if line]

    return {
        'title': paragraphs[0],
        'northStar': north_star,
        'opening': sections.get('__intro__', [])[:4],
        'values': values,
        'becoming': becoming,
        'nonNegotiables': sections.get('Non-negotiables', []),
        'question': ' '.join(sections.get('The question underneath all of this', [])),
    }


def _journal_sentence_records(paragraphs):
    records = []
    current_date = ''
    current_title = ''

    for index, paragraph in enumerate(paragraphs):
        paragraph_text = _collapse_whitespace(paragraph)
        if not paragraph_text:
            continue

        match = JOURNAL_DATE_RE.match(paragraph_text)
        if match:
            current_date = match.group('date')
            current_title = ''
            body = _collapse_whitespace(match.group('body'))
            if ': ' in body:
                possible_title, _ = body.split(': ', 1)
                if 0 < len(possible_title.split()) <= 8:
                    current_title = possible_title.strip()
            if not current_title:
                current_title = _truncate_text(body, limit=72)

        normalized = paragraph_text.replace('//', '. ')
        sentences = [segment.strip() for segment in re.split(r'(?<=[.!?])\s+', normalized) if segment.strip()]
        if not sentences:
            sentences = [paragraph_text]

        for sentence in sentences:
            records.append({
                'paragraph_index': index,
                'date': current_date,
                'title': current_title,
                'paragraph': paragraph_text,
                'sentence': sentence,
            })

    return records


def _journal_relation_tags(text):
    haystack = f" {_collapse_whitespace(text).lower()} "
    tags = set()
    for tag, hints in JOURNAL_RELATIONSHIP_HINTS.items():
        if any(hint in haystack for hint in hints):
            tags.add(tag)
    return tags


def _journal_public_context_score(text):
    haystack = _collapse_whitespace(text).lower()
    return sum(1 for hint in JOURNAL_PUBLIC_CONTEXT_HINTS if hint in haystack)


def _journal_date_sort_value(value):
    try:
        return datetime.strptime(value, '%d/%m/%Y')
    except (TypeError, ValueError):
        return datetime.min


def _extract_name_list(segment):
    cleaned = _collapse_whitespace(segment)
    if not cleaned:
        return []

    normalized = cleaned.replace(' & ', ', ').replace(' and ', ', ').replace(' n ', ', ')
    parts = [part.strip(" .,:;!?()[]{}\"'") for part in normalized.split(',') if part.strip()]
    candidates = []
    for part in parts:
        item = _collapse_whitespace(re.sub(r"(?:'s|’s)\b", '', part))
        item = re.sub(r'^(?:my|the|a|an)\s+', '', item, flags=re.IGNORECASE)
        if not item:
            continue
        if not re.fullmatch(r'[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?', item):
            continue
        if any(token in JOURNAL_PERSON_STOPWORDS for token in item.split()):
            continue
        if item in JOURNAL_PERSON_STOP_ENTITIES:
            continue
        candidates.append({
            'name': item,
            'fragment': part,
        })
    return candidates


def _is_possessive_reference(sentence, start_index):
    prefix = sentence[max(0, start_index - 24):start_index]
    return bool(re.search(r"[A-Za-z]+[’']s\s+$", prefix))


def _journal_name_context(sentence, raw_name, fallback=''):
    match = re.search(rf'\b{re.escape(raw_name)}\b', sentence, flags=re.IGNORECASE)
    if not match:
        return fallback or sentence
    start = max(0, match.start() - 48)
    end = min(len(sentence), match.end() + 96)
    return sentence[start:end]


def _journal_candidate_names(record):
    sentence = record.get('sentence', '')
    candidates = []

    for pattern, canonical in JOURNAL_FAMILY_ENTITY_PATTERNS:
        for match in pattern.finditer(sentence):
            if _is_possessive_reference(sentence, match.start()):
                continue
            candidates.append({
                'raw_name': canonical,
                'cue_tags': {'family'},
                'local_context': _journal_name_context(sentence, canonical, sentence[max(0, match.start() - 12):match.end() + 48]),
            })

    for pattern, cue_tags in JOURNAL_RELATION_ROLE_PATTERNS:
        for match in pattern.finditer(sentence):
            raw_name = match.group('name')
            candidates.append({
                'raw_name': raw_name,
                'cue_tags': set(cue_tags),
                'local_context': _journal_name_context(sentence, raw_name, match.group(0)),
            })

    for match in re.finditer(r'\b(?P<name>[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\s*\((?P<detail>[^)]+)\)', sentence):
        detail = _collapse_whitespace(match.group('detail')).lower()
        detail_tags = _journal_relation_tags(detail)
        if 'dating' in detail or 'romantic' in detail:
            detail_tags.add('romantic')
        if any(keyword in detail for keyword in ('cousin', 'aunt', 'dad', 'father', 'mother')):
            detail_tags.add('family')
        candidates.append({
            'raw_name': match.group('name'),
            'cue_tags': detail_tags,
            'local_context': f"{match.group('name')} ({match.group('detail')})",
        })

    list_cues = (
        ('by the name of ', {'mentor'}),
        ('the likes of ', {'friend'}),
        ('great people, ', {'friend', 'community'}),
        ('family catchup with ', {'family'}),
        ('met ', {'community'}),
        ('meeting with ', {'friend'}),
        ('conversation with ', {'friend'}),
        ('conversations with ', {'friend'}),
        ('chat with ', {'friend'}),
        ('long chat with ', {'friend'}),
        ('good chat with ', {'friend'}),
        ('great to speak to ', {'friend'}),
        ('deep conversation with ', {'friend'}),
        ('spoke to ', {'friend'}),
        ('speak to ', {'friend'}),
        ('texted ', {'romantic'}),
        ('bumped into ', {'community'}),
        ('talking to ', {'friend'}),
        ('spent the day with ', {'friend'}),
        ('spent time with ', {'friend'}),
        ('time with ', {'friend'}),
        ('evening with ', {'friend'}),
        ('night with ', {'friend'}),
        ('back with ', {'friend'}),
        ('went out with ', {'friend'}),
        ('coffee with ', {'friend'}),
        ('gig with ', {'friend'}),
        ('being with ', {'friend'}),
        ('getting on with ', {'friend'}),
        ('argument with ', {'romantic'}),
        ('walked with ', {'friend'}),
        ('walking with ', {'friend'}),
        ('cycling with ', {'friend'}),
        ('run with ', {'community'}),
        ('running with ', {'community'}),
        ('trail run with ', {'community'}),
        ('climbing with ', {'friend', 'community'}),
    )
    lower_sentence = sentence.lower()
    for cue, cue_tags in list_cues:
        if cue not in lower_sentence:
            continue
        start = lower_sentence.index(cue) + len(cue)
        fragment = sentence[start:]
        fragment = re.split(r'[.;!?]', fragment, maxsplit=1)[0]
        for item in _extract_name_list(fragment):
            candidates.append({
                'raw_name': item['name'],
                'cue_tags': set(cue_tags),
                'local_context': _journal_name_context(sentence, item['name'], item['fragment']),
            })

    if not candidates and _journal_relation_tags(sentence):
        for pattern in JOURNAL_DIRECT_NAME_PATTERNS:
            for match in pattern.finditer(sentence):
                raw_name = match.group('name')
                if raw_name in JOURNAL_PERSON_STOP_ENTITIES or raw_name in JOURNAL_PERSON_STOPWORDS:
                    continue
                candidates.append({
                    'raw_name': raw_name,
                    'cue_tags': set(),
                    'local_context': _journal_name_context(sentence, raw_name, match.group(0)),
                })

    if candidates:
        return candidates
    return []


def _normalize_journal_person_name(value):
    key = _collapse_whitespace(re.sub(r"(?:'s|’s)\b", '', value))
    return JOURNAL_SAFE_ALIAS_MAP.get(key, key)


def _journal_primary_relationship(tag_counts):
    family = tag_counts.get('family', 0)
    romantic = tag_counts.get('romantic', 0)
    work = tag_counts.get('work', 0)
    mentor = tag_counts.get('mentor', 0)
    friend = tag_counts.get('friend', 0)
    community = tag_counts.get('community', 0)

    core_scores = [(tag, tag_counts.get(tag, 0)) for tag in ('family', 'romantic', 'work') if tag_counts.get(tag, 0) > 0]
    core_scores.sort(key=lambda item: item[1], reverse=True)
    if core_scores:
        top_tag, top_score = core_scores[0]
        next_score = core_scores[1][1] if len(core_scores) > 1 else 0
        if next_score and next_score >= max(2, top_score * 0.75):
            return 'mixed'
        if top_tag == 'family':
            return 'family'
        if top_tag == 'romantic':
            return 'romantic'
        if top_tag == 'work':
            return 'work-friend' if friend else 'work'

    if mentor and friend:
        return 'mentor-friend'
    if mentor:
        return 'mentor'
    if friend:
        return 'friend'
    if community:
        return 'community'
    return 'acquaintance'


def _person_fact_score(snippet):
    text = _collapse_whitespace(snippet).lower()
    score = 0
    for keyword in (
        'greatest friend', 'dear friend', 'soulbond', 'falling for', 'dating', 'cousin', 'aunt', 'dad',
        'brother', 'grandparent', 'manager', 'colleague', 'opened up', 'deep conversation', 'prompted',
        'mentor', 'honest', 'love', 'good friend',
    ):
        if keyword in text:
            score += 2
    score += min(len(text) // 120, 2)
    return score


def _person_relation_label(value):
    return _collapse_whitespace(str(value or '').replace('-', ' ')) or 'acquaintance'


def _person_relation_to_dan(primary_relationship):
    if primary_relationship == 'family':
        return 'family member'
    if primary_relationship == 'romantic':
        return 'romantic partner'
    if primary_relationship == 'mentor-friend':
        return 'mentor and friend'
    if primary_relationship == 'mentor':
        return 'mentor'
    if primary_relationship == 'work-friend':
        return 'work friend'
    if primary_relationship == 'work':
        return 'colleague'
    if primary_relationship == 'friend':
        return 'friend'
    if primary_relationship == 'community':
        return 'community connection'
    if primary_relationship == 'mixed':
        return 'significant relationship'
    return 'acquaintance'


def _canonical_person_override(person):
    key = _collapse_whitespace(person.get('name', '')).lower()
    override = JOURNAL_CANONICAL_PERSON_ALIASES.get(key)
    if override:
        return dict(override)

    relationship_tags = set(person.get('relationshipTags', []) or [])
    if key == 'will' and 'family' in relationship_tags:
        return {
            'name': 'William Atkinson',
            'relation_to_dan': 'brother',
            'aliases': ['Will'],
        }
    return None


def _source_person_match(name, records):
    target = _collapse_whitespace(name).lower()
    if not target:
        return None
    for record in records or []:
        if _collapse_whitespace(record.get('name', '')).lower() == target:
            return record
    return None


def _person_article_lead(name, relation_to_dan, primary_relationship, mentions, mention_dates, self_name):
    relation_label = PERSON_RELATION_TO_DAN_LABELS.get(relation_to_dan or '', relation_to_dan or '')
    if relation_label:
        first = f"{name} is {self_name}'s {relation_label}."
    else:
        first = f"{name} is a person recorded in {self_name}'s knowledge base."

    context_label = _person_relation_label(primary_relationship)
    if mention_dates:
        date_span = mention_dates[0] if len(mention_dates) == 1 else f"{mention_dates[0]} to {mention_dates[-1]}"
        second = (
            f"The current record is compiled mainly from the Generous Slice journal, where {name} appears "
            f"as a {context_label} across {mentions} mention{'s' if mentions != 1 else ''} from {date_span}."
        )
    else:
        second = (
            f"The current record is compiled mainly from the Generous Slice journal, where {name} appears "
            f"as a {context_label} across {mentions} mention{'s' if mentions != 1 else ''}."
        )
    return f"{first} {second}".strip()


def _person_relation_summary(name, relation_to_dan, primary_relationship, mention_dates, aliases, self_name, personal_match, research_match):
    relation_label = PERSON_RELATION_TO_DAN_LABELS.get(relation_to_dan or '', relation_to_dan or 'person')
    lines = [
        f"{name} is currently treated as {self_name}'s {relation_label} in Akibwa."
        if relation_to_dan
        else f"{name} is currently treated as a {_person_relation_label(primary_relationship)} in {self_name}'s life."
    ]
    lines.append(f"Primary relationship signal: {_person_relation_label(primary_relationship)}.")
    if mention_dates:
        if len(mention_dates) == 1:
            lines.append(f"Current journal span: {mention_dates[0]}.")
        else:
            lines.append(f"Current journal span: {mention_dates[0]} to {mention_dates[-1]}.")
    if aliases:
        lines.append(f"Aliases seen in source text: {', '.join(aliases[:6])}.")
    if personal_match and personal_match.get('whyInfluential'):
        lines.append(f"Personal context: {personal_match.get('whyInfluential')}")
    if research_match and research_match.get('summary'):
        lines.append(f"Research note: {research_match.get('summary')}")
    return lines


def _person_context_notes(person, primary_relationship, aliases, personal_match, research_match):
    notes = [
        f"Primary relationship type: {_person_relation_label(primary_relationship)}.",
    ]
    tags = person.get('relationshipTags', []) or []
    if tags:
        notes.append(f"Context tags currently attached: {', '.join(_person_relation_label(tag) for tag in tags)}.")
    if person.get('mentions'):
        notes.append(f"Mentions currently tracked: {person.get('mentions')} in the journal-derived layer.")
    if aliases:
        notes.append(f"Known aliases: {', '.join(aliases[:6])}.")
    if personal_match and personal_match.get('relationship'):
        notes.append(f"Profile relationship signal: {_person_relation_label(personal_match.get('relationship'))}.")
    if research_match and research_match.get('role'):
        notes.append(f"Public role note: {research_match.get('role')}.")
    return notes


def _looks_like_publishable_person_name(name):
    cleaned = _collapse_whitespace(name)
    if not cleaned:
        return False
    lowered = cleaned.lower()
    if lowered in PERSON_PAGE_BLOCKLIST:
        return False
    if lowered in {_collapse_whitespace(item).lower() for item in JOURNAL_PERSON_STOP_ENTITIES}:
        return False
    tokens = cleaned.split()
    if any(token.lower() in {item.lower() for item in JOURNAL_PERSON_STOPWORDS} for token in tokens):
        return False
    return all(re.fullmatch(r"[A-Z][A-Za-z'’-]*", token) for token in tokens)


def _should_publish_person_page(name, relation_to_dan, primary_relationship, mentions, personal_match, research_match):
    if not _looks_like_publishable_person_name(name):
        return False
    if relation_to_dan in {'father', 'mother', 'brother', 'sister', 'uncle', 'aunt', 'grandfather', 'grandmother', 'cousin'}:
        return True
    if primary_relationship in {'romantic', 'mentor', 'mentor-friend', 'work', 'work-friend'}:
        return int(mentions or 0) >= 2
    if personal_match or research_match:
        return True
    if primary_relationship in {'friend', 'community', 'mixed'}:
        return int(mentions or 0) >= 3
    return int(mentions or 0) >= 4


def _build_people_layer(raw_people, profile_library, self_name):
    raw_directory = (raw_people or {}).get('directory', []) or []
    merged = {}

    for person in raw_directory:
        canonical = _canonical_person_override(person)
        name = _collapse_whitespace((canonical or {}).get('name') or person.get('name', ''))
        if not name:
            continue

        bucket = merged.setdefault(name, {
            'name': name,
            'slug': _slugify(name),
            'aliases': Counter(),
            'tagCounts': Counter(),
            'mentions': 0,
            'mentionDates': set(),
            'evidence': [],
            'snippetKeys': set(),
            'relationToDan': (canonical or {}).get('relation_to_dan', ''),
        })

        bucket['mentions'] += int(person.get('mentions', 0) or 0)

        if canonical:
            for alias in canonical.get('aliases', []):
                bucket['aliases'][alias] += 1

        for alias in [person.get('name', ''), *(person.get('aliases') or [])]:
            cleaned_alias = _collapse_whitespace(alias)
            if cleaned_alias:
                bucket['aliases'][cleaned_alias] += 1

        for tag in person.get('relationshipTags', []) or []:
            bucket['tagCounts'][tag] += 1

        primary = _collapse_whitespace(person.get('primaryRelationship', ''))
        if primary:
            bucket['tagCounts'][primary] += 2

        for date_value in (person.get('firstMentionDate', ''), person.get('lastMentionDate', '')):
            if date_value:
                bucket['mentionDates'].add(date_value)

        for item in person.get('evidence', []) or []:
            snippet = _collapse_whitespace(item.get('snippet', ''))
            if not snippet:
                continue
            snippet_key = snippet.lower()
            if snippet_key in bucket['snippetKeys']:
                continue
            bucket['snippetKeys'].add(snippet_key)
            bucket['evidence'].append({
                'date': item.get('date', ''),
                'title': item.get('title', ''),
                'snippet': snippet,
                'tags': sorted(item.get('tags', []) or []),
            })

    directory = []
    relationship_counts = Counter()
    personal_influences = profile_library.get('personalInfluences', []) or []
    research_notes = profile_library.get('researchNotes', []) or []

    for bucket in merged.values():
        mention_dates = sorted(bucket['mentionDates'], key=_journal_date_sort_value)
        evidence = sorted(
            bucket['evidence'],
            key=lambda item: (_person_fact_score(item.get('snippet', '')), _journal_date_sort_value(item.get('date', ''))),
            reverse=True,
        )

        aliases = [alias for alias, _ in bucket['aliases'].most_common() if alias and alias != bucket['name']]
        personal_match = _source_person_match(bucket['name'], personal_influences)
        research_match = _source_person_match(bucket['name'], research_notes)
        primary = _journal_primary_relationship(bucket['tagCounts'])
        relation_to_dan = bucket.get('relationToDan') or _person_relation_to_dan(primary)
        if relation_to_dan in {'father', 'mother', 'brother', 'sister', 'uncle', 'aunt', 'grandfather', 'grandmother', 'cousin'}:
            primary = 'family'
        if not _should_publish_person_page(bucket['name'], relation_to_dan, primary, bucket['mentions'], personal_match, research_match):
            continue
        relationship_counts[primary] += 1
        article_lead = _person_article_lead(
            bucket['name'],
            relation_to_dan,
            primary,
            bucket['mentions'],
            mention_dates,
            self_name,
        )
        relation_summary = _person_relation_summary(
            bucket['name'],
            relation_to_dan,
            primary,
            mention_dates,
            aliases,
            self_name,
            personal_match,
            research_match,
        )
        facts = []
        if aliases:
            facts.append(f"Canonical page built from source aliases such as {', '.join(aliases[:6])}.")
        if relation_to_dan:
            facts.append(f"Relation to Daniel: {PERSON_RELATION_TO_DAN_LABELS.get(relation_to_dan, relation_to_dan)}.")
        facts.append(f"Primary relationship type: {_person_relation_label(primary)}.")
        tags = sorted(tag for tag, count in bucket['tagCounts'].items() if count > 0)
        if tags:
            facts.append(f"Current context tags: {', '.join(_person_relation_label(tag) for tag in tags)}.")
        if mention_dates:
            if len(mention_dates) == 1:
                facts.append(f"Current journal span: {mention_dates[0]}.")
            else:
                facts.append(f"Current journal span: {mention_dates[0]} to {mention_dates[-1]}.")
        facts.append(f"Mentions currently tracked: {bucket['mentions']}.")
        if personal_match and personal_match.get('whyInfluential'):
            facts.append(f"Personal context note: {personal_match.get('whyInfluential')}")
        if research_match and research_match.get('summary'):
            facts.append(f"Research note: {research_match.get('summary')}")
        facts = facts[:6]
        context_notes = _person_context_notes(
            {
                'relationshipTags': tags,
                'mentions': bucket['mentions'],
            },
            primary,
            aliases,
            personal_match,
            research_match,
        )

        directory.append({
            'name': bucket['name'],
            'slug': bucket['slug'],
            'primaryRelationship': primary,
            'relationshipTags': sorted(tag for tag, count in bucket['tagCounts'].items() if count > 0),
            'relationToDan': relation_to_dan,
            'mentions': bucket['mentions'],
            'aliases': aliases,
            'articleLead': article_lead,
            'summary': article_lead,
            'relationSummary': relation_summary,
            'contextNotes': context_notes,
            'facts': facts,
            'evidence': evidence[:10],
            'firstMentionDate': mention_dates[0] if mention_dates else '',
            'lastMentionDate': mention_dates[-1] if mention_dates else '',
            'profileContext': personal_match or {},
            'researchContext': research_match or {},
            'infobox': {
                'relation_to_dan': relation_to_dan,
                'primary_relationship': primary,
                'mentions': bucket['mentions'],
                'aliases': aliases,
                'first_mention_date': mention_dates[0] if mention_dates else '',
                'last_mention_date': mention_dates[-1] if mention_dates else '',
            },
        })

    directory.sort(key=lambda item: (-int(item.get('mentions', 0)), -len(item.get('facts', [])), item.get('name', '')))
    featured = directory[:12]
    relationship_breakdown = [
        {'relationship': key, 'count': value}
        for key, value in sorted(relationship_counts.items(), key=lambda item: (-item[1], item[0]))
    ]

    return {
        'summary': {
            'person_count': len(directory),
            'featured_count': len(featured),
            'relationship_breakdown': relationship_breakdown,
        },
        'featured': featured,
        'directory': directory,
    }


def _extract_people_from_journal(paragraphs):
    sentence_records = _journal_sentence_records(paragraphs)
    people = {}

    for record in sentence_records:
        sentence = record.get('sentence', '')
        public_score = _journal_public_context_score(sentence)
        for candidate in _journal_candidate_names(record):
            raw_name = candidate.get('raw_name', '')
            cue_tags = set(candidate.get('cue_tags') or [])
            local_context = candidate.get('local_context', sentence)
            name = _normalize_journal_person_name(raw_name)
            if not name or name in JOURNAL_PERSON_STOP_ENTITIES:
                continue
            if any(token in JOURNAL_PERSON_STOPWORDS for token in name.split()):
                continue

            local_tags = _journal_relation_tags(local_context)
            all_tags = set(cue_tags) | set(local_tags)
            if public_score and public_score > (len(all_tags) + (1 if cue_tags else 0)):
                continue

            person = people.setdefault(name, {
                'name': name,
                'aliases': Counter(),
                'tagCounts': Counter(),
                'mentions': 0,
                'mentionDates': set(),
                'evidence': [],
                'snippetKeys': set(),
            })
            person['aliases'][raw_name] += 1
            person['mentions'] += 1
            if record.get('date'):
                person['mentionDates'].add(record['date'])
            for tag in all_tags:
                weight = 2 if tag in cue_tags or tag in {'family', 'romantic', 'work', 'mentor'} and tag in local_tags else 1
                person['tagCounts'][tag] += weight

            snippet = _truncate_text(sentence, limit=320)
            snippet_key = snippet.lower()
            if snippet_key not in person['snippetKeys'] and len(person['evidence']) < 12:
                person['snippetKeys'].add(snippet_key)
                person['evidence'].append({
                    'date': record.get('date', ''),
                    'title': record.get('title', ''),
                    'snippet': snippet,
                    'tags': sorted(all_tags),
                })

    directory = []
    relationship_counts = Counter()
    for person in people.values():
        if person['mentions'] <= 0:
            continue
        primary = _journal_primary_relationship(person['tagCounts'])
        if primary == 'acquaintance' and person['mentions'] < 2:
            continue
        relationship_counts[primary] += 1

        evidence = sorted(
            person['evidence'],
            key=lambda item: (_person_fact_score(item.get('snippet', '')), _journal_date_sort_value(item.get('date', ''))),
            reverse=True,
        )
        facts = []
        seen = set()
        for item in evidence:
            snippet = item.get('snippet', '')
            key = snippet.lower()
            if key in seen:
                continue
            seen.add(key)
            facts.append(snippet)
            if len(facts) >= 5:
                break

        aliases = [alias for alias, _ in person['aliases'].most_common() if alias != person['name']]
        core_contexts = [tag for tag in ('family', 'romantic', 'work') if person['tagCounts'].get(tag, 0) > 0]
        if len(core_contexts) > 1:
            summary = (
                f"{person['name']} appears across {', '.join(core_contexts)} contexts in the journal and may need disambiguation with future sources."
            )
        elif primary != 'mixed':
            summary = f"{person['name']} appears in the journal as a {primary.replace('-', ' ')} relationship."
        else:
            summary = f"{person['name']} appears across multiple relationship contexts in the journal and may need disambiguation with future sources."
        if facts:
            summary += f" {facts[0]}"

        mention_dates = sorted(person['mentionDates'], key=_journal_date_sort_value)
        directory.append({
            'name': person['name'],
            'slug': _slugify(person['name']),
            'primaryRelationship': primary,
            'relationshipTags': sorted(tag for tag, count in person['tagCounts'].items() if count > 0),
            'mentions': person['mentions'],
            'aliases': aliases,
            'summary': summary,
            'facts': facts,
            'evidence': person['evidence'][:8],
            'firstMentionDate': mention_dates[0] if mention_dates else '',
            'lastMentionDate': mention_dates[-1] if mention_dates else '',
        })

    directory.sort(key=lambda item: (item['mentions'], len(item['facts']), item['name']), reverse=True)
    featured = directory[:12]
    relationship_breakdown = [
        {'relationship': key, 'count': value}
        for key, value in sorted(relationship_counts.items(), key=lambda item: (-item[1], item[0]))
    ]

    return {
        'summary': {
            'person_count': len(directory),
            'featured_count': len(featured),
            'relationship_breakdown': relationship_breakdown,
        },
        'featured': featured,
        'directory': directory,
    }


def _parse_journal(path):
    paragraphs = _read_docx_paragraphs(path)
    entries = []
    index = 0

    while index < len(paragraphs):
        paragraph = paragraphs[index]
        match = JOURNAL_DATE_RE.match(paragraph)
        if not match:
            index += 1
            continue

        date_text = match.group('date')
        time_text = match.group('time') or '00:00'
        body = match.group('body').strip()
        title = ''
        text = body

        if ': ' in body:
            possible_title, remainder = body.split(': ', 1)
            if 0 < len(possible_title.split()) <= 6:
                title = possible_title.strip()
                text = remainder.strip()

        continuation = []
        next_index = index + 1
        while next_index < len(paragraphs) and not JOURNAL_DATE_RE.match(paragraphs[next_index]):
            continuation.append(paragraphs[next_index])
            next_index += 1

        if continuation:
            if not title and not text and 0 < len(body.split()) <= 8:
                title = body
                text = '\n\n'.join(continuation)
            elif not title and 0 < len(body.split()) <= 8 and body == text:
                title = body
                text = '\n\n'.join(continuation)
            else:
                pieces = [piece for piece in [text] + continuation if piece]
                text = '\n\n'.join(pieces)

        try:
            sort_key = datetime.strptime(f'{date_text} {time_text}', '%d/%m/%Y %H:%M')
        except ValueError:
            try:
                sort_key = datetime.strptime(date_text, '%d/%m/%Y')
            except ValueError:
                index = next_index
                continue

        entries.append({
            'date': date_text,
            'time': None if time_text == '00:00' and not match.group('time') else time_text,
            'title': title or '',
            'text': text,
            'sortKey': sort_key.isoformat(),
            'index': index,
        })
        index = next_index

    entries.sort(key=lambda item: item['index'], reverse=True)
    latest_entries = entries[:8]

    theme_rules = {
        'Rootedness': ('home', 'root', 'belong', 'local', 'compound', 'friendship'),
        'Embodiment': ('run', 'body', 'athlete', 'health', 'sleep', 'fitness', 'recovery'),
        'Dance': ('salsa', 'dance'),
        'Connection': ('date', 'girl', 'woman', 'friend', 'friends', 'love', 'hostel'),
        'Travel': ('travel', 'travelling', 'city', 'san francisco', 'hawaii', 'hostel'),
        'Work & AI': ('office', 'work', 'ai', 'code', 'national wealth fund'),
    }

    theme_counts = Counter()
    for entry in latest_entries:
        haystack = f"{entry['title']} {entry['text']}".lower()
        for theme, keywords in theme_rules.items():
            if any(keyword in haystack for keyword in keywords):
                theme_counts[theme] += 1

    return {
        'latestEntries': latest_entries,
        'themes': [theme for theme, _ in theme_counts.most_common(5)],
        'people': _extract_people_from_journal(paragraphs),
    }


def _parse_creative_doc(path):
    paragraphs = _read_docx_paragraphs(path)
    ignored_titles = {'formal writings', 'creative writing'}

    def looks_like_title(paragraph):
        cleaned = paragraph.rstrip(':').strip(' "\'')
        if not cleaned or cleaned.lower() in ignored_titles:
            return False, ''
        words = cleaned.split()
        is_title = len(words) <= 5 and (paragraph.endswith(':') or cleaned.istitle())
        return is_title, cleaned

    titles = []
    entries = []
    current_entry = None

    for paragraph in paragraphs:
        is_title, cleaned = looks_like_title(paragraph)
        if is_title:
            if cleaned not in titles:
                titles.append(cleaned)
            current_entry = {
                'title': cleaned,
                'paragraphs': [],
            }
            entries.append(current_entry)
            continue

        if current_entry is not None:
            current_entry['paragraphs'].append(paragraph)

    creative_entries = []
    for entry in entries:
        body_paragraphs = [paragraph for paragraph in entry['paragraphs'] if paragraph]
        body = '\n\n'.join(body_paragraphs).strip()
        excerpt = body_paragraphs[0] if body_paragraphs else ''
        creative_entries.append({
            'title': entry['title'],
            'excerpt': excerpt,
            'body': body,
            'paragraphs': body_paragraphs,
        })

    return {
        'titles': titles[:12],
        'entries': creative_entries[:24],
        'sampleLines': paragraphs[1:9],
    }


def _sheet_rows(sheet):
    return list(sheet.iter_rows(values_only=True))


def _row_dict(sheet):
    data = {}
    for row in _sheet_rows(sheet):
        if not row:
            continue
        key = row[0]
        if key:
            data[str(key)] = row[1:]
    return data


def _employment_logo_meta(employer):
    value = _collapse_whitespace(employer)
    mapping = {
        'National Wealth Fund': {'label': 'NWF', 'tone': 'forest'},
        'UK Infrastructure Bank Limited (trading as National Wealth Fund)': {'label': 'NWF', 'tone': 'forest'},
        'Leeds Building Society': {'label': 'LBS', 'tone': 'clay'},
        'Sky Betting & Gaming': {'label': 'SBG', 'tone': 'gold'},
        'Vanquis Bank': {'label': 'VB', 'tone': 'ink'},
        'Lloyds Banking Group': {'label': 'LBG', 'tone': 'forest'},
    }
    if value in mapping:
        return mapping[value]

    words = [word for word in re.split(r'[^A-Za-z0-9]+', value) if word]
    label = ''.join(word[0] for word in words[:3]).upper()[:3] or '?'
    return {'label': label, 'tone': 'ink'}


def _parse_employment_range(value):
    raw = _collapse_whitespace(value)
    match = EMPLOYMENT_RANGE_RE.match(raw)
    if not match:
        return {
            'label': raw,
            'startIso': '',
            'endIso': '',
            'isCurrent': 'present' in raw.lower(),
        }

    start = datetime.strptime(
        f"{match.group('start_month')} {match.group('start_year')}",
        '%b %Y',
    )
    start_iso = start.date().isoformat()

    if match.group('present'):
        return {
            'label': raw,
            'startIso': start_iso,
            'endIso': '',
            'isCurrent': True,
        }

    end = datetime.strptime(
        f"{match.group('end_month')} {match.group('end_year')}",
        '%b %Y',
    )
    return {
        'label': raw,
        'startIso': start_iso,
        'endIso': end.date().isoformat(),
        'isCurrent': False,
    }


def _parse_employment_history(sheet):
    history = []
    in_section = False

    for row in sheet.iter_rows(values_only=True):
        key = row[0] if row else None
        value = row[1] if len(row) > 1 else None

        if key == 'EMPLOYMENT HISTORY':
            in_section = True
            continue

        if not in_section:
            continue

        if key in ('LinkedIn Headline', 'LinkedIn Connections', 'Emergency Contact', 'EDUCATION'):
            break

        if not key or not value:
            continue

        range_meta = _parse_employment_range(str(key))
        if not range_meta['label']:
            continue

        raw_value = _collapse_whitespace(value)
        if ' — ' in raw_value:
            role, remainder = raw_value.split(' — ', 1)
        else:
            role, remainder = raw_value, ''

        employer = remainder
        location = ''
        if remainder and ', ' in remainder:
            employer, location = remainder.split(', ', 1)

        employer = employer or remainder or raw_value
        logo = _employment_logo_meta(employer)
        history.append({
            'range': range_meta['label'],
            'startIso': range_meta['startIso'],
            'endIso': range_meta['endIso'],
            'isCurrent': range_meta['isCurrent'],
            'role': role,
            'employer': employer,
            'location': location,
            'display': raw_value,
            'logoLabel': logo['label'],
            'logoTone': logo['tone'],
        })

    history.sort(key=lambda item: item.get('startIso') or '', reverse=True)
    return history


def _parse_inventory_workbook(path):
    workbook = _try_load_workbook(path)
    personal = _row_dict(workbook['Personal Info'])
    income = _row_dict(workbook['Income'])
    employment_history = _parse_employment_history(workbook['Personal Info'])
    recent_payslips = []

    for row in workbook['Income'].iter_rows(min_row=1, values_only=True):
        detail, amount, notes = row[:3]
        if isinstance(detail, str) and re.fullmatch(r'[A-Z][a-z]{2} \d{4} \(Period \d+\)', detail):
            recent_payslips.append({
                'label': detail,
                'netPay': round(_safe_float(amount), 2),
                'notes': str(notes or ''),
            })

    inventory_items = []
    for row in workbook['Inventory'].iter_rows(min_row=2, values_only=True):
        category, item, price, weight, *rest = row
        if not item or str(item).upper() == 'TOTAL':
            continue
        inventory_items.append({
            'category': str(category or ''),
            'item': str(item),
            'priceGbp': round(_safe_float(price), 2),
            'weightG': round(_safe_float(weight), 1),
            'status': str(rest[0]) if rest and rest[0] else '',
        })

    on_order = []
    for row in workbook['On Order'].iter_rows(min_row=2, values_only=True):
        category, item, price, weight = row[:4]
        if not item or str(item).upper() == 'TOTAL':
            continue
        on_order.append({
            'category': str(category or ''),
            'item': str(item),
            'priceGbp': round(_safe_float(price), 2),
            'weightG': round(_safe_float(weight), 1),
        })

    monthly_outgoings = []
    for row in workbook['Monthly Outgoings'].iter_rows(min_row=2, values_only=True):
        item, amount, notes = row[:3]
        if not item or str(item).upper().startswith('TOTAL') or not isinstance(amount, (int, float)):
            continue
        monthly_outgoings.append({
            'item': str(item),
            'amountGbp': round(float(amount), 2),
            'notes': str(notes or ''),
        })

    section = None
    investments_isa = []
    investments_taxable = []
    cash = []
    liabilities = []

    for row in workbook['Assets & Liabilities'].iter_rows(values_only=True):
        item, value, notes = row[:3]
        if item in (
            'INVESTMENTS — T212 S&S ISA',
            'INVESTMENTS — T212 Invest',
            'CASH',
            'LIABILITIES',
            'NET WORTH SUMMARY',
        ):
            section = item
            continue
        if not item or not isinstance(value, (int, float)):
            continue
        payload = {
            'item': str(item),
            'value': round(float(value), 2),
            'notes': str(notes or ''),
        }
        if section == 'INVESTMENTS — T212 S&S ISA':
            investments_isa.append(payload)
        elif section == 'INVESTMENTS — T212 Invest':
            investments_taxable.append(payload)
        elif section == 'CASH':
            cash.append(payload)
        elif section == 'LIABILITIES':
            liabilities.append(payload)

    category_weights = defaultdict(float)
    category_counts = defaultdict(int)
    for item in inventory_items:
        category_weights[item['category']] += item['weightG']
        category_counts[item['category']] += 1

    return {
        'profile': {
            'fullName': personal.get('Full Legal Name', [''])[0],
            'dateOfBirth': personal.get('Date of Birth', [''])[0],
            'age': personal.get('Age', [''])[0],
            'height': personal.get('Height', [''])[0],
            'currentLiving': personal.get('Current Living', [''])[0],
            'hometown': personal.get('Hometown', [''])[0],
        },
        'work': {
            'employer': income.get('Employer', [''])[0],
            'employerLegal': personal.get('Employer', [''])[0],
            'jobTitle': personal.get('Job Title', [''])[0],
            'linkedinTitle': personal.get('LinkedIn Title', [''])[0],
            'officeAddress': personal.get('Office Address', [''])[0],
            'startDateLinkedIn': personal.get('Start Date (LinkedIn)', [''])[0],
            'startDateContract': personal.get('Start Date (Contract)', [''])[0],
            'workingPattern': income.get('Working Pattern', [''])[0],
            'annualHoliday': income.get('Annual Holiday', [''])[0],
            'monthlyGrossSalary': round(_safe_float(income.get('Monthly Gross Salary', [0])[0]), 2),
            'annualGrossSalary': round(_safe_float(income.get('Annual Gross Salary', [0])[0]), 2),
            'annualBonusGross': round(_safe_float(income.get('Annual Bonus (2 months)', [0])[0]), 2),
            'totalAnnualCompGross': round(_safe_float(income.get('Total Annual Comp (gross)', [0])[0]), 2),
            'monthlyNetPay': round(_safe_float(income.get('Net Pay (full month)', [0])[0]), 2),
            'employerPension': round(_safe_float(income.get('Employer Pension (typical)', [0])[0]), 2),
            'recentPayslips': recent_payslips[:6],
            'history': employment_history,
        },
        'inventory': {
            'trackedItemCount': len(inventory_items),
            'trackedValueGbp': round(sum(item['priceGbp'] for item in inventory_items), 2),
            'trackedWeightG': round(sum(item['weightG'] for item in inventory_items), 1),
            'trackedWeightKg': round(sum(item['weightG'] for item in inventory_items) / 1000, 2),
            'items': inventory_items,
            'onOrder': on_order,
            'backpack': next((item for item in inventory_items if item['category'].lower() == 'backpack'), None),
            'heaviestItems': sorted(inventory_items, key=lambda item: item['weightG'], reverse=True)[:10],
            'mostExpensiveItems': sorted(inventory_items, key=lambda item: item['priceGbp'], reverse=True)[:10],
            'categoryBreakdown': sorted(
                (
                    {
                        'category': category,
                        'count': category_counts[category],
                        'weightG': round(weight, 1),
                    }
                    for category, weight in category_weights.items()
                    if category
                ),
                key=lambda item: item['weightG'],
                reverse=True,
            )[:8],
        },
        'money': {
            'monthlyOutgoings': monthly_outgoings,
            'monthlyOutgoingsTotal': round(sum(item['amountGbp'] for item in monthly_outgoings), 2),
            'isaHoldings': investments_isa,
            'taxableHoldings': investments_taxable,
            'cash': cash,
            'liabilities': liabilities,
            'isaTotal': round(sum(item['value'] for item in investments_isa), 2),
            'taxableInvestTotal': round(sum(item['value'] for item in investments_taxable), 2),
            'cashTotal': round(sum(item['value'] for item in cash), 2),
            'liabilitiesGbpTotal': round(
                sum(item['value'] for item in liabilities if 'usd' not in item['notes'].lower()),
                2,
            ),
            'medicalBillUsd': next(
                (item['value'] for item in liabilities if 'usd' in item['notes'].lower()),
                0,
            ),
            'medicalBillNotes': next(
                (item['notes'] for item in liabilities if 'usd' in item['notes'].lower()),
                '',
            ),
        },
    }


def _parse_finance_inventory_workbook(path):
    workbook = _try_load_workbook(path)
    watchlist = []
    for row in workbook['Purchase Watchlist'].iter_rows(min_row=2, values_only=True):
        category, item, status, priority, price, target, why, next_action, last_checked, notes = row[:10]
        if not item:
            continue
        watchlist.append({
            'category': str(category or ''),
            'item': str(item),
            'status': str(status or ''),
            'priority': str(priority or ''),
            'priceBudget': str(price or ''),
            'targetTrigger': str(target or ''),
            'whyItMatters': str(why or ''),
            'nextAction': str(next_action or ''),
            'lastChecked': str(last_checked or ''),
            'notes': str(notes or ''),
        })

    target_weight_kg = None
    for row in workbook['Raw Inventory'].iter_rows(min_row=1, max_row=10, values_only=True):
        for value in row:
            if isinstance(value, str):
                match = re.search(r'(\d+(?:\.\d+)?)\s*kg', value.lower())
                if match:
                    target_weight_kg = float(match.group(1))
                    break
        if target_weight_kg is not None:
            break

    return {
        'watchlist': watchlist[:8],
        'targetWeightKg': target_weight_kg,
    }


def _build_life_signals(snapshot):
    signals = []
    onebag = snapshot['onebag']
    money = snapshot['money']
    work = snapshot['work']
    journal = snapshot['journal']

    if onebag.get('targetWeightKg'):
        gap = round(onebag['trackedWeightKg'] - onebag['targetWeightKg'], 2)
        if gap > 0:
            signals.append({
                'label': 'Onebag gap',
                'text': f"Tracked setup is {gap:.2f} kg above the {onebag['targetWeightKg']:.1f} kg target.",
            })
        else:
            signals.append({
                'label': 'Onebag gap',
                'text': f"Tracked setup is within the {onebag['targetWeightKg']:.1f} kg target.",
            })

    if money.get('medicalBillUsd'):
        signals.append({
            'label': 'Medical bill',
            'text': f"Outstanding medical bill noted at ${money['medicalBillUsd']:,.2f}; insurance follow-up is still live.",
        })

    if work.get('workingPattern'):
        signals.append({
            'label': 'Work rhythm',
            'text': f"{work['workingPattern']} with typical net pay around £{work['monthlyNetPay']:,.2f} a month.",
        })

    if journal.get('themes'):
        signals.append({
            'label': 'Recent themes',
            'text': 'Recent journal entries keep circling around ' + ', '.join(journal['themes'][:3]).lower() + '.',
        })

    for watch_item in onebag.get('watchlist', []):
        if watch_item.get('priority', '').lower() == 'high':
            signals.append({
                'label': 'Watch item',
                'text': f"{watch_item['item']} is marked {watch_item['priority'].lower()} priority: {watch_item['nextAction'] or watch_item['whyItMatters']}.",
            })
            break

    return signals[:5]


def _build_life_snapshot():
    docs_dir = _docs_dir()
    if not docs_dir:
        return {
            'ok': False,
            'error': 'Akibwa Documents folder not found.',
            'searched': [candidate for candidate in _docs_candidates if candidate],
        }

    files = _life_files(docs_dir)

    missing = [path for path in files.values() if not os.path.exists(path)]
    if missing:
        return {
            'ok': False,
            'error': 'Some life dashboard source files are missing.',
            'documentsDir': docs_dir,
            'missing': missing,
        }

    inventory_data = _parse_inventory_workbook(files['inventoryWorkbook'])
    finance_data = _parse_finance_inventory_workbook(files['financeWorkbook'])
    soul_data = _parse_soul_doc(files['soulDoc'])
    journal_data = _parse_journal(files['journalDoc'])
    creative_data = _parse_creative_doc(files['creativeDoc'])
    running_data = _build_running_snapshot(docs_dir)
    profile_library = _load_profile_ingests(docs_dir)

    onebag = dict(inventory_data['inventory'])
    onebag['watchlist'] = finance_data['watchlist']
    onebag['targetWeightKg'] = finance_data['targetWeightKg']
    onebag['weightGapKg'] = (
        round(onebag['trackedWeightKg'] - finance_data['targetWeightKg'], 2)
        if finance_data['targetWeightKg'] is not None
        else None
    )

    money = dict(inventory_data['money'])
    money['investmentsTotal'] = round(money['isaTotal'] + money['taxableInvestTotal'], 2)
    try:
        market_performance = _build_market_performance(docs_dir, money)
    except Exception as exc:
        market_performance = {
            'ok': False,
            'error': str(exc),
            'positions': [],
            'series': [],
            'pointCount': 0,
            'asOfDate': '',
            'marketValueGbp': 0.0,
            'cashReserveGbp': 0.0,
            'totalWithCashGbp': 0.0,
            'marketDayChangeGbp': 0.0,
            'marketDayChangePct': 0.0,
            'totalDayChangeGbp': 0.0,
            'warnings': [],
            'rangeNote': 'Current-position market value over time. This is not a cost-basis return chart.',
            'priceSource': {
                'provider': 'Yahoo Finance chart data',
                'fxSymbol': MARKET_FX_SYMBOL,
            },
        }
    portfolio_history = []
    try:
        portfolio_history = _record_life_history(docs_dir, money, onebag, inventory_data['work'])
    except Exception:
        portfolio_history = list(_load_life_history(docs_dir).get('portfolio') or [])

    timestamps = [os.path.getmtime(path) for path in files.values()]
    if os.path.exists(running_data['storagePath']):
        timestamps.append(os.path.getmtime(running_data['storagePath']))
    timestamps.extend(path.stat().st_mtime for path in _profile_ingest_paths(docs_dir))
    snapshot = {
        'ok': True,
        'generatedAt': datetime.now().isoformat(timespec='seconds'),
        'sourceUpdatedAt': datetime.fromtimestamp(max(timestamps)).isoformat(timespec='seconds'),
        'documentsDir': docs_dir,
        'profile': inventory_data['profile'],
        'profileLibrary': profile_library,
        'work': inventory_data['work'],
        'onebag': onebag,
        'running': running_data,
        'money': money,
        'marketPerformance': market_performance,
        'portfolioHistory': portfolio_history,
        'soul': soul_data,
        'journal': journal_data,
        'people': journal_data.get('people', {}),
        'journalCapture': _journal_capture_capabilities(files['journalDoc'], _journal_capture_dir(docs_dir)),
        'creative': creative_data,
    }
    snapshot['people'] = _build_people_layer(
        journal_data.get('people', {}),
        profile_library,
        _collapse_whitespace(inventory_data.get('profile', {}).get('fullName', '')) or 'Daniel Atkinson',
    )
    snapshot['journal']['people'] = snapshot['people']
    snapshot['signals'] = _build_life_signals(snapshot)
    try:
        snapshot['aiContext'] = _build_ai_context_markdown(snapshot)
    except Exception as exc:
        snapshot['aiContext'] = {
            'directory': _ai_context_dir(docs_dir),
            'files': [],
            'error': str(exc),
        }
    if isinstance(snapshot.get('aiContext'), dict):
        snapshot['aiContext']['lastSyncedAt'] = _ai_context_sync_state.get('lastSyncedAt', '') or snapshot.get('generatedAt', '')
        snapshot['aiContext']['syncIntervalSeconds'] = AI_CONTEXT_SYNC_INTERVAL_SECONDS
        if _ai_context_sync_state.get('lastError'):
            snapshot['aiContext']['syncError'] = _ai_context_sync_state['lastError']
    return snapshot


def _sanitize_snapshot_for_hosting(snapshot):
    hosted = copy.deepcopy(snapshot)
    hosted['deploymentMode'] = 'hosted'
    hosted.pop('documentsDir', None)

    profile_library = hosted.get('profileLibrary')
    if isinstance(profile_library, dict):
        profile_library.pop('intimateRelationshipHistory', None)
        summary = profile_library.get('summary')
        if isinstance(summary, dict):
            summary.pop('intimate_relationship_count', None)

    running = hosted.get('running')
    if isinstance(running, dict):
        running.pop('storagePath', None)

    market = hosted.get('marketPerformance')
    if isinstance(market, dict):
        market.pop('cachePath', None)

    ai_context = hosted.get('aiContext') if isinstance(hosted.get('aiContext'), dict) else {}
    hosted['aiContext'] = {
        'directory': '/ai-context',
        'files': [name for name in ai_context.get('files', []) if isinstance(name, str)],
        'entrypoints': [name for name in ai_context.get('entrypoints', []) if isinstance(name, str)],
        'sources': [item for item in ai_context.get('sources', []) if isinstance(item, dict)],
        'graph': ai_context.get('graph', {}) if isinstance(ai_context.get('graph'), dict) else {},
        'site': {
            'homePath': ((ai_context.get('site') or {}).get('homePath', '') if isinstance(ai_context.get('site'), dict) else ''),
            'pages': [
                item for item in ((ai_context.get('site') or {}).get('pages', []) if isinstance(ai_context.get('site'), dict) else [])
                if isinstance(item, dict)
            ],
        },
        'extracted': {
            'directory': '/extracted',
            'files': [name for name in ((ai_context.get('extracted') or {}).get('files', []) if isinstance(ai_context.get('extracted'), dict) else []) if isinstance(name, str)],
            'entrypoints': [name for name in ((ai_context.get('extracted') or {}).get('entrypoints', []) if isinstance(ai_context.get('extracted'), dict) else []) if isinstance(name, str)],
        },
        'lastSyncedAt': ai_context.get('lastSyncedAt', '') or hosted.get('generatedAt', ''),
        'syncIntervalSeconds': ai_context.get('syncIntervalSeconds', AI_CONTEXT_SYNC_INTERVAL_SECONDS),
        'readmeUrl': '/ai-context/README.md',
    }

    hosted['journalCapture'] = {
        'canAppend': False,
        'appendReason': 'Hosted mode is read-only. Use the local dashboard on your Mac to write journal files.',
        'canCreateNoteFile': False,
        'noteDirectory': 'Local dashboard only',
        'canServerTranscribe': False,
        'transcribeReasons': [
            'Hosted mode is read-only. Use the local dashboard on your Mac to capture and transcribe notes.',
        ],
    }
    return hosted


def _render_life_config_js(config):
    payload = json.dumps(config, ensure_ascii=False, indent=2)
    return f'window.AKIBWA_LIFE_CONFIG = {payload};\n'


def _copy_hosted_asset(source_name, output_dir):
    source_path = Path(_public_dir, source_name)
    if not source_path.exists():
        return
    target_path = Path(output_dir, source_name)
    target_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source_path, target_path)


def _export_cloudflare_site(output_dir):
    snapshot = _build_life_snapshot()
    if not snapshot.get('ok'):
        raise RuntimeError(snapshot.get('error') or 'The life snapshot could not be built.')

    docs_dir = snapshot.get('documentsDir')
    if not docs_dir:
        raise RuntimeError('Akibwa Documents folder not found.')

    output_path = Path(output_dir)
    if not output_path.is_absolute():
        output_path = Path(_repo_dir, output_dir)
    output_path = output_path.resolve()
    if output_path.exists():
        shutil.rmtree(output_path)
    output_path.mkdir(parents=True, exist_ok=True)

    hosted_snapshot = _sanitize_snapshot_for_hosting(snapshot)
    hosted_config = {
        'mode': 'hosted',
        'dataUrl': '/data/life.json',
        'noteFileUrl': '',
        'transcribeUrl': '',
        'readOnly': True,
    }

    life_html = _load_life_html()
    _write_text_atomic(output_path / 'index.html', life_html)
    _write_text_atomic(output_path / 'life.html', life_html)
    _write_text_atomic(output_path / 'life-config.js', _render_life_config_js(hosted_config))
    _write_json_atomic(output_path / 'data' / 'life.json', hosted_snapshot)

    for asset_name in ('life.css', 'life.js', 'privacy.html'):
        _copy_hosted_asset(asset_name, output_path)

    ai_context_dir = Path(_ai_context_dir(docs_dir))
    exported_ai_context = output_path / 'ai-context'
    if ai_context_dir.exists():
        shutil.copytree(ai_context_dir, exported_ai_context, dirs_exist_ok=True)

    _write_text_atomic(
        output_path / '_headers',
        (
            "/*\n"
            "  X-Frame-Options: DENY\n"
            "  X-Content-Type-Options: nosniff\n"
            "  Referrer-Policy: same-origin\n"
            "  X-Robots-Tag: noindex, nofollow\n"
            "  Content-Security-Policy: default-src 'self'; script-src 'self'; style-src 'self' 'unsafe-inline'; img-src 'self' data: https:; connect-src 'self'; font-src 'self' data:; frame-ancestors 'none'; base-uri 'self'\n"
            "\n"
            "/index.html\n"
            "  Cache-Control: no-store\n"
            "\n"
            "/life.html\n"
            "  Cache-Control: no-store\n"
            "\n"
            "/data/*\n"
            "  Cache-Control: no-store\n"
            "\n"
            "/ai-context/*\n"
            "  Cache-Control: no-store\n"
        ),
    )
    _write_text_atomic(
        output_path / '_redirects',
        (
            "/ /index.html 200\n"
            "/life /life.html 200\n"
        ),
    )
    return {
        'outputDir': str(output_path),
        'snapshotPath': str(output_path / 'data' / 'life.json'),
        'aiContextDir': str(exported_ai_context),
    }


def _load_life_html():
    with open(os.path.join(_public_dir, 'life.html'), encoding='utf-8') as f:
        return f.read()


class DashboardHandler(http.server.SimpleHTTPRequestHandler):

    def __init__(self, *args, **kwargs):
        self._cache_control_sent = False
        super().__init__(*args, **kwargs)

    def send_response(self, code, message=None):
        self._cache_control_sent = False
        super().send_response(code, message)

    def send_header(self, keyword, value):
        if str(keyword).lower() == 'cache-control':
            self._cache_control_sent = True
        super().send_header(keyword, value)

    def end_headers(self):
        if not self._cache_control_sent:
            super().send_header('Cache-Control', 'no-store, no-cache, must-revalidate')
            super().send_header('Pragma', 'no-cache')
            super().send_header('Expires', '0')
            self._cache_control_sent = True
        super().end_headers()

    def _send_file(self, path, cache_control=None):
        target = Path(path)
        if not target.exists() or not target.is_file():
            self.send_error(404, 'File not found.')
            return

        body = target.read_bytes()
        self.send_response(200)
        self.send_header('Content-Type', self.guess_type(str(target)))
        self.send_header('Content-Length', str(len(body)))
        if cache_control:
            self.send_header('Cache-Control', cache_control)
        self.end_headers()
        self.wfile.write(body)

    def _send_json(self, payload, status=200):
        body = json.dumps(payload, ensure_ascii=False).encode('utf-8')
        self.send_response(status)
        self.send_header('Content-Type', 'application/json; charset=utf-8')
        self.send_header('Cache-Control', 'no-store')
        self.send_header('Content-Length', str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def _read_json_body(self):
        try:
            length = int(self.headers.get('Content-Length', '0'))
        except ValueError as exc:
            raise RuntimeError('Invalid Content-Length header.') from exc

        raw = self.rfile.read(length) if length > 0 else b''
        if not raw:
            return {}
        try:
            return json.loads(raw.decode('utf-8'))
        except json.JSONDecodeError as exc:
            raise RuntimeError('Request body must be valid JSON.') from exc

    def do_GET(self):
        parsed = urllib.parse.urlparse(self.path)
        params = urllib.parse.parse_qs(parsed.query)

        if parsed.path == '/api/life.json':
            snapshot = _build_life_snapshot()
            self._send_json(snapshot, status=200 if snapshot.get('ok') else 500)
            return

        if parsed.path == '/callback':
            self.redirect('/life.html')
            return

        if parsed.path in ('', '/'):
            self.redirect('/life.html')
            return

        if parsed.path == '/life.html':
            body = _load_life_html().encode('utf-8')
            self.send_response(200)
            self.send_header('Content-Type', 'text/html; charset=utf-8')
            self.send_header('Content-Length', str(len(body)))
            self.end_headers()
            self.wfile.write(body)
            return

        if parsed.path.startswith('/assets/'):
            asset_path = Path(_dist_dir, parsed.path.lstrip('/')).resolve()
            assets_root = (Path(_dist_dir) / 'assets').resolve()
            if assets_root in asset_path.parents and asset_path.exists():
                self._send_file(asset_path, cache_control='no-store')
                return

        super().do_GET()

    def do_POST(self):
        parsed = urllib.parse.urlparse(self.path)

        if parsed.path == '/api/journal/append':
            try:
                payload = self._read_json_body()
                docs_dir = _docs_dir()
                if not docs_dir:
                    raise RuntimeError('Akibwa Documents folder not found.')
                files = _life_files(docs_dir)
                entry = _append_journal_entry(
                    files['journalDoc'],
                    payload.get('title', ''),
                    payload.get('content', ''),
                    source=payload.get('source', 'typed'),
                )
                snapshot = _build_life_snapshot()
                self._send_json({
                    'ok': True,
                    'message': 'Journal entry saved to the main journal doc.',
                    'entry': entry,
                    'snapshot': snapshot,
                })
            except RuntimeError as exc:
                self._send_json({'ok': False, 'error': str(exc)}, status=400)
            except Exception as exc:
                self._send_json({'ok': False, 'error': str(exc)}, status=500)
            return

        if parsed.path == '/api/journal/note-file':
            try:
                payload = self._read_json_body()
                docs_dir = _docs_dir()
                if not docs_dir:
                    raise RuntimeError('Akibwa Documents folder not found.')
                entry = _create_journal_note_file(
                    docs_dir,
                    payload.get('title', ''),
                    payload.get('content', ''),
                    source=payload.get('source', 'typed'),
                )
                snapshot = _build_life_snapshot()
                self._send_json({
                    'ok': True,
                    'message': 'Journal note saved as a text file.',
                    'entry': entry,
                    'snapshot': snapshot,
                })
            except RuntimeError as exc:
                self._send_json({'ok': False, 'error': str(exc)}, status=400)
            except Exception as exc:
                self._send_json({'ok': False, 'error': str(exc)}, status=500)
            return

        if parsed.path == '/api/journal/transcribe':
            try:
                payload = self._read_json_body()
                transcript = _transcribe_audio_blob(
                    payload.get('audioBase64', ''),
                    payload.get('mimeType', ''),
                    language=payload.get('language', 'en'),
                )
                self._send_json({
                    'ok': True,
                    'transcript': transcript,
                })
            except RuntimeError as exc:
                self._send_json({'ok': False, 'error': str(exc)}, status=400)
            except Exception as exc:
                self._send_json({'ok': False, 'error': str(exc)}, status=500)
            return

        self._send_json({'ok': False, 'error': 'Unsupported endpoint.'}, status=404)

    def redirect(self, location):
        self.send_response(302)
        self.send_header('Location', location)
        self.end_headers()

    def log_message(self, fmt, *args):
        pass


def open_browser():
    import time

    time.sleep(0.9)
    webbrowser.open(f'http://127.0.0.1:{PORT}/life.html')


def _parse_args(argv=None):
    parser = argparse.ArgumentParser(description='Serve the local Akibwa dashboard or export a hosted snapshot.')
    parser.add_argument(
        '--export-cloudflare',
        metavar='DIR',
        help='Write a static hosted export for Cloudflare Pages into DIR.',
    )
    parser.add_argument(
        '--no-browser',
        action='store_true',
        help='Do not open the browser automatically when serving locally.',
    )
    return parser.parse_args(argv)


def main(argv=None):
    args = _parse_args(argv)
    os.chdir(_public_dir)

    try:
        _sync_ai_context(force=True)
    except Exception as exc:
        _ai_context_sync_state['lastError'] = str(exc)

    if args.export_cloudflare:
        result = _export_cloudflare_site(args.export_cloudflare)
        print('Hosted export ready:')
        print(f"  -> {result['outputDir']}")
        print(f"  -> {result['snapshotPath']}")
        print(f"  -> {result['aiContextDir']}")
        return

    threading.Thread(target=_ai_context_sync_loop, daemon=True).start()
    if not args.no_browser:
        threading.Thread(target=open_browser, daemon=True).start()

    print('━' * 46)
    print("  Dan's Local Dashboard")
    print(f'  ->  http://127.0.0.1:{PORT}/life.html')
    print()
    print('  Keep this window open while using')
    print('  the dashboard. Press Ctrl+C to stop.')
    print('━' * 46)

    try:
        with http.server.HTTPServer(('127.0.0.1', PORT), DashboardHandler) as httpd:
            httpd.serve_forever()
    except KeyboardInterrupt:
        print('\nDashboard stopped.')
    except OSError as e:
        if 'Address already in use' in str(e):
            print(f'\nPort {PORT} is already in use.')
            print(f'Open http://127.0.0.1:{PORT}/life.html in your browser,')
            print('or stop the other process first.')
        else:
            raise


if __name__ == '__main__':
    main()
